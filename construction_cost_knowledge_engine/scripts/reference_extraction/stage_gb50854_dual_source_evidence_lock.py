from __future__ import annotations

import csv
import hashlib
import re
import sys
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, Iterable, List, Sequence, Tuple
from xml.sax.saxutils import escape as xml_escape


STAGE_NAME = "GB50854_DUAL_SOURCE_EVIDENCE_LOCK_1"
ENGINE_DIR = "construction_cost_knowledge_engine"
RUN_DIR_REL = Path(
    "construction_cost_knowledge_engine/data/private/reference_extraction/runs/GB50854_DUAL_SOURCE_EVIDENCE_LOCK_1"
)
DOCS_REL = Path("construction_cost_knowledge_engine/docs/reference_extraction")
NATIONAL_PDF_REL = Path(
    "construction_cost_knowledge_engine/data/private/reference_extraction/source_standards/国家标准/房屋建筑与装饰工程工程量计算标准.pdf"
)
DOCX_PROXY_REL = Path(
    "construction_cost_knowledge_engine/data/private/reference_extraction/source_standards/房屋建筑与装饰工程工程量计算标准 GB_T 50854-2024.docx"
)
BASELINE_REL = Path(
    "construction_cost_knowledge_engine/data/private/reference_extraction/runs/GB50854_2024_stageB_docx_full"
)


SOURCE_ROLE_HEADERS = [
    "source_id",
    "object_type",
    "display_name",
    "actual_path",
    "file_name",
    "sha256",
    "file_size_bytes",
    "page_count",
    "paragraph_count",
    "table_count",
    "text_layer_status",
    "readable_status",
    "source_role",
    "authority_status",
    "record_count",
    "review_status",
    "remark",
]

RELATION_HEADERS = [
    "relation_id",
    "authority_document_id",
    "extraction_document_id",
    "derived_artifact_id",
    "relation_type",
    "verification_status",
    "verification_method",
    "conflict_resolution_rule",
    "review_status",
    "remark",
]

INTEGRITY_HEADERS = [
    "check_id",
    "check_name",
    "expected",
    "actual",
    "pass_fail",
    "blocking_if_fail",
    "evidence",
]

SAMPLE_HEADERS = [
    "sample_id",
    "sample_category",
    "appendix_code",
    "appendix_name",
    "bill_reference_id",
    "bill_code_9",
    "bill_name",
    "unit",
    "project_feature_raw",
    "quantity_calculation_rule",
    "work_content_raw",
    "source_heading_path",
    "source_table_index",
    "source_row_index",
    "authority_document_id",
    "authority_pdf_page_no",
    "authority_verification_status",
    "verification_method",
    "review_status",
    "remark",
]

CONFLICT_HEADERS = [
    "conflict_id",
    "conflict_type",
    "authority_document_id",
    "extraction_document_id",
    "derived_artifact_id",
    "bill_reference_id",
    "bill_code_9",
    "field_name",
    "authority_value",
    "proxy_value",
    "baseline_value",
    "conflict_status",
    "conflict_resolution_rule",
    "review_status",
    "remark",
]

BACKLOG_HEADERS = [
    "backlog_id",
    "derived_artifact_id",
    "authority_document_id",
    "bill_reference_id",
    "appendix_code",
    "appendix_name",
    "bill_code_9",
    "bill_name",
    "source_heading_path",
    "source_table_index",
    "authority_pdf_page_no",
    "authority_verification_status",
    "verification_method",
    "required_action",
    "review_status",
]

GATE_HEADERS = [
    "authority_source_status",
    "extraction_proxy_status",
    "baseline_integrity_status",
    "sample_verification_status",
    "authority_conflict_count",
    "pending_evidence_link_count",
    "approved_count",
    "final_status",
]


def project_root_from_here() -> Path:
    current = Path(__file__).resolve()
    for parent in [current, *current.parents]:
        if (parent / ENGINE_DIR).exists():
            return parent
    return current.parents[3]


def nstr(value: Any) -> str:
    return "" if value is None else str(value)


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def read_csv(path: Path) -> List[Dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as fh:
        return list(csv.DictReader(fh))


def write_csv(path: Path, headers: Sequence[str], rows: Sequence[Dict[str, Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as fh:
        writer = csv.DictWriter(fh, fieldnames=headers, extrasaction="ignore")
        writer.writeheader()
        for row in rows:
            writer.writerow({field: row.get(field, "") for field in headers})


def pdf_profile(path: Path) -> Dict[str, Any]:
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        page_count = len(reader.pages)
        samples: List[str] = []
        for page in reader.pages[: min(5, page_count)]:
            try:
                text = page.extract_text() or ""
            except Exception:
                text = ""
            if text.strip():
                samples.append(text.strip())
        text_layer_status = "text_present" if samples else "text_not_detected"
        return {
            "page_count": page_count,
            "text_layer_status": text_layer_status,
            "readable_status": "readable",
        }
    except Exception as exc:
        return {
            "page_count": "",
            "text_layer_status": f"text_layer_error:{type(exc).__name__}",
            "readable_status": f"not_readable:{type(exc).__name__}",
        }


def docx_profile(path: Path) -> Dict[str, Any]:
    try:
        from docx import Document

        document = Document(str(path))
        return {
            "paragraph_count": len(document.paragraphs),
            "table_count": len(document.tables),
            "readable_status": "readable",
        }
    except Exception as exc:
        return {
            "paragraph_count": "",
            "table_count": "",
            "readable_status": f"not_readable:{type(exc).__name__}",
        }


def append_check(
    rows: List[Dict[str, Any]],
    check_name: str,
    expected: str,
    actual: Any,
    passed: bool,
    blocking: bool,
    evidence: str,
) -> None:
    rows.append(
        {
            "check_id": f"GB50854-INT-{len(rows) + 1:03d}",
            "check_name": check_name,
            "expected": expected,
            "actual": actual,
            "pass_fail": "pass" if passed else "fail",
            "blocking_if_fail": str(blocking).lower(),
            "evidence": evidence,
        }
    )


def source_role_registry(
    project_root: Path,
    pdf_profile_row: Dict[str, Any],
    docx_profile_row: Dict[str, Any],
    baseline_summary: Dict[str, Any],
) -> List[Dict[str, Any]]:
    pdf_path = project_root / NATIONAL_PDF_REL
    docx_path = project_root / DOCX_PROXY_REL
    baseline_dir = project_root / BASELINE_REL
    return [
        {
            "source_id": "GB50854_AUTHORITY_PDF_2024",
            "object_type": "source_document",
            "display_name": "GB/T 50854-2024 authority PDF",
            "actual_path": str(pdf_path),
            "file_name": pdf_path.name,
            "sha256": sha256_file(pdf_path),
            "file_size_bytes": pdf_path.stat().st_size,
            "page_count": pdf_profile_row["page_count"],
            "paragraph_count": "",
            "table_count": "",
            "text_layer_status": pdf_profile_row["text_layer_status"],
            "readable_status": pdf_profile_row["readable_status"],
            "source_role": "authority_source",
            "authority_status": "official_standard_evidence",
            "record_count": "",
            "review_status": "locked",
            "remark": "Official PDF is the authority source. Do not OCR or mutate it in this stage.",
        },
        {
            "source_id": "GB50854_EXTRACTION_PROXY_DOCX_2024",
            "object_type": "source_document",
            "display_name": "GB/T 50854-2024 DOCX extraction proxy",
            "actual_path": str(docx_path),
            "file_name": docx_path.name,
            "sha256": sha256_file(docx_path),
            "file_size_bytes": docx_path.stat().st_size,
            "page_count": "not_applicable_docx",
            "paragraph_count": docx_profile_row["paragraph_count"],
            "table_count": docx_profile_row["table_count"],
            "text_layer_status": "docx_structured_text_present"
            if docx_profile_row["readable_status"] == "readable"
            else "docx_structured_text_unknown",
            "readable_status": docx_profile_row["readable_status"],
            "source_role": "extraction_proxy",
            "authority_status": "non_authoritative_structured_source",
            "record_count": "",
            "review_status": "locked_for_extraction_reuse",
            "remark": "Used only as structured extraction proxy; conflict resolution rule is official_pdf_wins.",
        },
        {
            "source_id": "GB50854_BASELINE_472_DERIVED_REFERENCE",
            "object_type": "derived_artifact",
            "display_name": "GB/T 50854 472 bill item baseline",
            "actual_path": str(baseline_dir),
            "file_name": baseline_dir.name,
            "sha256": "",
            "file_size_bytes": "",
            "page_count": "",
            "paragraph_count": "",
            "table_count": "",
            "text_layer_status": "",
            "readable_status": "readable",
            "source_role": "derived_reference_candidate",
            "authority_status": "derived_pending_reference",
            "record_count": baseline_summary["bill_item_count"],
            "review_status": "pending",
            "remark": "Derived from extraction proxy and governed by authority PDF evidence. Records stay pending.",
        },
    ]


def source_relationship() -> List[Dict[str, Any]]:
    return [
        {
            "relation_id": "GB50854-REL-001",
            "authority_document_id": "GB50854_AUTHORITY_PDF_2024",
            "extraction_document_id": "GB50854_EXTRACTION_PROXY_DOCX_2024",
            "derived_artifact_id": "",
            "relation_type": "verifies",
            "verification_status": "pending_manual_visual_samples",
            "verification_method": "manual_visual_page_sample",
            "conflict_resolution_rule": "official_pdf_wins",
            "review_status": "pending",
            "remark": "PDF has no detected machine text layer; this stage locks role relationship without claiming row-level text verification.",
        },
        {
            "relation_id": "GB50854-REL-002",
            "authority_document_id": "",
            "extraction_document_id": "GB50854_EXTRACTION_PROXY_DOCX_2024",
            "derived_artifact_id": "GB50854_BASELINE_472_DERIVED_REFERENCE",
            "relation_type": "derived",
            "verification_status": "baseline_integrity_passed",
            "verification_method": "docx_structural_table_extraction_integrity_check",
            "conflict_resolution_rule": "official_pdf_wins",
            "review_status": "pending",
            "remark": "472 bill items and 161 context rules are derived from the extraction proxy and remain pending reference candidates.",
        },
        {
            "relation_id": "GB50854-REL-003",
            "authority_document_id": "GB50854_AUTHORITY_PDF_2024",
            "extraction_document_id": "",
            "derived_artifact_id": "GB50854_BASELINE_472_DERIVED_REFERENCE",
            "relation_type": "authoritative_evidence_for",
            "verification_status": "pending_evidence_link_backlog",
            "verification_method": "manual_visual_page_sample",
            "conflict_resolution_rule": "official_pdf_wins",
            "review_status": "pending",
            "remark": "Official PDF is authoritative evidence for the baseline. Row-level official page links are intentionally backlogged.",
        },
    ]


def baseline_integrity(project_root: Path) -> Tuple[List[Dict[str, Any]], Dict[str, Any]]:
    base = project_root / BASELINE_REL
    bill_path = base / "bill_item_reference_all_candidate.csv"
    rule_path = base / "bill_context_rules_all.csv"
    profile_path = base / "docx_full_profile.csv"
    bills = read_csv(bill_path)
    rules = read_csv(rule_path)
    profiles = read_csv(profile_path)
    bill_codes = [row.get("bill_code_9", "") for row in bills]
    duplicate_codes = sorted({code for code in bill_codes if bill_codes.count(code) > 1})
    invalid_codes = [code for code in bill_codes if not re.fullmatch(r"\d{9}", code or "")]
    review_statuses = sorted({row.get("review_status", "") for row in bills})
    work_content_nonempty = sum(1 for row in bills if row.get("work_content_raw"))
    quantity_rule_nonempty = sum(1 for row in bills if row.get("quantity_calculation_rule"))
    source_hashes = sorted({row.get("source_file_hash", "").lower() for row in bills + rules if row.get("source_file_hash")})
    docx_hash = sha256_file(project_root / DOCX_PROXY_REL).lower()
    profile = profiles[0] if profiles else {}

    checks: List[Dict[str, Any]] = []
    append_check(checks, "bill_item_count", "472", len(bills), len(bills) == 472, True, str(bill_path))
    append_check(checks, "context_rule_count", "161", len(rules), len(rules) == 161, True, str(rule_path))
    append_check(checks, "duplicate_bill_code_9", "0", len(duplicate_codes), len(duplicate_codes) == 0, True, "; ".join(duplicate_codes))
    append_check(checks, "invalid_bill_code_9", "0", len(invalid_codes), len(invalid_codes) == 0, True, "; ".join(invalid_codes))
    append_check(checks, "review_status_all_pending", "pending", "; ".join(review_statuses), review_statuses == ["pending"], True, "bill item review_status")
    append_check(checks, "work_content_nonempty_count", "> 0", work_content_nonempty, work_content_nonempty > 0, False, "work_content_raw field")
    append_check(checks, "quantity_rule_nonempty_count", "> 0", quantity_rule_nonempty, quantity_rule_nonempty > 0, False, "quantity_calculation_rule field")
    append_check(
        checks,
        "baseline_source_hash_matches_docx_proxy",
        docx_hash,
        "; ".join(source_hashes),
        source_hashes == [docx_hash],
        True,
        "baseline source_file_hash fields",
    )
    append_check(checks, "docx_profile_candidate_count", "472", profile.get("candidate_count", ""), profile.get("candidate_count") == "472", True, str(profile_path))
    append_check(checks, "docx_profile_context_rule_count", "161", profile.get("context_rule_count", ""), profile.get("context_rule_count") == "161", True, str(profile_path))

    summary = {
        "bill_item_count": len(bills),
        "context_rule_count": len(rules),
        "duplicate_count": len(duplicate_codes),
        "invalid_count": len(invalid_codes),
        "review_statuses": review_statuses,
        "work_content_nonempty": work_content_nonempty,
        "quantity_rule_nonempty": quantity_rule_nonempty,
        "baseline_integrity_passed": all(row["pass_fail"] == "pass" for row in checks if row["blocking_if_fail"] == "true"),
        "bills": bills,
        "rules": rules,
    }
    return checks, summary


def sample_review_rows(bills: Sequence[Dict[str, str]]) -> List[Dict[str, Any]]:
    categories = {
        "A": "appendix_A",
        "B": "appendix_B",
        "C": "appendix_C",
        "D": "appendix_D",
        "E": "appendix_E",
        "F": "appendix_F",
        "G": "appendix_G",
        "H": "appendix_H",
        "J": "appendix_J",
        "K": "appendix_K",
        "L": "decoration_related_appendix_L",
        "M": "decoration_related_appendix_M",
        "N": "decoration_related_appendix_N",
        "P": "decoration_related_appendix_P",
        "Q": "decoration_related_appendix_Q",
        "R": "measure_item_appendix_R",
    }
    selected: List[Dict[str, str]] = []
    seen = set()
    for row in bills:
        appendix = row.get("appendix_code", "")
        if appendix in categories and appendix not in seen:
            selected.append(row)
            seen.add(appendix)
    output: List[Dict[str, Any]] = []
    for index, row in enumerate(selected, start=1):
        output.append(
            {
                "sample_id": f"GB50854-SAMPLE-{index:03d}",
                "sample_category": categories.get(row.get("appendix_code", ""), "appendix_sample"),
                "appendix_code": row.get("appendix_code", ""),
                "appendix_name": row.get("appendix_name", ""),
                "bill_reference_id": row.get("bill_reference_id", ""),
                "bill_code_9": row.get("bill_code_9", ""),
                "bill_name": row.get("bill_name", ""),
                "unit": row.get("unit", ""),
                "project_feature_raw": row.get("project_feature_raw", ""),
                "quantity_calculation_rule": row.get("quantity_calculation_rule", ""),
                "work_content_raw": row.get("work_content_raw", ""),
                "source_heading_path": row.get("source_heading_path", ""),
                "source_table_index": row.get("source_table_index", ""),
                "source_row_index": row.get("source_row_index", ""),
                "authority_document_id": "GB50854_AUTHORITY_PDF_2024",
                "authority_pdf_page_no": "",
                "authority_verification_status": "pending_evidence_link",
                "verification_method": "manual_visual_page_sample",
                "review_status": "pending",
                "remark": "Stratified sample selected from derived baseline; official PDF page/table visual evidence link remains to be added.",
            }
        )
    return output


def evidence_backlog_rows(bills: Sequence[Dict[str, str]]) -> List[Dict[str, Any]]:
    output: List[Dict[str, Any]] = []
    for index, row in enumerate(bills, start=1):
        output.append(
            {
                "backlog_id": f"GB50854-EVIDENCE-BACKLOG-{index:03d}",
                "derived_artifact_id": "GB50854_BASELINE_472_DERIVED_REFERENCE",
                "authority_document_id": "GB50854_AUTHORITY_PDF_2024",
                "bill_reference_id": row.get("bill_reference_id", ""),
                "appendix_code": row.get("appendix_code", ""),
                "appendix_name": row.get("appendix_name", ""),
                "bill_code_9": row.get("bill_code_9", ""),
                "bill_name": row.get("bill_name", ""),
                "source_heading_path": row.get("source_heading_path", ""),
                "source_table_index": row.get("source_table_index", ""),
                "authority_pdf_page_no": "",
                "authority_verification_status": "pending_evidence_link",
                "verification_method": "manual_visual_page_sample",
                "required_action": "Add official PDF page/table visual evidence link; do not claim automatic text verification.",
                "review_status": "pending",
            }
        )
    return output


def xml_text(value: Any) -> str:
    return xml_escape(nstr(value), {"\n": "&#10;", "\r": ""})


def excel_col(index: int) -> str:
    number = index + 1
    label = ""
    while number:
        number, remainder = divmod(number - 1, 26)
        label = chr(65 + remainder) + label
    return label


def cell_xml(row_index: int, col_index: int, value: Any, header: bool = False) -> str:
    ref = f"{excel_col(col_index)}{row_index}"
    style = ' s="1"' if header else ""
    if value is None or nstr(value) == "":
        return f'<c r="{ref}"{style}/>'
    return f'<c r="{ref}"{style} t="inlineStr"><is><t xml:space="preserve">{xml_text(value)}</t></is></c>'


def column_width(field: str, rows: Sequence[Dict[str, Any]]) -> float:
    if field in {"actual_path", "remark", "project_feature_raw", "quantity_calculation_rule", "work_content_raw", "required_action"}:
        return 58.0
    if field in {"sha256", "source_heading_path", "evidence"}:
        return 42.0
    max_len = len(field)
    for row in rows[:200]:
        max_len = max(max_len, len(nstr(row.get(field, ""))))
    return min(max(max_len + 2, 10), 34)


def sheet_xml(headers: Sequence[str], rows: Sequence[Dict[str, Any]]) -> str:
    max_row = max(len(rows) + 1, 1)
    max_col = max(len(headers), 1)
    last_ref = f"{excel_col(max_col - 1)}{max_row}"
    cols = "".join(
        f'<col min="{index}" max="{index}" width="{column_width(field, rows):.2f}" customWidth="1"/>'
        for index, field in enumerate(headers, start=1)
    )
    header_cells = "".join(cell_xml(1, index, field, header=True) for index, field in enumerate(headers))
    row_xml = [f'<row r="1" spans="1:{max_col}">{header_cells}</row>']
    for row_number, row in enumerate(rows, start=2):
        cells = "".join(cell_xml(row_number, index, row.get(field, "")) for index, field in enumerate(headers))
        row_xml.append(f'<row r="{row_number}" spans="1:{max_col}">{cells}</row>')
    return f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <dimension ref="A1:{last_ref}"/>
  <sheetViews><sheetView workbookViewId="0"><pane ySplit="1" topLeftCell="A2" activePane="bottomLeft" state="frozen"/><selection pane="bottomLeft" activeCell="A2" sqref="A2"/></sheetView></sheetViews>
  <sheetFormatPr defaultRowHeight="15"/>
  <cols>{cols}</cols>
  <sheetData>{''.join(row_xml)}</sheetData>
  <autoFilter ref="A1:{last_ref}"/>
</worksheet>'''


def styles_xml() -> str:
    return '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<styleSheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
  <fonts count="2">
    <font><sz val="11"/><color theme="1"/><name val="Calibri"/><family val="2"/></font>
    <font><b/><sz val="11"/><color rgb="FFFFFFFF"/><name val="Calibri"/><family val="2"/></font>
  </fonts>
  <fills count="3">
    <fill><patternFill patternType="none"/></fill>
    <fill><patternFill patternType="gray125"/></fill>
    <fill><patternFill patternType="solid"><fgColor rgb="FF1F4E78"/><bgColor indexed="64"/></patternFill></fill>
  </fills>
  <borders count="1"><border><left/><right/><top/><bottom/><diagonal/></border></borders>
  <cellStyleXfs count="1"><xf numFmtId="0" fontId="0" fillId="0" borderId="0"/></cellStyleXfs>
  <cellXfs count="2"><xf numFmtId="0" fontId="0" fillId="0" borderId="0" xfId="0"/><xf numFmtId="0" fontId="1" fillId="2" borderId="0" xfId="0" applyFont="1" applyFill="1"/></cellXfs>
  <cellStyles count="1"><cellStyle name="Normal" xfId="0" builtinId="0"/></cellStyles>
</styleSheet>'''


def write_xlsx(path: Path, sheets: Sequence[Tuple[str, Sequence[str], Sequence[Dict[str, Any]]]]) -> None:
    workbook_sheets = []
    workbook_rels = []
    content_overrides = []
    for index, (name, _headers, _rows) in enumerate(sheets, start=1):
        safe_name = name[:31]
        workbook_sheets.append(f'<sheet name="{xml_text(safe_name)}" sheetId="{index}" r:id="rId{index}"/>')
        workbook_rels.append(
            f'<Relationship Id="rId{index}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" Target="worksheets/sheet{index}.xml"/>'
        )
        content_overrides.append(
            f'<Override PartName="/xl/worksheets/sheet{index}.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>'
        )
    style_rel_id = len(sheets) + 1
    workbook_rels.append(
        f'<Relationship Id="rId{style_rel_id}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/>'
    )
    now = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
    workbook_xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"><sheets>{''.join(workbook_sheets)}</sheets></workbook>'''
    workbook_rels_xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">{''.join(workbook_rels)}</Relationships>'''
    rels_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/>
  <Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" Target="docProps/core.xml"/>
  <Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" Target="docProps/app.xml"/>
</Relationships>'''
    content_types = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>
  <Override PartName="/xl/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.styles+xml"/>
  <Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>
  <Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>
  {''.join(content_overrides)}
</Types>'''
    core_xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:dcterms="http://purl.org/dc/terms/" xmlns:dcmitype="http://purl.org/dc/dcmitype/" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
  <dc:title>GB50854 Dual Source Evidence Review</dc:title>
  <dc:creator>AI Construction System</dc:creator>
  <cp:lastModifiedBy>AI Construction System</cp:lastModifiedBy>
  <dcterms:created xsi:type="dcterms:W3CDTF">{now}</dcterms:created>
  <dcterms:modified xsi:type="dcterms:W3CDTF">{now}</dcterms:modified>
</cp:coreProperties>'''
    app_xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties" xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">
  <Application>AI Construction System</Application>
  <HeadingPairs><vt:vector size="2" baseType="variant"><vt:variant><vt:lpstr>Worksheets</vt:lpstr></vt:variant><vt:variant><vt:i4>{len(sheets)}</vt:i4></vt:variant></vt:vector></HeadingPairs>
  <TitlesOfParts><vt:vector size="{len(sheets)}" baseType="lpstr">{''.join(f"<vt:lpstr>{xml_text(name[:31])}</vt:lpstr>" for name, _headers, _rows in sheets)}</vt:vector></TitlesOfParts>
</Properties>'''
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", content_types)
        archive.writestr("_rels/.rels", rels_xml)
        archive.writestr("docProps/core.xml", core_xml)
        archive.writestr("docProps/app.xml", app_xml)
        archive.writestr("xl/workbook.xml", workbook_xml)
        archive.writestr("xl/_rels/workbook.xml.rels", workbook_rels_xml)
        archive.writestr("xl/styles.xml", styles_xml())
        for index, (_name, headers, rows) in enumerate(sheets, start=1):
            archive.writestr(f"xl/worksheets/sheet{index}.xml", sheet_xml(headers, rows))


def md_table(headers: Sequence[str], rows: Sequence[Dict[str, Any]]) -> str:
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join("---" for _ in headers) + " |"]
    for row in rows:
        lines.append("| " + " | ".join(nstr(row.get(header, "")).replace("\n", " ").replace("|", "\\|") for header in headers) + " |")
    return "\n".join(lines)


def write_report(
    path: Path,
    registry_rows: Sequence[Dict[str, Any]],
    relationship_rows: Sequence[Dict[str, Any]],
    integrity_rows: Sequence[Dict[str, Any]],
    sample_rows: Sequence[Dict[str, Any]],
    conflict_rows: Sequence[Dict[str, Any]],
    backlog_rows: Sequence[Dict[str, Any]],
    gate_row: Dict[str, Any],
) -> None:
    content = f"""# Stage {STAGE_NAME} Report

## Final Status

{gate_row['final_status']}

This stage locks the GB/T 50854-2024 dual-source evidence model only. It does not OCR the authority PDF, does not mutate the PDF or DOCX, does not modify the 472-row baseline, does not parse A01/A02/A03, does not execute Mapping, and does not modify Web.

## Source Roles

{md_table(['source_id', 'source_role', 'authority_status', 'sha256', 'page_count', 'paragraph_count', 'table_count', 'readable_status'], registry_rows)}

## Source Relationships

{md_table(RELATION_HEADERS, relationship_rows)}

## Baseline Integrity

{md_table(INTEGRITY_HEADERS, integrity_rows)}

## Stratified Authority Sample

- sample rows: {len(sample_rows)}
- appendix coverage: {', '.join(row['appendix_code'] for row in sample_rows)}
- verification method: manual_visual_page_sample
- sample status: pending_evidence_link

## Authority Conflicts

- conflict rows: {len(conflict_rows)}
- conflict resolution rule: official_pdf_wins

## Evidence Link Backlog

- pending evidence link rows: {len(backlog_rows)}
- reason: authority PDF text layer was not detected; row-level official PDF page/table links must be added by visual evidence review.

## Gate

{md_table(GATE_HEADERS, [gate_row])}
"""
    path.write_text(content, encoding="utf-8")


def write_governance_doc(path: Path, gate_row: Dict[str, Any]) -> None:
    content = f"""# GB50854 Dual Source Governance

Stage: `{STAGE_NAME}`

## Roles

- `GB50854_AUTHORITY_PDF_2024`: `source_role = authority_source`, `authority_status = official_standard_evidence`.
- `GB50854_EXTRACTION_PROXY_DOCX_2024`: `source_role = extraction_proxy`, `authority_status = non_authoritative_structured_source`.
- `GB50854_BASELINE_472_DERIVED_REFERENCE`: `source_role = derived_reference_candidate`, `authority_status = derived_pending_reference`.

The DOCX is a structured extraction proxy only. It must not be treated as the authority source for conflicts.

## Conflict Rule

All source relationships use:

`conflict_resolution_rule = official_pdf_wins`

When the authority PDF and extraction proxy or derived baseline conflict, the authority PDF governs. A conflict must be recorded in `gb50854_authority_conflicts.csv`; it must not be silently resolved in place.

## Evidence Link Policy

The authority PDF currently has no detected machine text layer in this stage. Therefore the stage does not claim automatic row-level text verification and does not rerun OCR. Records without an authority page/table link stay:

`authority_verification_status = pending_evidence_link`

Future work should add `official_pdf_page_no` and table/page visual evidence incrementally through the backlog. Unlinked records must not be shown as verified.

## Web Display Policy

Future Web screens should display source roles separately:

- authority PDF evidence;
- extraction proxy source;
- derived pending reference baseline.

The Web must not present the DOCX proxy as the authority source and must not hide pending evidence links.

## Supplemental Gate

`gb50854_dual_source_gate.csv` final status:

`{gate_row['final_status']}`
"""
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")


def run(project_root: Path) -> Dict[str, Path]:
    run_dir = project_root / RUN_DIR_REL
    docs_dir = project_root / DOCS_REL
    if run_dir.exists() and any(run_dir.iterdir()):
        raise FileExistsError(f"Refusing to overwrite existing run directory: {run_dir}")
    run_dir.mkdir(parents=True, exist_ok=True)

    pdf_path = project_root / NATIONAL_PDF_REL
    docx_path = project_root / DOCX_PROXY_REL
    if not pdf_path.exists():
        raise FileNotFoundError(pdf_path)
    if not docx_path.exists():
        raise FileNotFoundError(docx_path)

    pdf_prof = pdf_profile(pdf_path)
    docx_prof = docx_profile(docx_path)
    integrity_rows, baseline_summary = baseline_integrity(project_root)
    registry_rows = source_role_registry(project_root, pdf_prof, docx_prof, baseline_summary)
    relationship_rows = source_relationship()
    sample_rows = sample_review_rows(baseline_summary["bills"])
    backlog_rows = evidence_backlog_rows(baseline_summary["bills"])
    conflict_rows: List[Dict[str, Any]] = []

    authority_source_status = "pass" if pdf_prof["readable_status"] == "readable" else "fail"
    extraction_proxy_status = "pass" if docx_prof["readable_status"] == "readable" else "fail"
    baseline_integrity_status = "pass" if baseline_summary["baseline_integrity_passed"] else "fail"
    sample_verification_status = "pending_evidence_link_backlog" if backlog_rows else "pass"
    approved_count = 0
    if authority_source_status != "pass" or extraction_proxy_status != "pass" or conflict_rows:
        final_status = "blocked_gb50854_source_conflict"
    elif baseline_integrity_status != "pass":
        final_status = "blocked_gb50854_baseline_integrity_failed"
    elif backlog_rows:
        final_status = "gb50854_baseline_ready_with_evidence_backlog"
    else:
        final_status = "gb50854_dual_source_ready_for_building_parse"

    gate_row = {
        "authority_source_status": authority_source_status,
        "extraction_proxy_status": extraction_proxy_status,
        "baseline_integrity_status": baseline_integrity_status,
        "sample_verification_status": sample_verification_status,
        "authority_conflict_count": len(conflict_rows),
        "pending_evidence_link_count": len(backlog_rows),
        "approved_count": approved_count,
        "final_status": final_status,
    }

    role_csv = run_dir / "gb50854_source_role_registry.csv"
    relationship_csv = run_dir / "gb50854_source_relationship.csv"
    integrity_csv = run_dir / "gb50854_baseline_integrity_check.csv"
    sample_csv = run_dir / "gb50854_authority_sample_review.csv"
    conflict_csv = run_dir / "gb50854_authority_conflicts.csv"
    backlog_csv = run_dir / "gb50854_evidence_link_backlog.csv"
    gate_csv = run_dir / "gb50854_dual_source_gate.csv"
    workbook = run_dir / "GB50854_Dual_Source_Evidence_Review.xlsx"
    report = run_dir / "stage_gb50854_dual_source_evidence_lock_report.md"
    governance_doc = docs_dir / "GB50854_DUAL_SOURCE_GOVERNANCE.md"

    write_csv(role_csv, SOURCE_ROLE_HEADERS, registry_rows)
    write_csv(relationship_csv, RELATION_HEADERS, relationship_rows)
    write_csv(integrity_csv, INTEGRITY_HEADERS, integrity_rows)
    write_csv(sample_csv, SAMPLE_HEADERS, sample_rows)
    write_csv(conflict_csv, CONFLICT_HEADERS, conflict_rows)
    write_csv(backlog_csv, BACKLOG_HEADERS, backlog_rows)
    write_csv(gate_csv, GATE_HEADERS, [gate_row])
    write_xlsx(
        workbook,
        [
            ("source_role_registry", SOURCE_ROLE_HEADERS, registry_rows),
            ("source_relationship", RELATION_HEADERS, relationship_rows),
            ("baseline_integrity", INTEGRITY_HEADERS, integrity_rows),
            ("authority_sample_review", SAMPLE_HEADERS, sample_rows),
            ("authority_conflicts", CONFLICT_HEADERS, conflict_rows),
            ("evidence_link_backlog", BACKLOG_HEADERS, backlog_rows),
            ("dual_source_gate", GATE_HEADERS, [gate_row]),
        ],
    )
    write_report(report, registry_rows, relationship_rows, integrity_rows, sample_rows, conflict_rows, backlog_rows, gate_row)
    write_governance_doc(governance_doc, gate_row)

    return {
        "run_dir": run_dir,
        "report": report,
        "relationship_csv": relationship_csv,
        "gate_csv": gate_csv,
        "governance_doc": governance_doc,
    }


def main() -> int:
    project_root = project_root_from_here()
    outputs = run(project_root)
    print(f"{STAGE_NAME} complete")
    for key, path in outputs.items():
        print(f"{key}: {path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
