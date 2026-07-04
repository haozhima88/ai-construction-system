#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Stage B-DOCX-2 full DOCX extraction for GB/T 50854-2024.

This script extends the successful Appendix A DOCX-first parser to all
appendices in the building and decoration quantity calculation standard.
It only writes reference-candidate CSV/Markdown files. It does not write
any database, migration, pipeline, approval, internal price library, or
quota-to-bill mapping output.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import re
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P


DEFAULT_DOCX_PATH = Path(
    r"E:\workspace\01_Projects\ai-construction-system\construction_cost_knowledge_engine\data\private\reference_extraction\source_standards\房屋建筑与装饰工程工程量计算标准 GB_T 50854-2024.docx"
)

SOURCE_TYPE = "official_bill_standard_docx"
SOURCE_NAME = "房屋建筑与装饰工程工程量计算标准 GB/T 50854-2024"
REVIEW_STATUS = "pending"
BILL_EXTRACTION_METHOD = "docx_table"
BILL_CODE_POLICY = "标准表提供一至九位编码；实际工程清单十二位编码的十至十二位应按具体工程项目名称和项目特征另行编制，本阶段不生成完整十二位编码。"

EXPECTED_APPENDICES = {
    "A": "土石方工程",
    "B": "地基处理与边坡支护工程",
    "C": "桩基工程",
    "D": "砌筑工程",
    "E": "混凝土及钢筋混凝土工程",
    "F": "金属结构工程",
    "G": "木结构工程",
    "H": "门窗工程",
    "J": "屋面及防水工程",
    "K": "保温、隔热、防腐工程",
    "L": "楼地面装饰工程",
    "M": "墙、柱面装饰与隔断、幕墙工程",
    "N": "天棚工程",
    "P": "油漆、涂料、裱糊工程",
    "Q": "其他装饰工程",
    "R": "措施项目",
}

PROFILE_FIELDS = [
    "profile_id",
    "source_type",
    "source_name",
    "source_file",
    "source_file_hash",
    "file_exists",
    "file_size_bytes",
    "paragraph_count",
    "nonempty_paragraph_count",
    "table_count",
    "body_block_count",
    "expected_appendix_count",
    "detected_appendix_count",
    "bill_table_count",
    "candidate_count",
    "context_rule_count",
    "can_read_text",
    "can_read_tables",
    "contains_expected_appendices",
    "extraction_method",
    "remark",
]

REGISTRY_FIELDS = [
    "registry_id",
    "source_type",
    "source_name",
    "source_file",
    "source_file_hash",
    "source_heading_path",
    "appendix_code",
    "appendix_name",
    "section_code",
    "section_name",
    "table_code",
    "table_name",
    "table_base_code",
    "source_table_index",
    "start_paragraph_index",
    "end_paragraph_index",
    "extraction_method",
    "extraction_confidence",
    "parse_issue",
    "remark",
]

CANDIDATE_FIELDS = [
    "bill_reference_id",
    "source_type",
    "source_name",
    "source_file",
    "source_file_hash",
    "source_heading_path",
    "source_table_index",
    "source_row_index",
    "appendix_code",
    "appendix_name",
    "section_code",
    "section_name",
    "table_code",
    "table_name",
    "table_base_code",
    "bill_code_9",
    "bill_code_full_policy",
    "bill_name",
    "project_feature_raw",
    "project_feature_items_json",
    "unit",
    "quantity_calculation_rule",
    "work_content_raw",
    "work_content_items_json",
    "keywords",
    "extraction_method",
    "extraction_confidence",
    "review_status",
    "reviewer",
    "remark",
]

RULE_FIELDS = [
    "rule_id",
    "source_type",
    "source_name",
    "source_file",
    "source_file_hash",
    "source_heading_path",
    "appendix_code",
    "appendix_name",
    "rule_code",
    "rule_text",
    "related_bill_codes",
    "extraction_method",
    "extraction_confidence",
    "review_status",
    "remark",
]

ISSUE_FIELDS = [
    "issue_id",
    "source_location",
    "appendix_code",
    "section_code",
    "table_code",
    "bill_code_9",
    "issue_type",
    "issue_detail",
    "severity",
    "suggested_action",
]


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def norm(value: Any) -> str:
    if value is None:
        return ""
    text = str(value).replace("\r", "\n").replace("\u3000", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n+", "\n", text)
    return text.strip()


def compact(value: Any) -> str:
    return re.sub(r"\s+", " ", norm(value)).strip()


def has_trailing_toc_page_number(text: str) -> bool:
    return bool(re.search(r"\s+\d+$", compact(text)))


def validate_bill_code(code: str) -> bool:
    return bool(re.fullmatch(r"\d{9}", compact(code)))


def is_guangdong_quota_code(value: str) -> bool:
    return bool(re.search(r"\bA1-1-\d+(?:-\d+)?\b", compact(value)))


def table_rows(table) -> List[List[str]]:
    return [[norm(cell.text) for cell in row.cells] for row in table.rows]


def compress_adjacent_duplicates(values: Sequence[str]) -> List[str]:
    result: List[str] = []
    for value in values:
        value = norm(value)
        if result and result[-1] == value:
            continue
        result.append(value)
    return result


def table_base_code(rows: Sequence[Sequence[str]]) -> str:
    for row in rows[1:]:
        cells = compress_adjacent_duplicates(row)
        for cell in cells[:3]:
            if validate_bill_code(cell):
                return compact(cell)[:6]
    return ""


def split_numbered_items(raw: str) -> List[str]:
    text = compact(raw)
    if not text:
        return []
    pattern = r"(?:^|\s)(\d+[\.．、]\s*.*?)(?=\s+\d+[\.．、]\s*|$)"
    matches = [compact(m.group(1)) for m in re.finditer(pattern, text)]
    if matches:
        return matches
    return [text]


def keywords(*parts: str) -> str:
    text = " ".join(parts)
    tokens = [t for t in re.split(r"[\s、，；;（）()《》\"“”]+", text) if len(t) >= 2]
    seen = set()
    result = []
    for token in tokens:
        if token not in seen:
            seen.add(token)
            result.append(token)
    return ";".join(result[:12])


def parse_appendix_heading(text: str) -> Optional[Tuple[str, str]]:
    m = re.fullmatch(r"附录([A-Z])\s+(.+)", compact(text))
    if not m:
        return None
    code, name = m.groups()
    if code in EXPECTED_APPENDICES and name == EXPECTED_APPENDICES[code]:
        return code, name
    return None


def parse_section_heading(text: str) -> Optional[Tuple[str, str]]:
    text = compact(text)
    if has_trailing_toc_page_number(text):
        return None
    m = re.fullmatch(r"([A-Z]\.\d+)\s+(.+)", text)
    if not m:
        return None
    code, name = m.groups()
    if code[0] not in EXPECTED_APPENDICES:
        return None
    return code, name


def parse_bill_table_caption(text: str) -> Optional[Dict[str, str]]:
    text = compact(text)
    m = re.fullmatch(r"(表([A-Z]\.\d+\.\d+))\s*(.*?)（编码[:：]?\s*(\d{6})）", text)
    if not m:
        return None
    table_code, naked_code, table_name, base_code = m.groups()
    appendix_code = naked_code.split(".", 1)[0]
    if appendix_code not in EXPECTED_APPENDICES:
        return None
    return {
        "appendix_code": appendix_code,
        "table_code": table_code,
        "table_name": compact(table_name),
        "table_base_code": base_code,
    }


def parse_any_table_caption(text: str) -> Optional[str]:
    text = compact(text)
    if re.match(r"^表[A-Z]\.\d+\.\d+(?:-\d+)?\s+", text):
        return text
    return None


def parse_rule_heading(text: str) -> Optional[Tuple[str, str]]:
    text = compact(text)
    m = re.match(r"^([A-Z]\.\d+\.\d+)\s+(.+)", text)
    if not m:
        return None
    code, body = m.groups()
    if code[0] not in EXPECTED_APPENDICES:
        return None
    return code, body


def iter_body_blocks(doc: Document) -> List[Dict[str, Any]]:
    blocks: List[Dict[str, Any]] = []
    paragraph_index = 0
    table_index = 0
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            para = doc.paragraphs[paragraph_index]
            blocks.append(
                {
                    "kind": "paragraph",
                    "block_index": len(blocks),
                    "paragraph_index": paragraph_index,
                    "text": norm(para.text),
                }
            )
            paragraph_index += 1
        elif isinstance(child, CT_Tbl):
            table = doc.tables[table_index]
            blocks.append(
                {
                    "kind": "table",
                    "block_index": len(blocks),
                    "table_index": table_index,
                    "rows": table_rows(table),
                }
            )
            table_index += 1
    return blocks


def make_issue(
    issues: List[Dict[str, Any]],
    source_location: str,
    appendix_code: str,
    section_code: str,
    table_code: str,
    bill_code_9: str,
    issue_type: str,
    detail: str,
    severity: str,
    action: str,
) -> None:
    issues.append(
        {
            "issue_id": f"ISSUE_GB50854_FULL_{len(issues) + 1:04d}",
            "source_location": source_location,
            "appendix_code": appendix_code,
            "section_code": section_code,
            "table_code": table_code,
            "bill_code_9": bill_code_9,
            "issue_type": issue_type,
            "issue_detail": detail,
            "severity": severity,
            "suggested_action": action,
        }
    )


def row_dict(
    registry_id: str,
    docx_path: Path,
    source_hash: str,
    heading_path: str,
    appendix_code: str,
    appendix_name: str,
    section_code: str = "",
    section_name: str = "",
    table_code: str = "",
    table_name: str = "",
    table_base_code: str = "",
    source_table_index: Any = "",
    start_paragraph_index: Any = "",
    end_paragraph_index: Any = "",
    method: str = "docx_paragraph_and_table_index",
    confidence: str = "0.95",
    parse_issue: str = "",
    remark: str = "",
) -> Dict[str, Any]:
    return {
        "registry_id": registry_id,
        "source_type": SOURCE_TYPE,
        "source_name": SOURCE_NAME,
        "source_file": docx_path.name,
        "source_file_hash": source_hash,
        "source_heading_path": heading_path,
        "appendix_code": appendix_code,
        "appendix_name": appendix_name,
        "section_code": section_code,
        "section_name": section_name,
        "table_code": table_code,
        "table_name": table_name,
        "table_base_code": table_base_code,
        "source_table_index": source_table_index,
        "start_paragraph_index": start_paragraph_index,
        "end_paragraph_index": end_paragraph_index,
        "extraction_method": method,
        "extraction_confidence": confidence,
        "parse_issue": parse_issue,
        "remark": remark,
    }


def scan_structure(
    docx_path: Path,
    source_hash: str,
    doc: Document,
    issues: List[Dict[str, Any]],
) -> Tuple[List[Dict[str, Any]], Dict[int, Dict[str, Any]], List[Dict[str, Any]]]:
    blocks = iter_body_blocks(doc)
    appendix_entries: List[Dict[str, Any]] = []
    section_entries: List[Dict[str, Any]] = []
    table_entries: List[Dict[str, Any]] = []
    bill_table_by_index: Dict[int, Dict[str, Any]] = {}
    pending_caption: Optional[Dict[str, Any]] = None
    current_appendix: Optional[Dict[str, Any]] = None
    current_section: Optional[Dict[str, Any]] = None

    for block in blocks:
        if block["kind"] == "paragraph":
            text = compact(block["text"])
            if not text:
                continue
            appendix = parse_appendix_heading(text)
            if appendix:
                code, name = appendix
                current_appendix = {
                    "appendix_code": code,
                    "appendix_name": name,
                    "paragraph_index": block["paragraph_index"],
                    "block_index": block["block_index"],
                }
                current_section = None
                appendix_entries.append(current_appendix)
                pending_caption = None
                continue
            section = parse_section_heading(text)
            if section and current_appendix and section[0].startswith(current_appendix["appendix_code"] + "."):
                code, name = section
                current_section = {
                    "appendix_code": current_appendix["appendix_code"],
                    "appendix_name": current_appendix["appendix_name"],
                    "section_code": code,
                    "section_name": name,
                    "paragraph_index": block["paragraph_index"],
                    "block_index": block["block_index"],
                }
                section_entries.append(current_section)
                pending_caption = None
                continue
            caption = parse_bill_table_caption(text)
            if caption and current_appendix and current_section:
                pending_caption = {
                    **caption,
                    "paragraph_index": block["paragraph_index"],
                    "block_index": block["block_index"],
                    "section_code": current_section["section_code"],
                    "section_name": current_section["section_name"],
                    "appendix_name": current_appendix["appendix_name"],
                }
                continue
        else:
            rows = block["rows"]
            base = table_base_code(rows)
            if not base:
                continue
            table_index = block["table_index"]
            meta: Dict[str, Any]
            parse_issue = ""
            confidence = "0.96"
            if pending_caption and pending_caption["table_base_code"] == base:
                meta = dict(pending_caption)
            else:
                parse_issue = "bill_table_caption_not_immediately_matched"
                confidence = "0.78"
                appendix_code = current_appendix["appendix_code"] if current_appendix else ""
                appendix_name = current_appendix["appendix_name"] if current_appendix else ""
                section_code = current_section["section_code"] if current_section else ""
                section_name = current_section["section_name"] if current_section else ""
                meta = {
                    "appendix_code": appendix_code,
                    "appendix_name": appendix_name,
                    "section_code": section_code,
                    "section_name": section_name,
                    "table_code": "",
                    "table_name": "",
                    "table_base_code": base,
                    "paragraph_index": "",
                    "block_index": block["block_index"],
                }
                make_issue(
                    issues,
                    f"table={table_index}",
                    appendix_code,
                    section_code,
                    "",
                    "",
                    "table_not_found",
                    f"Bill table with base code {base} was detected without a matching caption.",
                    "medium",
                    "Review registry row and table caption alignment.",
                )
            heading = (
                f"附录{meta['appendix_code']} {meta['appendix_name']} > "
                f"{meta['section_code']} {meta['section_name']} > "
                f"{meta['table_code']} {meta['table_name']}（编码：{base}）"
            )
            table_meta = {
                **meta,
                "source_heading_path": heading,
                "source_table_index": table_index,
                "table_base_code": base,
            }
            bill_table_by_index[table_index] = table_meta
            table_entries.append(
                {
                    **table_meta,
                    "parse_issue": parse_issue,
                    "extraction_confidence": confidence,
                }
            )
            pending_caption = None

    appendix_start_by_code = {row["appendix_code"]: row for row in appendix_entries}
    detected = set(appendix_start_by_code)
    for code, name in EXPECTED_APPENDICES.items():
        if code not in detected:
            make_issue(
                issues,
                "appendix_scan",
                code,
                "",
                "",
                "",
                "table_not_found",
                f"Expected appendix heading not found: 附录{code} {name}.",
                "critical",
                "Verify DOCX structure before benchmark validation.",
            )

    max_para = len(doc.paragraphs) - 1
    appendix_end: Dict[str, Any] = {}
    for idx, app in enumerate(appendix_entries):
        next_start = appendix_entries[idx + 1]["paragraph_index"] if idx + 1 < len(appendix_entries) else max_para + 1
        appendix_end[app["appendix_code"]] = next_start - 1

    section_end: Dict[Tuple[str, str], Any] = {}
    sections_by_appendix: Dict[str, List[Dict[str, Any]]] = defaultdict(list)
    for section in section_entries:
        sections_by_appendix[section["appendix_code"]].append(section)
    for app_code, sections in sections_by_appendix.items():
        for idx, section in enumerate(sections):
            next_start = sections[idx + 1]["paragraph_index"] if idx + 1 < len(sections) else appendix_end.get(app_code, max_para) + 1
            section_end[(app_code, section["section_code"])] = next_start - 1

    registry: List[Dict[str, Any]] = []
    for app in appendix_entries:
        heading = f"附录{app['appendix_code']} {app['appendix_name']}"
        registry.append(
            row_dict(
                f"GB50854_{app['appendix_code']}_APPENDIX",
                docx_path,
                source_hash,
                heading,
                app["appendix_code"],
                app["appendix_name"],
                section_code=app["appendix_code"],
                section_name=app["appendix_name"],
                start_paragraph_index=app["paragraph_index"],
                end_paragraph_index=appendix_end.get(app["appendix_code"], ""),
                confidence="0.96",
                remark="appendix",
            )
        )

        for section in sections_by_appendix.get(app["appendix_code"], []):
            heading = f"附录{app['appendix_code']} {app['appendix_name']} > {section['section_code']} {section['section_name']}"
            is_context = section["section_name"] == "其他规定"
            registry.append(
                row_dict(
                    f"GB50854_{section['section_code'].replace('.', '_')}_{'RULES' if is_context else 'SECTION'}",
                    docx_path,
                    source_hash,
                    heading,
                    app["appendix_code"],
                    app["appendix_name"],
                    section_code=section["section_code"],
                    section_name=section["section_name"],
                    start_paragraph_index=section["paragraph_index"],
                    end_paragraph_index=section_end.get((app["appendix_code"], section["section_code"]), ""),
                    confidence="0.96",
                    remark="context_rules" if is_context else "section",
                )
            )

            for table in [t for t in table_entries if t["section_code"] == section["section_code"]]:
                registry.append(
                    row_dict(
                        f"GB50854_{table['table_code'].replace('表', '').replace('.', '_')}_TABLE",
                        docx_path,
                        source_hash,
                        table["source_heading_path"],
                        app["appendix_code"],
                        app["appendix_name"],
                        section_code=section["section_code"],
                        section_name=section["section_name"],
                        table_code=table["table_code"],
                        table_name=table["table_name"],
                        table_base_code=table["table_base_code"],
                        source_table_index=table["source_table_index"],
                        start_paragraph_index=table.get("paragraph_index", ""),
                        end_paragraph_index=table.get("paragraph_index", ""),
                        confidence=table["extraction_confidence"],
                        parse_issue=table["parse_issue"],
                        remark="bill_item_table",
                    )
                )

    return registry, bill_table_by_index, blocks


def normalized_bill_cells(cells: Sequence[str]) -> Tuple[List[str], str]:
    raw = [norm(cell) for cell in cells]
    compressed = compress_adjacent_duplicates(raw)
    code_pos = next((idx for idx, value in enumerate(compressed[:3]) if validate_bill_code(value)), None)
    if code_pos is None:
        return compressed, "missing_bill_code"
    if code_pos > 0:
        compressed = compressed[code_pos:]
    if len(compressed) in (4, 6):
        return compressed, ""
    if len(compressed) > 6:
        return compressed[:6], f"extra_columns_after_merge_compaction:{len(compressed)}"
    return compressed, f"unexpected_column_count:{len(compressed)}"


def extract_candidates(
    docx_path: Path,
    source_hash: str,
    doc: Document,
    bill_table_by_index: Dict[int, Dict[str, Any]],
    issues: List[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    candidates: List[Dict[str, Any]] = []
    seen_codes: Dict[str, str] = {}
    for table_index, meta in sorted(bill_table_by_index.items()):
        rows = table_rows(doc.tables[table_index])
        for row_index, cells in enumerate(rows[1:], start=1):
            compressed, shape_issue = normalized_bill_cells(cells)
            if not any(compressed):
                continue
            source_location = f"table={table_index};row={row_index}"
            if len(compressed) == 6:
                bill_code, bill_name, feature, unit, quantity_rule, work_content = compressed
                source_lacks_feature = False
                source_lacks_quantity_rule = False
            elif len(compressed) == 4:
                bill_code, bill_name, unit, work_content = compressed
                feature = ""
                quantity_rule = ""
                source_lacks_feature = True
                source_lacks_quantity_rule = True
            else:
                padded = list(compressed) + [""] * (6 - len(compressed))
                bill_code, bill_name, feature, unit, quantity_rule, work_content = padded[:6]
                source_lacks_feature = False
                source_lacks_quantity_rule = False

            bill_code = compact(bill_code)
            bill_name = norm(bill_name)
            feature = norm(feature)
            unit = norm(unit)
            quantity_rule = norm(quantity_rule)
            work_content = norm(work_content)

            if shape_issue:
                make_issue(
                    issues,
                    source_location,
                    meta["appendix_code"],
                    meta["section_code"],
                    meta["table_code"],
                    bill_code,
                    "multiline_feature_parse_issue",
                    f"Column normalization issue: {shape_issue}. Raw cell count={len(cells)}.",
                    "low",
                    "Review row shape against original Word table; raw fields were preserved where possible.",
                )

            all_text = " ".join([bill_code, bill_name, feature, unit, quantity_rule, work_content])
            if is_guangdong_quota_code(all_text):
                make_issue(
                    issues,
                    source_location,
                    meta["appendix_code"],
                    meta["section_code"],
                    meta["table_code"],
                    bill_code,
                    "A1_code_mixed_in",
                    "Detected Guangdong quota code pattern A1-1-* in candidate row.",
                    "critical",
                    "Do not use this candidate until source contamination is resolved.",
                )
            if not bill_code:
                make_issue(issues, source_location, meta["appendix_code"], meta["section_code"], meta["table_code"], bill_code, "missing_bill_code", "Missing bill code.", "high", "Verify source row before benchmark validation.")
            elif not validate_bill_code(bill_code):
                make_issue(issues, source_location, meta["appendix_code"], meta["section_code"], meta["table_code"], bill_code, "invalid_bill_code", "Bill code is not 9 digits.", "high", "Do not use row until corrected.")
            elif bill_code in seen_codes:
                make_issue(issues, source_location, meta["appendix_code"], meta["section_code"], meta["table_code"], bill_code, "duplicate_bill_code", f"Duplicate bill_code_9 also seen at {seen_codes[bill_code]}.", "high", "Resolve duplicate before downstream use.")
            else:
                seen_codes[bill_code] = source_location

            if not bill_name:
                make_issue(issues, source_location, meta["appendix_code"], meta["section_code"], meta["table_code"], bill_code, "missing_bill_name", "Missing bill name.", "high", "Verify source table.")
            if not unit:
                make_issue(issues, source_location, meta["appendix_code"], meta["section_code"], meta["table_code"], bill_code, "missing_unit", "Missing unit.", "high", "Manual QA required.")
            if not feature:
                make_issue(
                    issues,
                    source_location,
                    meta["appendix_code"],
                    meta["section_code"],
                    meta["table_code"],
                    bill_code,
                    "missing_project_feature",
                    "Project feature is blank because the source table does not provide a project feature column." if source_lacks_feature else "Missing project feature.",
                    "low" if source_lacks_feature else "medium",
                    "Confirm source-table design during manual QA." if source_lacks_feature else "Manual QA required.",
                )
            if not quantity_rule:
                make_issue(
                    issues,
                    source_location,
                    meta["appendix_code"],
                    meta["section_code"],
                    meta["table_code"],
                    bill_code,
                    "missing_quantity_rule",
                    "Quantity calculation rule is blank because the source table does not provide a quantity calculation rule column." if source_lacks_quantity_rule else "Missing quantity calculation rule.",
                    "low" if source_lacks_quantity_rule else "high",
                    "Confirm source-table design during manual QA." if source_lacks_quantity_rule else "Manual QA required.",
                )
            if not work_content:
                make_issue(issues, source_location, meta["appendix_code"], meta["section_code"], meta["table_code"], bill_code, "missing_work_content", "Missing work content.", "medium", "Manual QA required.")
            if meta["section_name"] == "其他规定":
                make_issue(issues, source_location, meta["appendix_code"], meta["section_code"], meta["table_code"], bill_code, "context_rule_mixed_into_bill_item", "Candidate row was extracted from an other-rules section.", "critical", "Remove this row from bill item candidates.")

            feature_items = split_numbered_items(feature)
            work_items = split_numbered_items(work_content)
            completeness_ok = validate_bill_code(bill_code) and bill_name and unit and work_content
            if meta["appendix_code"] != "R":
                completeness_ok = completeness_ok and feature and quantity_rule
            confidence = "0.96" if completeness_ok and not shape_issue else "0.84"
            remark = "DOCX table extraction candidate; not enterprise final standard name."
            if source_lacks_feature or source_lacks_quantity_rule:
                remark = "Source table provides only project code, project name, unit, and work content; blank feature/rule fields require manual QA."
            candidates.append(
                {
                    "bill_reference_id": f"GB50854_2024_{meta['appendix_code']}_{bill_code}",
                    "source_type": SOURCE_TYPE,
                    "source_name": SOURCE_NAME,
                    "source_file": docx_path.name,
                    "source_file_hash": source_hash,
                    "source_heading_path": meta["source_heading_path"],
                    "source_table_index": table_index,
                    "source_row_index": row_index,
                    "appendix_code": meta["appendix_code"],
                    "appendix_name": meta["appendix_name"],
                    "section_code": meta["section_code"],
                    "section_name": meta["section_name"],
                    "table_code": meta["table_code"],
                    "table_name": meta["table_name"],
                    "table_base_code": meta["table_base_code"],
                    "bill_code_9": bill_code,
                    "bill_code_full_policy": BILL_CODE_POLICY,
                    "bill_name": bill_name,
                    "project_feature_raw": feature,
                    "project_feature_items_json": json.dumps(feature_items, ensure_ascii=False),
                    "unit": unit,
                    "quantity_calculation_rule": quantity_rule,
                    "work_content_raw": work_content,
                    "work_content_items_json": json.dumps(work_items, ensure_ascii=False),
                    "keywords": keywords(bill_name, feature, work_content),
                    "extraction_method": BILL_EXTRACTION_METHOD,
                    "extraction_confidence": confidence,
                    "review_status": REVIEW_STATUS,
                    "reviewer": "",
                    "remark": remark,
                }
            )
    return candidates


def serialize_context_table(rows: Sequence[Sequence[str]]) -> str:
    lines: List[str] = []
    for row in rows:
        clean = [compact(cell) for cell in row]
        if any(clean):
            lines.append(" | ".join(clean))
    return "\n".join(lines)


def extract_context_rules(
    docx_path: Path,
    source_hash: str,
    blocks: Sequence[Dict[str, Any]],
    candidates: Sequence[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    codes_by_appendix: Dict[str, List[str]] = defaultdict(list)
    for row in candidates:
        if validate_bill_code(row["bill_code_9"]):
            codes_by_appendix[row["appendix_code"]].append(row["bill_code_9"])

    rules: List[Dict[str, Any]] = []
    current_appendix_code = ""
    current_appendix_name = ""
    current_section_code = ""
    current_section_name = ""
    current_rule: Optional[Dict[str, Any]] = None
    pending_context_caption = ""

    def flush() -> None:
        nonlocal current_rule
        if not current_rule:
            return
        rule_text = norm("\n".join(current_rule["parts"]))
        if not rule_text:
            current_rule = None
            return
        explicit_codes = sorted(set(re.findall(r"\b\d{9}\b", rule_text)))
        related = explicit_codes
        if not related and current_rule["appendix_code"] in codes_by_appendix:
            related = []
        rules.append(
            {
                "rule_id": f"GB50854_{current_rule['appendix_code']}_{current_rule['rule_code'].replace('.', '_').replace('-', '_')}",
                "source_type": SOURCE_TYPE,
                "source_name": SOURCE_NAME,
                "source_file": docx_path.name,
                "source_file_hash": source_hash,
                "source_heading_path": current_rule["source_heading_path"],
                "appendix_code": current_rule["appendix_code"],
                "appendix_name": current_rule["appendix_name"],
                "rule_code": current_rule["rule_code"],
                "rule_text": rule_text,
                "related_bill_codes": ";".join(related),
                "extraction_method": current_rule["extraction_method"],
                "extraction_confidence": current_rule["extraction_confidence"],
                "review_status": REVIEW_STATUS,
                "remark": current_rule["remark"],
            }
        )
        current_rule = None

    for block in blocks:
        if block["kind"] == "paragraph":
            text = compact(block["text"])
            if not text:
                continue
            app = parse_appendix_heading(text)
            if app:
                flush()
                current_appendix_code, current_appendix_name = app
                current_section_code = ""
                current_section_name = ""
                pending_context_caption = ""
                continue
            sec = parse_section_heading(text)
            if sec and current_appendix_code and sec[0].startswith(current_appendix_code + "."):
                flush()
                current_section_code, current_section_name = sec
                pending_context_caption = ""
                continue
            if current_section_name != "其他规定":
                continue
            table_caption = parse_any_table_caption(text)
            if table_caption:
                pending_context_caption = table_caption
                if current_rule:
                    current_rule["parts"].append(table_caption)
                continue
            parsed_rule = parse_rule_heading(text)
            if parsed_rule and parsed_rule[0].startswith(current_section_code + "."):
                flush()
                rule_code, body = parsed_rule
                current_rule = {
                    "appendix_code": current_appendix_code,
                    "appendix_name": current_appendix_name,
                    "rule_code": rule_code,
                    "parts": [body],
                    "source_heading_path": f"附录{current_appendix_code} {current_appendix_name} > {current_section_code} {current_section_name}",
                    "extraction_method": "docx_paragraph",
                    "extraction_confidence": "0.93",
                    "remark": f"source_paragraph_index={block['paragraph_index']}; context rule only; excluded from bill_item_reference candidates.",
                }
            elif current_rule:
                current_rule["parts"].append(text)
            else:
                current_rule = {
                    "appendix_code": current_appendix_code,
                    "appendix_name": current_appendix_name,
                    "rule_code": f"{current_section_code}.UNNUMBERED_{block['paragraph_index']}",
                    "parts": [text],
                    "source_heading_path": f"附录{current_appendix_code} {current_appendix_name} > {current_section_code} {current_section_name}",
                    "extraction_method": "docx_paragraph",
                    "extraction_confidence": "0.72",
                    "remark": f"source_paragraph_index={block['paragraph_index']}; unnumbered context text; excluded from bill_item_reference candidates.",
                }
        else:
            if current_section_name != "其他规定":
                continue
            rows = block["rows"]
            if table_base_code(rows):
                continue
            serialized = serialize_context_table(rows)
            if not serialized:
                continue
            caption = pending_context_caption or f"context_table_{block['table_index']}"
            if current_rule:
                current_rule["parts"].append(f"{caption}\n{serialized}")
                current_rule["extraction_method"] = "docx_paragraph_and_context_table"
                current_rule["extraction_confidence"] = "0.88"
            else:
                current_rule = {
                    "appendix_code": current_appendix_code,
                    "appendix_name": current_appendix_name,
                    "rule_code": f"{current_section_code}.TABLE_{block['table_index']}",
                    "parts": [caption, serialized],
                    "source_heading_path": f"附录{current_appendix_code} {current_appendix_name} > {current_section_code} {current_section_name}",
                    "extraction_method": "docx_context_table",
                    "extraction_confidence": "0.78",
                    "remark": f"source_table_index={block['table_index']}; context table only; excluded from bill_item_reference candidates.",
                }
            pending_context_caption = ""
    flush()
    return rules


def build_profile(
    docx_path: Path,
    source_hash: str,
    doc: Document,
    blocks: Sequence[Dict[str, Any]],
    registry: Sequence[Dict[str, Any]],
    candidates: Sequence[Dict[str, Any]],
    rules: Sequence[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    nonempty_paragraphs = [p for p in doc.paragraphs if compact(p.text)]
    detected_appendices = sorted({row["appendix_code"] for row in registry if row["remark"] == "appendix"})
    expected_found = sorted(set(EXPECTED_APPENDICES).intersection(detected_appendices))
    return [
        {
            "profile_id": "GB50854_2024_DOCX_FULL",
            "source_type": SOURCE_TYPE,
            "source_name": SOURCE_NAME,
            "source_file": docx_path.name,
            "source_file_hash": source_hash,
            "file_exists": str(docx_path.exists()).lower(),
            "file_size_bytes": docx_path.stat().st_size if docx_path.exists() else "",
            "paragraph_count": len(doc.paragraphs),
            "nonempty_paragraph_count": len(nonempty_paragraphs),
            "table_count": len(doc.tables),
            "body_block_count": len(blocks),
            "expected_appendix_count": len(EXPECTED_APPENDICES),
            "detected_appendix_count": len(detected_appendices),
            "bill_table_count": sum(1 for row in registry if row["remark"] == "bill_item_table"),
            "candidate_count": len(candidates),
            "context_rule_count": len(rules),
            "can_read_text": str(bool(nonempty_paragraphs)).lower(),
            "can_read_tables": str(len(doc.tables) > 0).lower(),
            "contains_expected_appendices": ";".join(expected_found),
            "extraction_method": "docx_first_structural_table_parser",
            "remark": "Full Appendix A-R DOCX-first extraction; no OCR and no PDF parsing used.",
        }
    ]


def write_csv(path: Path, fields: Sequence[str], rows: Sequence[Dict[str, Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8-sig") as fh:
        writer = csv.DictWriter(fh, fieldnames=fields)
        writer.writeheader()
        for row in rows:
            writer.writerow({field: row.get(field, "") for field in fields})


def count_table(counter: Counter) -> str:
    lines = ["| Item | Count |", "|---|---:|"]
    for key in sorted(counter):
        lines.append(f"| {key or '(blank)'} | {counter[key]} |")
    return "\n".join(lines)


def completeness_counts(rows: Sequence[Dict[str, Any]], fields: Sequence[str]) -> Dict[str, int]:
    return {field: sum(1 for row in rows if not compact(row.get(field, ""))) for field in fields}


def write_report(
    path: Path,
    docx_path: Path,
    profile: Sequence[Dict[str, Any]],
    registry: Sequence[Dict[str, Any]],
    candidates: Sequence[Dict[str, Any]],
    rules: Sequence[Dict[str, Any]],
    issues: Sequence[Dict[str, Any]],
) -> None:
    appendix_counts = Counter(row["appendix_code"] for row in candidates)
    rule_counts = Counter(row["appendix_code"] for row in rules)
    issue_counts = Counter(row["issue_type"] for row in issues)
    severity_counts = Counter(row["severity"] for row in issues)
    registry_appendices = sorted({row["appendix_code"] for row in registry if row["remark"] == "appendix"})
    bill_tables_by_appendix = Counter(row["appendix_code"] for row in registry if row["remark"] == "bill_item_table")
    completeness = completeness_counts(
        candidates,
        [
            "bill_code_9",
            "bill_name",
            "appendix_code",
            "section_code",
            "table_code",
            "unit",
            "project_feature_raw",
            "quantity_calculation_rule",
            "work_content_raw",
        ],
    )
    invalid_codes = [row["bill_code_9"] for row in candidates if not validate_bill_code(row["bill_code_9"])]
    duplicate_codes = [code for code, count in Counter(row["bill_code_9"] for row in candidates).items() if code and count > 1]
    quota_codes = [row["bill_code_9"] for row in candidates if is_guangdong_quota_code(" ".join(str(row.get(field, "")) for field in CANDIDATE_FIELDS))]
    non_pending = [row["bill_code_9"] for row in candidates if row["review_status"] != REVIEW_STATUS]
    context_mixed = [row for row in issues if row["issue_type"] == "context_rule_mixed_into_bill_item"]
    blank_feature_rows = [row for row in candidates if not compact(row.get("project_feature_raw", ""))]
    blank_quantity_rows = [row for row in candidates if not compact(row.get("quantity_calculation_rule", ""))]
    blank_feature_sample = "; ".join(
        f"{row['bill_code_9']} {row['appendix_code']} {row['table_code']} {row['bill_name']}"
        for row in blank_feature_rows[:8]
    ) or "none"
    blank_quantity_sample = "; ".join(
        f"{row['bill_code_9']} {row['appendix_code']} {row['table_code']} {row['bill_name']}"
        for row in blank_quantity_rows[:8]
    ) or "none"
    blocking_issues = [
        row
        for row in issues
        if row["severity"] in {"critical", "high"}
        and row["issue_type"]
        not in {"missing_project_feature", "missing_quantity_rule", "remark_empty_only"}
    ]
    go = (
        not invalid_codes
        and not duplicate_codes
        and not quota_codes
        and not non_pending
        and not context_mixed
        and not blocking_issues
        and set(EXPECTED_APPENDICES).issubset(set(registry_appendices))
    )

    lines = [
        "# Stage B-DOCX-2 Full Extraction Report - GB/T 50854-2024",
        "",
        "## 1. Task Scope",
        "",
        "Full DOCX-first extraction of GB/T 50854-2024 Appendix A-R bill item reference candidates. This run writes only reference-candidate CSV and Markdown outputs. It does not write databases, migrations, existing pipeline code, approved records, internal_price_library, quota_to_bill_mapping, PDF parses, OCR output, or A1-1-* to bill_code mappings.",
        "",
        "## 2. Input File Profile",
        "",
        f"- source_file: `{docx_path}`",
        f"- source_file_hash: `{profile[0]['source_file_hash']}`",
        f"- file_size_bytes: {profile[0]['file_size_bytes']}",
        f"- paragraph_count: {profile[0]['paragraph_count']}",
        f"- nonempty_paragraph_count: {profile[0]['nonempty_paragraph_count']}",
        f"- table_count: {profile[0]['table_count']}",
        f"- body_block_count: {profile[0]['body_block_count']}",
        f"- can_read_text: {profile[0]['can_read_text']}",
        f"- can_read_tables: {profile[0]['can_read_tables']}",
        "",
        "## 3. Parser Reuse from Appendix A",
        "",
        "- Reused the Appendix A DOCX-first rules: Word table extraction, six-column bill item mapping, 9-digit `bill_code_9` validation, raw project feature/work content preservation, `pending` review status, and separate context-rule output.",
        "- Generalized the Appendix A hardcoded table metadata into paragraph/table-order scanning for all expected appendices.",
        "- Added merged-cell normalization by collapsing adjacent duplicate cells, matching the Word-table behavior observed outside Appendix A.",
        "- Kept `其他规定` content out of `bill_item_reference_all_candidate.csv` and wrote it to `bill_context_rules_all.csv`.",
        "",
        "## 4. Appendix Registry Summary",
        "",
        f"- expected_appendix_count: {len(EXPECTED_APPENDICES)}",
        f"- detected_appendix_count: {len(registry_appendices)}",
        f"- bill_item_table_count: {sum(1 for row in registry if row['remark'] == 'bill_item_table')}",
        count_table(bill_tables_by_appendix),
        "",
        "## 5. Bill Item Extraction Summary",
        "",
        f"- total_bill_item_count: {len(candidates)}",
        count_table(appendix_counts),
        "",
        "## 6. Context Rules Summary",
        "",
        f"- total_context_rule_count: {len(rules)}",
        count_table(rule_counts),
        "",
        "## 7. Field Completeness Check",
        "",
        "| Field | Blank Count |",
        "|---|---:|",
    ]
    for field, count in completeness.items():
        lines.append(f"| {field} | {count} |")
    remark_blank = sum(1 for row in candidates if not compact(row.get("remark", "")))
    lines.extend(
        [
            f"| remark | {remark_blank} |",
            "",
            "`remark` is auxiliary only and is not treated as a blocking extraction field.",
            "",
            "## 8. Bill Code Validation",
            "",
            f"- all_bill_code_9_are_9_digits: {str(not invalid_codes).lower()}",
            f"- duplicate_bill_code_9_count: {len(duplicate_codes)}",
            f"- A1_1_quota_code_mixed_in_count: {len(quota_codes)}",
            f"- non_pending_review_status_count: {len(non_pending)}",
            f"- context_rule_mixed_into_bill_item_count: {len(context_mixed)}",
            "",
            "## 9. Issues and Risks",
            "",
            f"- issue_count: {len(issues)}",
            f"- severity_counts: {json.dumps(dict(severity_counts), ensure_ascii=False, sort_keys=True)}",
            count_table(issue_counts) if issues else "No extraction issues were generated by structural checks.",
            "",
            f"- blank_project_feature_rows: {len(blank_feature_rows)}",
            f"- blank_project_feature_sample: {blank_feature_sample}",
            f"- blank_quantity_calculation_rule_rows: {len(blank_quantity_rows)}",
            f"- blank_quantity_calculation_rule_sample: {blank_quantity_sample}",
            "",
            "Known source-structure risks: Appendix R measures table provides only project code, project name, unit, and work content, so its `project_feature_raw` and `quantity_calculation_rule` blanks are preserved as source-truth blanks. Table E.5.1 also contains three source rows with blank project feature text. These rows require manual QA, but the parser does not invent missing source content.",
            "",
            "## 10. Recommended WorkBuddy Benchmark Slices",
            "",
            "- 附录 B 地基处理与边坡支护工程",
            "- 附录 E 混凝土及钢筋混凝土工程",
            "- 附录 R 措施项目",
            "",
            "## 11. Manual QA Checklist",
            "",
            "- Verify `bill_code_9` is a 9-digit standard bill code.",
            "- Verify `bill_name` comes from the source text.",
            "- Verify `unit` is complete.",
            "- Verify `project_feature_raw` is complete where the source table provides that column.",
            "- Verify `quantity_calculation_rule` is complete where the source table provides that column.",
            "- Verify `work_content_raw` is complete.",
            "- Verify `其他规定` did not enter bill item candidates.",
            "- Verify no `A1-1-*` Guangdong quota code was mixed in.",
            "- Verify all `review_status` values are `pending`.",
            "",
            "## 12. Go / No-Go Recommendation for Benchmark Validation",
            "",
            "Go for WorkBuddy benchmark validation slices. Structural blockers were not detected; Appendix R source-column limitations require targeted manual QA." if go else "No-Go for benchmark validation until the blocking issues above are resolved.",
            "",
            f"Generated at: {datetime.now().isoformat(timespec='seconds')}",
        ]
    )
    path.write_text("\n".join(lines), encoding="utf-8")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Extract GB/T 50854-2024 full bill item references from DOCX.")
    parser.add_argument("--docx", type=Path, default=DEFAULT_DOCX_PATH)
    parser.add_argument(
        "--project-root",
        type=Path,
        default=Path(__file__).resolve().parents[2],
        help="construction_cost_knowledge_engine project root",
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    docx_path = args.docx
    project_root = args.project_root
    output_dir = project_root / "data" / "private" / "reference_extraction" / "runs" / "GB50854_2024_stageB_docx_full"
    output_dir.mkdir(parents=True, exist_ok=True)

    issues: List[Dict[str, Any]] = []
    if not docx_path.exists():
        make_issue(
            issues,
            str(docx_path),
            "",
            "",
            "",
            "",
            "table_not_found",
            "DOCX file not found at required path.",
            "critical",
            "Place the DOCX at the required project path and rerun.",
        )
        write_csv(output_dir / "bill_extraction_issues_all.csv", ISSUE_FIELDS, issues)
        raise SystemExit(f"DOCX not found: {docx_path}")

    source_hash = sha256_file(docx_path)
    doc = Document(str(docx_path))
    registry, bill_table_by_index, blocks = scan_structure(docx_path, source_hash, doc, issues)
    candidates = extract_candidates(docx_path, source_hash, doc, bill_table_by_index, issues)
    rules = extract_context_rules(docx_path, source_hash, blocks, candidates)
    profile = build_profile(docx_path, source_hash, doc, blocks, registry, candidates, rules)

    write_csv(output_dir / "docx_full_profile.csv", PROFILE_FIELDS, profile)
    write_csv(output_dir / "bill_appendix_registry_all.csv", REGISTRY_FIELDS, registry)
    write_csv(output_dir / "bill_item_reference_all_candidate.csv", CANDIDATE_FIELDS, candidates)
    write_csv(output_dir / "bill_context_rules_all.csv", RULE_FIELDS, rules)
    write_csv(output_dir / "bill_extraction_issues_all.csv", ISSUE_FIELDS, issues)
    write_report(output_dir / "stageB_docx_full_report.md", docx_path, profile, registry, candidates, rules, issues)

    print(f"profile_rows={len(profile)}")
    print(f"registry_rows={len(registry)}")
    print(f"bill_table_rows={len(bill_table_by_index)}")
    print(f"candidate_rows={len(candidates)}")
    print(f"context_rule_rows={len(rules)}")
    print(f"issue_rows={len(issues)}")
    print("candidate_appendix_counts=" + json.dumps(dict(Counter(r["appendix_code"] for r in candidates)), ensure_ascii=False, sort_keys=True))
    print("issue_counts=" + json.dumps(dict(Counter(r["issue_type"] for r in issues)), ensure_ascii=False, sort_keys=True))
    print("output_dir=" + str(output_dir))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
