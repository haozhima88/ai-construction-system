#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Stage SOURCE-BASELINE-LOCK-1.

Generate two independent, human-reviewable source baseline packages:

1. GB/T 50854-2024 full DOCX parse review package.
2. GD2018 normalized Excel full quota parse review package.

This script reads source DOCX/XLSX files directly. It does not write databases,
migrations, schemas, mappings, approvals, internal price libraries, enterprise
templates, or bill_code values back into quota references.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import re
import sys
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple

from docx import Document
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, PatternFill


SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

import stageB_docx_extract_gb50854_full as gb  # noqa: E402


PROJECT_ROOT_DEFAULT = Path(r"E:\workspace\01_Projects\ai-construction-system")
ENGINE_REL = Path("construction_cost_knowledge_engine")
RUNS_REL = ENGINE_REL / "data" / "private" / "reference_extraction" / "runs"
DOCS_REF_REL = ENGINE_REL / "docs" / "reference_extraction"
SOURCE_DOCX_REL = (
    ENGINE_REL
    / "data"
    / "private"
    / "reference_extraction"
    / "source_standards"
    / "房屋建筑与装饰工程工程量计算标准 GB_T 50854-2024.docx"
)
SOURCE_EXCEL_REL = (
    ENGINE_REL
    / "data"
    / "private"
    / "reference_extraction"
    / "source_excels"
    / "广东省房屋建筑与装饰工程综合定额2018_normalized.xlsx"
)
OUTPUT_ROOT_REL = RUNS_REL / "SOURCE_BASELINE_LOCK_1"
GB_OUTPUT_REL = OUTPUT_ROOT_REL / "GB50854_2024_full_standard_parse_review"
GD_OUTPUT_REL = OUTPUT_ROOT_REL / "GD2018_normalized_full_quota_parse_review"

STAGE_NAME = "SOURCE_BASELINE_LOCK_1"
REVIEW_STATUS = "pending"
GD_SOURCE_TYPE = "reference_excel_normalized_from_third_party"
GD_SOURCE_NAME = "广东省房屋建筑与装饰工程综合定额2018"
GD_SOURCE_TRUST_LEVEL = "L1"
GD_VERIFICATION_STATUS = "structure_checked"

GB_BILL_FIELDS = [
    "bill_reference_id",
    "source_type",
    "source_name",
    "source_file",
    "source_file_hash",
    "source_heading_path",
    "source_table_index",
    "source_row_index",
    "appendix_code",
    "appendix_name",
    "section_code",
    "section_name",
    "table_code",
    "table_name",
    "table_base_code",
    "bill_code_9",
    "bill_code_full_policy",
    "bill_name",
    "project_feature_raw",
    "project_feature_items_json",
    "unit",
    "quantity_calculation_rule",
    "work_content_raw",
    "work_content_items_json",
    "keywords",
    "extraction_method",
    "extraction_confidence",
    "review_status",
    "reviewer",
    "remark",
]

GB_RULE_FIELDS = [
    "rule_id",
    "source_type",
    "source_name",
    "source_file",
    "source_file_hash",
    "source_heading_path",
    "appendix_code",
    "appendix_name",
    "section_code",
    "section_name",
    "rule_code",
    "rule_text",
    "related_bill_codes",
    "extraction_method",
    "extraction_confidence",
    "review_status",
    "remark",
]

GB_SOURCE_BLOCK_FIELDS = [
    "block_id",
    "block_order",
    "block_type",
    "heading_path",
    "paragraph_index",
    "table_index",
    "table_row_index",
    "table_col_index",
    "text",
    "normalized_text",
    "detected_role",
    "extraction_note",
]

GB_ISSUE_FIELDS = [
    "issue_id",
    "source_location",
    "appendix_code",
    "section_code",
    "table_code",
    "bill_code_9",
    "issue_type",
    "issue_detail",
    "severity",
    "suggested_action",
]

GD_BASE_FIELDS = [
    "source_row_id",
    "source_file",
    "source_file_hash",
    "source_sheet",
    "source_excel_row",
    "source_code_raw",
    "source_code_normalized",
    "raw_name",
    "raw_spec_model",
    "raw_unit",
    "raw_main_material_factor",
    "raw_quantity",
    "raw_main_material_price",
    "raw_labor_fee",
    "raw_material_fee",
    "raw_machine_fee",
    "raw_management_fee",
    "raw_total_fee",
    "raw_main_material_total",
    "code_prefix",
    "chapter_guess",
    "section_guess",
    "is_quota_item",
    "parse_issue",
    "review_status",
    "remark",
]

GD_QUOTA_FIELDS = [
    "reference_id",
    "source_type",
    "source_name",
    "source_file",
    "source_file_hash",
    "source_sheet",
    "source_excel_row",
    "source_code",
    "raw_name",
    "quota_name_candidate",
    "quota_feature_text_candidate",
    "unit",
    "raw_spec_model",
    "raw_main_material_factor",
    "raw_quantity",
    "raw_main_material_price",
    "raw_labor_fee",
    "raw_material_fee",
    "raw_machine_fee",
    "raw_management_fee",
    "raw_total_fee",
    "raw_main_material_total",
    "code_prefix",
    "chapter_guess",
    "section_guess",
    "extraction_confidence",
    "source_trust_level",
    "verification_status",
    "review_status",
    "reviewer",
    "remark",
]

GD_PRICING_FIELDS = [
    "pricing_review_id",
    "source_code",
    "raw_name",
    "unit",
    "raw_labor_fee",
    "raw_material_fee",
    "raw_machine_fee",
    "raw_management_fee",
    "raw_total_fee",
    "has_any_pricing",
    "missing_pricing_fields",
    "review_status",
    "remark",
]

GD_SECTION_FIELDS = [
    "section_inventory_id",
    "chapter_guess",
    "section_guess",
    "code_prefix",
    "first_source_code",
    "last_source_code",
    "row_count",
    "sample_names",
    "confidence",
    "remark",
]

GD_ISSUE_FIELDS = [
    "issue_id",
    "source_row_id",
    "source_excel_row",
    "source_code_raw",
    "source_code_normalized",
    "issue_type",
    "issue_detail",
    "severity",
    "suggested_action",
]

MANIFEST_FIELDS = [
    "stage_name",
    "artifact_name",
    "expected_path",
    "exists",
    "file_size_bytes",
    "row_count",
    "sha256",
    "created_or_modified_time",
    "source_file",
    "can_regenerate",
    "backup_required",
    "backup_path",
    "status",
    "remark",
]

FIELD_ALIASES = {
    "source_code": ["source_code", "定额编号", "项目编码", "编码"],
    "raw_name": ["raw_name", "项目名称", "定额名称", "名称"],
    "raw_spec_model": ["raw_spec_model", "规格型号", "规格", "型号"],
    "raw_unit": ["raw_unit", "计量单位", "单位"],
    "raw_main_material_factor": ["raw_main_material_factor", "主材系数"],
    "raw_quantity": ["raw_quantity", "工程量", "数量"],
    "raw_main_material_price": ["raw_main_material_price", "主材单价"],
    "raw_labor_fee": ["raw_labor_fee", "人工费"],
    "raw_material_fee": ["raw_material_fee", "材料费"],
    "raw_machine_fee": ["raw_machine_fee", "机具费", "机械费"],
    "raw_management_fee": ["raw_management_fee", "管理费"],
    "raw_total_fee": ["raw_total_fee", "合计", "综合单价", "单价"],
    "raw_main_material_total": ["raw_main_material_total", "主材合价", "主材合计"],
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--project-root", type=Path, default=PROJECT_ROOT_DEFAULT)
    return parser.parse_args()


def sha256_file(path: Path, upper: bool = False) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    value = digest.hexdigest()
    return value.upper() if upper else value


def norm(value: Any) -> str:
    if value is None:
        return ""
    text = str(value).replace("\r", "\n")
    text = re.sub(r"[ \t\u3000]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def compact(value: Any) -> str:
    return re.sub(r"\s+", "", norm(value))


def safe_header(value: Any, index: int) -> str:
    cleaned = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff]+", "_", norm(value))
    cleaned = cleaned.strip("_") or f"blank_{index:02d}"
    return f"raw_extra_col{index:02d}_{cleaned}"


def normalize_source_code(value: Any) -> str:
    text = norm(value).upper()
    for char in ["－", "—", "–", "﹣", "−"]:
        text = text.replace(char, "-")
    text = re.sub(r"\s+", "", text)
    return text


def is_quota_code(value: Any) -> bool:
    return bool(re.fullmatch(r"[A-Z]\d*-\d+(?:-\d+){1,2}", normalize_source_code(value)))


def is_supplemental_code(value: str) -> bool:
    return bool(re.fullmatch(r"[A-Z]\d*-\d+-\d+-\d+", normalize_source_code(value)))


def code_prefix(value: str) -> str:
    parts = normalize_source_code(value).split("-")
    return "-".join(parts[:2]) if len(parts) >= 2 else normalize_source_code(value)


def chapter_guess(value: str) -> str:
    return normalize_source_code(value).split("-", 1)[0] if value else ""


def natural_code_key(value: str) -> Tuple[Any, ...]:
    parts = re.split(r"(\d+)", normalize_source_code(value))
    key: List[Any] = []
    for part in parts:
        if part.isdigit():
            key.append(int(part))
        else:
            key.append(part)
    return tuple(key)


def light_name(value: str) -> str:
    return re.sub(r"\s+", " ", norm(value))


def write_csv(path: Path, fields: Sequence[str], rows: Iterable[Dict[str, Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(fields), extrasaction="ignore")
        writer.writeheader()
        for row in rows:
            writer.writerow({field: row.get(field, "") for field in fields})


def read_csv(path: Path) -> List[Dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def csv_row_count(path: Path) -> int:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return sum(1 for _ in csv.DictReader(handle))


def rel(path: Path, root: Path) -> str:
    return path.relative_to(root).as_posix()


def worksheet_safe(value: Any) -> Any:
    if isinstance(value, str) and len(value) > 32767:
        return value[:32740] + "\n...[TRUNCATED_FOR_XLSX_CELL_LIMIT]"
    return value


def write_xlsx(path: Path, sheet_specs: Sequence[Tuple[str, Sequence[str], Sequence[Dict[str, Any]]]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    workbook = Workbook()
    default = workbook.active
    workbook.remove(default)
    header_fill = PatternFill(fill_type="solid", fgColor="D9EAF7")
    for sheet_name, fields, rows in sheet_specs:
        ws = workbook.create_sheet(title=sheet_name[:31])
        ws.append(list(fields))
        for cell in ws[1]:
            cell.font = Font(bold=True)
            cell.fill = header_fill
            cell.alignment = Alignment(vertical="top", wrap_text=True)
        for row in rows:
            ws.append([worksheet_safe(row.get(field, "")) for field in fields])
        ws.freeze_panes = "A2"
        for idx, field in enumerate(fields, start=1):
            width = min(max(len(field) + 2, 12), 42)
            ws.column_dimensions[ws.cell(row=1, column=idx).column_letter].width = width
        for row in ws.iter_rows(min_row=2):
            for cell in row:
                cell.alignment = Alignment(vertical="top", wrap_text=True)
    workbook.save(path)


def extract_section_from_heading(heading_path: str) -> Tuple[str, str]:
    for part in reversed([p.strip() for p in heading_path.split(">")]):
        match = re.match(r"([A-Z]\.\d+)\s+(.+)$", part)
        if match:
            return match.group(1), match.group(2)
    return "", ""


def augment_rules(rules: Sequence[Dict[str, Any]]) -> List[Dict[str, Any]]:
    output: List[Dict[str, Any]] = []
    for row in rules:
        section_code, section_name = extract_section_from_heading(row.get("source_heading_path", ""))
        output.append({**row, "section_code": section_code, "section_name": section_name})
    return output


def detect_paragraph_role(text: str, current_section_name: str) -> Tuple[str, str]:
    clean = compact(text)
    if not clean:
        return "empty", "other"
    if gb.parse_appendix_heading(clean):
        return "heading", "appendix_heading"
    if gb.parse_section_heading(clean):
        return "heading", "section_heading"
    if gb.parse_bill_table_caption(clean) or gb.parse_any_table_caption(clean):
        return "heading", "table_title"
    if current_section_name == "其他规定" or gb.parse_rule_heading(clean):
        return "paragraph", "context_rule"
    if clean.startswith("注") or clean.startswith("注："):
        return "paragraph", "note"
    if "术语" in clean or "定义" in clean:
        return "paragraph", "term_definition"
    return "paragraph", "general_clause"


def build_source_blocks(blocks: Sequence[Dict[str, Any]], bill_table_by_index: Dict[int, Dict[str, Any]]) -> List[Dict[str, Any]]:
    output: List[Dict[str, Any]] = []
    block_order = 0
    current_appendix_code = ""
    current_appendix_name = ""
    current_section_code = ""
    current_section_name = ""
    heading_path = ""

    def add_row(
        block_type: str,
        text: str,
        detected_role: str,
        note: str,
        paragraph_index: Any = "",
        table_index: Any = "",
        table_row_index: Any = "",
        table_col_index: Any = "",
        heading: Optional[str] = None,
    ) -> None:
        nonlocal block_order
        block_order += 1
        output.append(
            {
                "block_id": f"GB50854_BLOCK_{block_order:06d}",
                "block_order": block_order,
                "block_type": block_type,
                "heading_path": heading if heading is not None else heading_path,
                "paragraph_index": paragraph_index,
                "table_index": table_index,
                "table_row_index": table_row_index,
                "table_col_index": table_col_index,
                "text": text,
                "normalized_text": compact(text),
                "detected_role": detected_role,
                "extraction_note": note,
            }
        )

    for block in blocks:
        if block["kind"] == "paragraph":
            text = norm(block.get("text"))
            clean = compact(text)
            app = gb.parse_appendix_heading(clean)
            if app:
                current_appendix_code, current_appendix_name = app
                current_section_code = ""
                current_section_name = ""
                heading_path = f"附录{current_appendix_code} {current_appendix_name}"
            else:
                sec = gb.parse_section_heading(clean)
                if sec and current_appendix_code and sec[0].startswith(current_appendix_code + "."):
                    current_section_code, current_section_name = sec
                    heading_path = f"附录{current_appendix_code} {current_appendix_name} > {current_section_code} {current_section_name}"
                else:
                    caption = gb.parse_bill_table_caption(clean)
                    if caption:
                        heading_path = (
                            f"附录{current_appendix_code} {current_appendix_name} > "
                            f"{current_section_code} {current_section_name} > "
                            f"{caption['table_code']} {caption['table_name']}（编码：{caption['table_base_code']}）"
                        )
            block_type, role = detect_paragraph_role(text, current_section_name)
            add_row(block_type, text, role, "docx_body_paragraph", paragraph_index=block.get("paragraph_index", ""))
        else:
            table_index = block.get("table_index", "")
            rows = block.get("rows", [])
            table_meta = bill_table_by_index.get(table_index)
            table_heading = table_meta.get("source_heading_path") if table_meta else heading_path
            table_role = "bill_item_table" if table_meta else "context_rule" if current_section_name == "其他规定" else "other"
            for row_idx, row_values in enumerate(rows):
                row_text = " | ".join(norm(cell) for cell in row_values)
                add_row(
                    "table_row",
                    row_text,
                    table_role,
                    "docx_table_row",
                    table_index=table_index,
                    table_row_index=row_idx,
                    heading=table_heading,
                )
                for col_idx, cell_text in enumerate(row_values):
                    text = norm(cell_text)
                    add_row(
                        "table_cell" if text else "empty",
                        text,
                        table_role,
                        "docx_table_cell",
                        table_index=table_index,
                        table_row_index=row_idx,
                        table_col_index=col_idx,
                        heading=table_heading,
                    )
    return output


def add_gb_baseline_issues(
    issues: List[Dict[str, Any]],
    candidates: Sequence[Dict[str, Any]],
    rules: Sequence[Dict[str, Any]],
) -> None:
    seen = Counter(row.get("bill_code_9", "") for row in candidates)
    existing_keys = {(row.get("bill_code_9", ""), row.get("issue_type", "")) for row in issues}

    def maybe_issue(row: Dict[str, Any], issue_type: str, detail: str, severity: str, action: str) -> None:
        key = (row.get("bill_code_9", ""), issue_type)
        if key in existing_keys:
            return
        gb.make_issue(
            issues,
            "baseline_validation",
            row.get("appendix_code", ""),
            row.get("section_code", ""),
            row.get("table_code", ""),
            row.get("bill_code_9", ""),
            issue_type,
            detail,
            severity,
            action,
        )
        existing_keys.add(key)

    for row in candidates:
        code = row.get("bill_code_9", "")
        if not re.fullmatch(r"\d{9}", code or ""):
            maybe_issue(row, "invalid_bill_code", "bill_code_9 is not 9 digits.", "critical", "Fix parser or source before baseline confirmation.")
        if seen[code] > 1:
            maybe_issue(row, "duplicate_bill_code", "Duplicate bill_code_9 in baseline candidates.", "critical", "Resolve duplicate before downstream mapping.")
        if not row.get("quantity_calculation_rule"):
            maybe_issue(row, "missing_quantity_rule", "Quantity calculation rule is blank in parsed source row.", "medium", "Verify against original standard; preserve source blanks only if table truly lacks the column.")
        if not row.get("work_content_raw"):
            maybe_issue(row, "missing_work_content", "Work content is blank in parsed source row.", "high", "Review original DOCX table before source lock.")
        if row.get("review_status") != REVIEW_STATUS:
            maybe_issue(row, "non_pending_review_status", "Review status is not pending.", "critical", "Baseline must not approve rows.")
    for row in rules:
        if not row.get("rule_text"):
            gb.make_issue(issues, "baseline_validation", row.get("appendix_code", ""), row.get("section_code", ""), "", "", "missing_context_rule_text", "Context rule text is blank.", "medium", "Verify context rule extraction.")


def build_gb_package(project_root: Path) -> Dict[str, Any]:
    docx_path = project_root / SOURCE_DOCX_REL
    if not docx_path.exists():
        raise SystemExit(f"blocked_missing_inputs: {docx_path}")
    source_hash = sha256_file(docx_path, upper=True)
    doc = Document(str(docx_path))
    issues: List[Dict[str, Any]] = []
    registry, bill_table_by_index, blocks = gb.scan_structure(docx_path, source_hash, doc, issues)
    candidates = gb.extract_candidates(docx_path, source_hash, doc, bill_table_by_index, issues)
    rules = augment_rules(gb.extract_context_rules(docx_path, source_hash, blocks, candidates))
    source_blocks = build_source_blocks(blocks, bill_table_by_index)
    add_gb_baseline_issues(issues, candidates, rules)

    output_dir = project_root / GB_OUTPUT_REL
    output_dir.mkdir(parents=True, exist_ok=True)
    write_csv(output_dir / "gb50854_bill_items_full_review.csv", GB_BILL_FIELDS, candidates)
    write_csv(output_dir / "gb50854_context_rules_full_review.csv", GB_RULE_FIELDS, rules)
    write_csv(output_dir / "gb50854_appendix_table_registry_full.csv", gb.REGISTRY_FIELDS, registry)
    write_csv(output_dir / "gb50854_source_blocks_full.csv", GB_SOURCE_BLOCK_FIELDS, source_blocks)
    write_csv(output_dir / "gb50854_full_parse_issues.csv", GB_ISSUE_FIELDS, issues)

    summary_rows = build_gb_summary_rows(docx_path, source_hash, doc, registry, candidates, rules, source_blocks, issues)
    try:
        write_xlsx(
            output_dir / "GB50854_2024_full_standard_parse_review.xlsx",
            [
                ("bill_items_full", GB_BILL_FIELDS, candidates),
                ("context_rules_full", GB_RULE_FIELDS, rules),
                ("appendix_table_registry", gb.REGISTRY_FIELDS, registry),
                ("source_blocks_full", GB_SOURCE_BLOCK_FIELDS, source_blocks),
                ("parse_issues", GB_ISSUE_FIELDS, issues),
                ("summary", ["metric_name", "metric_value", "remark"], summary_rows),
            ],
        )
    except Exception as exc:  # pragma: no cover - escalates as stage blocker.
        raise SystemExit(f"blocked_xlsx_generation_failed: GB50854 workbook generation failed: {exc}") from exc

    write_gb_report(output_dir / "gb50854_full_standard_parse_review_report.md", docx_path, source_hash, doc, registry, candidates, rules, source_blocks, issues)
    return {
        "docx_path": docx_path,
        "source_hash": source_hash,
        "output_dir": output_dir,
        "registry": registry,
        "candidates": candidates,
        "rules": rules,
        "source_blocks": source_blocks,
        "issues": issues,
        "xlsx_path": output_dir / "GB50854_2024_full_standard_parse_review.xlsx",
    }


def build_gb_summary_rows(
    docx_path: Path,
    source_hash: str,
    doc: Document,
    registry: Sequence[Dict[str, Any]],
    candidates: Sequence[Dict[str, Any]],
    rules: Sequence[Dict[str, Any]],
    source_blocks: Sequence[Dict[str, Any]],
    issues: Sequence[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    return [
        {"metric_name": "source_file", "metric_value": docx_path.name, "remark": ""},
        {"metric_name": "source_file_hash", "metric_value": source_hash, "remark": "SHA256"},
        {"metric_name": "file_size_bytes", "metric_value": docx_path.stat().st_size, "remark": ""},
        {"metric_name": "paragraph_count", "metric_value": len(doc.paragraphs), "remark": ""},
        {"metric_name": "table_count", "metric_value": len(doc.tables), "remark": ""},
        {"metric_name": "registry_rows", "metric_value": len(registry), "remark": ""},
        {"metric_name": "bill_item_rows", "metric_value": len(candidates), "remark": ""},
        {"metric_name": "context_rule_rows", "metric_value": len(rules), "remark": ""},
        {"metric_name": "source_block_rows", "metric_value": len(source_blocks), "remark": ""},
        {"metric_name": "parse_issue_rows", "metric_value": len(issues), "remark": ""},
        {"metric_name": "missing_quantity_calculation_rule", "metric_value": sum(1 for row in candidates if not row.get("quantity_calculation_rule")), "remark": "Preserves source blanks; requires manual review."},
        {"metric_name": "missing_work_content", "metric_value": sum(1 for row in candidates if not row.get("work_content_raw")), "remark": ""},
        {"metric_name": "non_pending_review_status", "metric_value": sum(1 for row in candidates if row.get("review_status") != REVIEW_STATUS), "remark": ""},
        {"metric_name": "approved_rows", "metric_value": sum(1 for row in candidates if row.get("review_status") == "approved"), "remark": ""},
    ]


def write_gb_report(
    path: Path,
    docx_path: Path,
    source_hash: str,
    doc: Document,
    registry: Sequence[Dict[str, Any]],
    candidates: Sequence[Dict[str, Any]],
    rules: Sequence[Dict[str, Any]],
    source_blocks: Sequence[Dict[str, Any]],
    issues: Sequence[Dict[str, Any]],
) -> None:
    invalid_codes = sum(1 for row in candidates if not re.fullmatch(r"\d{9}", row.get("bill_code_9", "") or ""))
    duplicate_codes = sum(1 for _, count in Counter(row.get("bill_code_9", "") for row in candidates).items() if count > 1)
    missing_quantity = sum(1 for row in candidates if not row.get("quantity_calculation_rule"))
    missing_work = sum(1 for row in candidates if not row.get("work_content_raw"))
    issue_counts = Counter(row.get("issue_type", "") for row in issues)
    lines = [
        "# GB50854 Full Standard Parse Review Report",
        "",
        "## 1. Task Scope",
        "",
        "Read the GB/T 50854-2024 DOCX source directly and generate an independent full-standard parse review package. This package preserves bill item text, quantity calculation rules, work content, context rules, and source block structure for manual confirmation.",
        "",
        "## 2. Source File Profile",
        "",
        f"- source_file: `{docx_path}`",
        f"- source_file_hash: `{source_hash}`",
        f"- file_size_bytes: {docx_path.stat().st_size}",
        f"- paragraph_count: {len(doc.paragraphs)}",
        f"- table_count: {len(doc.tables)}",
        "",
        "## 3. Full Document Block Coverage",
        "",
        f"- source_block_rows: {len(source_blocks)}",
        f"- heading_rows: {sum(1 for row in source_blocks if row.get('block_type') == 'heading')}",
        f"- paragraph_rows: {sum(1 for row in source_blocks if row.get('block_type') == 'paragraph')}",
        f"- table_row_rows: {sum(1 for row in source_blocks if row.get('block_type') == 'table_row')}",
        f"- table_cell_or_empty_rows: {sum(1 for row in source_blocks if row.get('block_type') in {'table_cell', 'empty'})}",
        "",
        "## 4. Bill Item Extraction Summary",
        "",
        f"- bill_item_rows: {len(candidates)}",
        f"- invalid_bill_code_9: {invalid_codes}",
        f"- duplicate_bill_code_9: {duplicate_codes}",
        f"- review_status_pending_rows: {sum(1 for row in candidates if row.get('review_status') == REVIEW_STATUS)}",
        "",
        "## 5. Context Rules Summary",
        "",
        f"- context_rule_rows: {len(rules)}",
        f"- context_rules_with_text: {sum(1 for row in rules if row.get('rule_text'))}",
        "",
        "## 6. Quantity Calculation Rule Completeness",
        "",
        f"- missing_quantity_calculation_rule_rows: {missing_quantity}",
        "Rows with blanks are preserved as source-text blanks and must be manually confirmed. The parser does not invent missing quantity rules.",
        "",
        "## 7. Work Content Completeness",
        "",
        f"- missing_work_content_rows: {missing_work}",
        "Work content is retained from source table cells and is never intentionally omitted.",
        "",
        "## 8. Parse Issues",
        "",
        f"- parse_issue_rows: {len(issues)}",
        f"- issue_type_counts: {json.dumps(dict(issue_counts), ensure_ascii=False, sort_keys=True)}",
        "",
        "## 9. Not Approved / Not Final Statement",
        "",
        "All rows remain `pending`. This stage does not generate approved records, enterprise standard names, internal price library data, database rows, or quota-to-bill mapping.",
        "",
        "## 10. Human Review Guidance",
        "",
        "- Confirm every `bill_code_9`, `bill_name`, `project_feature_raw`, `unit`, `quantity_calculation_rule`, and `work_content_raw` against the DOCX source.",
        "- Pay special attention to rows where the source table lacks quantity calculation rule or project feature columns.",
        "- Review `gb50854_source_blocks_full.csv` when auditors need to trace a row back to document structure.",
        "",
    ]
    path.write_text("\n".join(lines), encoding="utf-8")


def header_key(value: Any) -> str:
    return re.sub(r"\s+", "", norm(value)).lower()


def build_column_map(headers: Sequence[str]) -> Dict[str, int]:
    alias_lookup: Dict[str, str] = {}
    for normalized_name, aliases in FIELD_ALIASES.items():
        for alias in aliases:
            alias_lookup[header_key(alias)] = normalized_name
    result: Dict[str, int] = {}
    for idx, header in enumerate(headers):
        name = alias_lookup.get(header_key(header))
        if name and name not in result:
            result[name] = idx
    return result


def cell_at(row: Sequence[Any], idx: Optional[int]) -> Any:
    if idx is None or idx >= len(row):
        return ""
    return row[idx]


def issue_text(parts: Sequence[str]) -> str:
    return ";".join(part for part in parts if part)


def gd_parse_issue_for(row: Dict[str, Any]) -> str:
    issues: List[str] = []
    code = row.get("source_code_normalized", "")
    raw_name = row.get("raw_name", "")
    raw_unit = row.get("raw_unit", "")
    is_quota = row.get("is_quota_item") == "true"
    if not code:
        issues.append("blank_source_code")
    elif not is_quota:
        issues.append("non_quota_or_invalid_source_code")
    if is_quota and not raw_name:
        issues.append("missing_name")
    if is_quota and not raw_unit:
        issues.append("missing_unit")
    if is_quota and is_supplemental_code(code):
        issues.append("supplemental_source_code")
    return issue_text(issues)


def build_gd_all_rows(excel_path: Path, excel_hash: str, ws: Any, headers: Sequence[str]) -> Tuple[List[Dict[str, Any]], List[str]]:
    column_map = build_column_map(headers)
    extra_fields = [safe_header(header, idx + 1) for idx, header in enumerate(headers)]
    rows: List[Dict[str, Any]] = []
    for excel_row, values in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
        if not any(value is not None and norm(value) != "" for value in values):
            continue
        raw_by_field = {field: cell_at(values, column_map.get(field)) for field in FIELD_ALIASES}
        raw_code = norm(raw_by_field.get("source_code"))
        code = normalize_source_code(raw_code)
        raw_name = norm(raw_by_field.get("raw_name"))
        valid_quota = is_quota_code(code) and raw_name and "小计" not in raw_name and "说明" not in raw_name
        row: Dict[str, Any] = {
            "source_row_id": f"GD2018_NORMALIZED_ROW_{len(rows) + 1:06d}",
            "source_file": excel_path.name,
            "source_file_hash": excel_hash,
            "source_sheet": ws.title,
            "source_excel_row": excel_row,
            "source_code_raw": raw_code,
            "source_code_normalized": code,
            "raw_name": raw_name,
            "raw_spec_model": norm(raw_by_field.get("raw_spec_model")),
            "raw_unit": norm(raw_by_field.get("raw_unit")),
            "raw_main_material_factor": raw_by_field.get("raw_main_material_factor", ""),
            "raw_quantity": raw_by_field.get("raw_quantity", ""),
            "raw_main_material_price": raw_by_field.get("raw_main_material_price", ""),
            "raw_labor_fee": raw_by_field.get("raw_labor_fee", ""),
            "raw_material_fee": raw_by_field.get("raw_material_fee", ""),
            "raw_machine_fee": raw_by_field.get("raw_machine_fee", ""),
            "raw_management_fee": raw_by_field.get("raw_management_fee", ""),
            "raw_total_fee": raw_by_field.get("raw_total_fee", ""),
            "raw_main_material_total": raw_by_field.get("raw_main_material_total", ""),
            "code_prefix": code_prefix(code) if code else "",
            "chapter_guess": chapter_guess(code) if code else "",
            "section_guess": code_prefix(code) if code else "",
            "is_quota_item": str(valid_quota).lower(),
            "review_status": REVIEW_STATUS,
            "remark": "normalized_excel_source_row;not_official_truth;not_mapping;not_enterprise_standard_name",
        }
        for idx, field in enumerate(extra_fields):
            row[field] = cell_at(values, idx)
        row["parse_issue"] = gd_parse_issue_for(row)
        rows.append(row)
    return rows, extra_fields


def build_gd_quota_rows(all_rows: Sequence[Dict[str, Any]]) -> List[Dict[str, Any]]:
    quota_rows = [row for row in all_rows if row.get("is_quota_item") == "true"]
    quota_rows.sort(key=lambda row: natural_code_key(row.get("source_code_normalized", "")))
    output: List[Dict[str, Any]] = []
    for row in quota_rows:
        code = row.get("source_code_normalized", "")
        output.append(
            {
                "reference_id": f"GD2018_NORMALIZED_QUOTA_{code}",
                "source_type": GD_SOURCE_TYPE,
                "source_name": GD_SOURCE_NAME,
                "source_file": row.get("source_file", ""),
                "source_file_hash": row.get("source_file_hash", ""),
                "source_sheet": row.get("source_sheet", ""),
                "source_excel_row": row.get("source_excel_row", ""),
                "source_code": code,
                "raw_name": row.get("raw_name", ""),
                "quota_name_candidate": light_name(row.get("raw_name", "")),
                "quota_feature_text_candidate": light_name(row.get("raw_spec_model", "") or row.get("raw_name", "")),
                "unit": row.get("raw_unit", ""),
                "raw_spec_model": row.get("raw_spec_model", ""),
                "raw_main_material_factor": row.get("raw_main_material_factor", ""),
                "raw_quantity": row.get("raw_quantity", ""),
                "raw_main_material_price": row.get("raw_main_material_price", ""),
                "raw_labor_fee": row.get("raw_labor_fee", ""),
                "raw_material_fee": row.get("raw_material_fee", ""),
                "raw_machine_fee": row.get("raw_machine_fee", ""),
                "raw_management_fee": row.get("raw_management_fee", ""),
                "raw_total_fee": row.get("raw_total_fee", ""),
                "raw_main_material_total": row.get("raw_main_material_total", ""),
                "code_prefix": row.get("code_prefix", ""),
                "chapter_guess": row.get("chapter_guess", ""),
                "section_guess": row.get("section_guess", ""),
                "extraction_confidence": "0.84" if not is_supplemental_code(code) else "0.78",
                "source_trust_level": GD_SOURCE_TRUST_LEVEL,
                "verification_status": GD_VERIFICATION_STATUS,
                "review_status": REVIEW_STATUS,
                "reviewer": "",
                "remark": "quota candidate description only;not_enterprise_standard_name;not_approved;no_bill_code_generated",
            }
        )
    return output


def build_gd_pricing_rows(quota_rows: Sequence[Dict[str, Any]]) -> List[Dict[str, Any]]:
    pricing_fields = ["raw_labor_fee", "raw_material_fee", "raw_machine_fee", "raw_management_fee", "raw_total_fee"]
    rows: List[Dict[str, Any]] = []
    for row in quota_rows:
        missing = [field for field in pricing_fields if row.get(field) in ("", None)]
        has_any = any(row.get(field) not in ("", None) for field in pricing_fields)
        rows.append(
            {
                "pricing_review_id": f"GD2018_PRICING_{row['source_code']}",
                "source_code": row["source_code"],
                "raw_name": row["raw_name"],
                "unit": row["unit"],
                "raw_labor_fee": row.get("raw_labor_fee", ""),
                "raw_material_fee": row.get("raw_material_fee", ""),
                "raw_machine_fee": row.get("raw_machine_fee", ""),
                "raw_management_fee": row.get("raw_management_fee", ""),
                "raw_total_fee": row.get("raw_total_fee", ""),
                "has_any_pricing": str(has_any).lower(),
                "missing_pricing_fields": ";".join(missing),
                "review_status": REVIEW_STATUS,
                "remark": "pricing fields preserved from normalized Excel;do_not_generate_internal_price_library",
            }
        )
    return rows


def build_gd_section_inventory(quota_rows: Sequence[Dict[str, Any]]) -> List[Dict[str, Any]]:
    grouped: Dict[str, List[Dict[str, Any]]] = defaultdict(list)
    for row in quota_rows:
        grouped[row.get("code_prefix", "")].append(row)
    rows: List[Dict[str, Any]] = []
    for idx, prefix in enumerate(sorted(grouped, key=natural_code_key), start=1):
        group = sorted(grouped[prefix], key=lambda row: natural_code_key(row.get("source_code", "")))
        rows.append(
            {
                "section_inventory_id": f"GD2018_SECTION_{idx:04d}",
                "chapter_guess": group[0].get("chapter_guess", ""),
                "section_guess": prefix,
                "code_prefix": prefix,
                "first_source_code": group[0].get("source_code", ""),
                "last_source_code": group[-1].get("source_code", ""),
                "row_count": len(group),
                "sample_names": ";".join(row.get("raw_name", "") for row in group[:5]),
                "confidence": "0.82",
                "remark": "section inventory generated from source_code prefix; human confirmation required for official section names",
            }
        )
    return rows


def build_gd_issues(all_rows: Sequence[Dict[str, Any]], quota_rows: Sequence[Dict[str, Any]], extra_fields: Sequence[str]) -> List[Dict[str, Any]]:
    issues: List[Dict[str, Any]] = []

    def add(row: Dict[str, Any], issue_type: str, detail: str, severity: str, action: str) -> None:
        issues.append(
            {
                "issue_id": f"ISSUE_GD2018_BASELINE_{len(issues) + 1:05d}",
                "source_row_id": row.get("source_row_id", ""),
                "source_excel_row": row.get("source_excel_row", ""),
                "source_code_raw": row.get("source_code_raw", ""),
                "source_code_normalized": row.get("source_code_normalized", ""),
                "issue_type": issue_type,
                "issue_detail": detail,
                "severity": severity,
                "suggested_action": action,
            }
        )

    for row in all_rows:
        for issue_type in [value for value in row.get("parse_issue", "").split(";") if value]:
            add(row, issue_type, f"All-rows parse issue: {issue_type}.", "medium", "Review source row in normalized Excel.")
    for row in quota_rows:
        if row.get("review_status") != REVIEW_STATUS:
            add(row, "non_pending_review_status", "Quota review_status is not pending.", "critical", "Baseline rows must remain pending.")
        if not row.get("unit"):
            add(row, "missing_unit", "Quota unit is blank.", "high", "Review normalized Excel source row.")
        if not row.get("raw_name"):
            add(row, "missing_name", "Quota raw_name is blank.", "high", "Review normalized Excel source row.")
        if not any(row.get(field) not in ("", None) for field in ["raw_labor_fee", "raw_material_fee", "raw_machine_fee", "raw_management_fee", "raw_total_fee"]):
            add(row, "missing_all_pricing_fields", "All pricing fields are blank.", "high", "Review normalized Excel source row.")
    if extra_fields:
        issues.append(
            {
                "issue_id": f"ISSUE_GD2018_BASELINE_{len(issues) + 1:05d}",
                "source_row_id": "",
                "source_excel_row": "",
                "source_code_raw": "",
                "source_code_normalized": "",
                "issue_type": "original_fields_preserved_as_raw_extra",
                "issue_detail": f"Original workbook fields preserved as {len(extra_fields)} raw_extra_colNN_* columns in all_rows_review.",
                "severity": "low",
                "suggested_action": "Use raw_extra columns for source audit when normalized field aliases are insufficient.",
            }
        )
    return issues


def build_gd_summary_rows(
    excel_path: Path,
    excel_hash: str,
    all_rows: Sequence[Dict[str, Any]],
    quota_rows: Sequence[Dict[str, Any]],
    pricing_rows: Sequence[Dict[str, Any]],
    inventory: Sequence[Dict[str, Any]],
    issues: Sequence[Dict[str, Any]],
    extra_fields: Sequence[str],
) -> List[Dict[str, Any]]:
    return [
        {"metric_name": "source_file", "metric_value": excel_path.name, "remark": ""},
        {"metric_name": "source_file_hash", "metric_value": excel_hash, "remark": "SHA256"},
        {"metric_name": "file_size_bytes", "metric_value": excel_path.stat().st_size, "remark": ""},
        {"metric_name": "all_rows", "metric_value": len(all_rows), "remark": "non-empty data rows below header"},
        {"metric_name": "quota_item_rows", "metric_value": len(quota_rows), "remark": ""},
        {"metric_name": "pricing_review_rows", "metric_value": len(pricing_rows), "remark": ""},
        {"metric_name": "section_inventory_rows", "metric_value": len(inventory), "remark": ""},
        {"metric_name": "parse_issue_rows", "metric_value": len(issues), "remark": ""},
        {"metric_name": "raw_extra_field_count", "metric_value": len(extra_fields), "remark": "original workbook fields preserved"},
        {"metric_name": "blank_source_code_normalized", "metric_value": sum(1 for row in all_rows if not row.get("source_code_normalized")), "remark": ""},
        {"metric_name": "blank_raw_name", "metric_value": sum(1 for row in all_rows if not row.get("raw_name")), "remark": ""},
        {"metric_name": "blank_raw_unit", "metric_value": sum(1 for row in all_rows if not row.get("raw_unit")), "remark": ""},
        {"metric_name": "quota_non_pending_review_status", "metric_value": sum(1 for row in quota_rows if row.get("review_status") != REVIEW_STATUS), "remark": ""},
        {"metric_name": "quota_approved_rows", "metric_value": sum(1 for row in quota_rows if row.get("review_status") == "approved"), "remark": ""},
    ]


def build_gd_package(project_root: Path) -> Dict[str, Any]:
    excel_path = project_root / SOURCE_EXCEL_REL
    if not excel_path.exists():
        raise SystemExit(f"blocked_missing_inputs: {excel_path}")
    excel_hash = sha256_file(excel_path)
    workbook = load_workbook(excel_path, read_only=True, data_only=False)
    visible_sheets = [ws for ws in workbook.worksheets if ws.sheet_state == "visible"]
    if not visible_sheets:
        raise SystemExit("blocked_missing_inputs: normalized Excel has no visible sheets")
    ws = visible_sheets[0]
    headers = [norm(cell.value) for cell in next(ws.iter_rows(min_row=1, max_row=1))]
    all_rows, extra_fields = build_gd_all_rows(excel_path, excel_hash, ws, headers)
    quota_rows = build_gd_quota_rows(all_rows)
    pricing_rows = build_gd_pricing_rows(quota_rows)
    inventory = build_gd_section_inventory(quota_rows)
    issues = build_gd_issues(all_rows, quota_rows, extra_fields)
    output_dir = project_root / GD_OUTPUT_REL
    output_dir.mkdir(parents=True, exist_ok=True)
    all_rows_fields = GD_BASE_FIELDS + extra_fields
    write_csv(output_dir / "gd2018_normalized_all_rows_review.csv", all_rows_fields, all_rows)
    write_csv(output_dir / "gd2018_normalized_quota_items_full_review.csv", GD_QUOTA_FIELDS, quota_rows)
    write_csv(output_dir / "gd2018_normalized_section_inventory_full.csv", GD_SECTION_FIELDS, inventory)
    write_csv(output_dir / "gd2018_normalized_pricing_fields_full_review.csv", GD_PRICING_FIELDS, pricing_rows)
    write_csv(output_dir / "gd2018_normalized_full_parse_issues.csv", GD_ISSUE_FIELDS, issues)
    summary_rows = build_gd_summary_rows(excel_path, excel_hash, all_rows, quota_rows, pricing_rows, inventory, issues, extra_fields)
    try:
        write_xlsx(
            output_dir / "GD2018_normalized_full_quota_parse_review.xlsx",
            [
                ("all_rows_review", all_rows_fields, all_rows),
                ("quota_items_full", GD_QUOTA_FIELDS, quota_rows),
                ("pricing_fields_full", GD_PRICING_FIELDS, pricing_rows),
                ("section_inventory_full", GD_SECTION_FIELDS, inventory),
                ("parse_issues", GD_ISSUE_FIELDS, issues),
                ("summary", ["metric_name", "metric_value", "remark"], summary_rows),
            ],
        )
    except Exception as exc:  # pragma: no cover - escalates as stage blocker.
        raise SystemExit(f"blocked_xlsx_generation_failed: GD2018 workbook generation failed: {exc}") from exc
    write_gd_report(output_dir / "gd2018_normalized_full_quota_parse_review_report.md", excel_path, excel_hash, ws, headers, all_rows, quota_rows, pricing_rows, inventory, issues, extra_fields)
    return {
        "excel_path": excel_path,
        "source_hash": excel_hash,
        "output_dir": output_dir,
        "all_rows": all_rows,
        "quota_rows": quota_rows,
        "pricing_rows": pricing_rows,
        "inventory": inventory,
        "issues": issues,
        "extra_fields": extra_fields,
        "xlsx_path": output_dir / "GD2018_normalized_full_quota_parse_review.xlsx",
    }


def write_gd_report(
    path: Path,
    excel_path: Path,
    excel_hash: str,
    ws: Any,
    headers: Sequence[str],
    all_rows: Sequence[Dict[str, Any]],
    quota_rows: Sequence[Dict[str, Any]],
    pricing_rows: Sequence[Dict[str, Any]],
    inventory: Sequence[Dict[str, Any]],
    issues: Sequence[Dict[str, Any]],
    extra_fields: Sequence[str],
) -> None:
    pricing_complete = sum(1 for row in pricing_rows if row.get("has_any_pricing") == "true")
    issue_counts = Counter(row.get("issue_type", "") for row in issues)
    lines = [
        "# GD2018 Normalized Full Quota Parse Review Report",
        "",
        "## 1. Task Scope",
        "",
        "Read the normalized GD2018 Excel source directly and generate a complete engineering parse review package. This package preserves all original workbook fields and adds engineering review columns.",
        "",
        "## 2. Source Excel Profile",
        "",
        f"- source_file: `{excel_path}`",
        f"- source_file_hash: `{excel_hash}`",
        f"- source_sheet: `{ws.title}`",
        f"- max_row: {ws.max_row}",
        f"- max_column: {ws.max_column}",
        f"- original_headers: {json.dumps(list(headers), ensure_ascii=False)}",
        "",
        "## 3. All Rows Review Summary",
        "",
        f"- all_rows_review_rows: {len(all_rows)}",
        f"- blank_source_code_normalized_rows: {sum(1 for row in all_rows if not row.get('source_code_normalized'))}",
        f"- blank_raw_name_rows: {sum(1 for row in all_rows if not row.get('raw_name'))}",
        f"- blank_raw_unit_rows: {sum(1 for row in all_rows if not row.get('raw_unit'))}",
        "",
        "## 4. Quota Item Summary",
        "",
        f"- quota_item_rows: {len(quota_rows)}",
        f"- section_inventory_rows: {len(inventory)}",
        f"- supplemental_source_code_rows: {sum(1 for row in quota_rows if is_supplemental_code(row.get('source_code', '')))}",
        "",
        "## 5. Original Field Preservation",
        "",
        f"- original_field_count: {len(headers)}",
        f"- raw_extra_field_count: {len(extra_fields)}",
        "- Original workbook fields are preserved as `raw_extra_colNN_*` columns in `gd2018_normalized_all_rows_review.csv`.",
        "",
        "## 6. Pricing Field Preservation",
        "",
        f"- pricing_review_rows: {len(pricing_rows)}",
        f"- rows_with_any_pricing: {pricing_complete}",
        f"- rows_missing_all_pricing: {len(pricing_rows) - pricing_complete}",
        "",
        "## 7. Name Governance",
        "",
        "- `quota_name_candidate` is not the final enterprise standard name.",
        "- It is a quota candidate description lightly cleaned from the normalized Excel source.",
        "- A phrase such as `人工挖沟槽土方 一、二类土 深度在4m内` must remain a quota candidate description until reviewed.",
        "- This stage must not generate approved enterprise `standard_name` values.",
        "",
        "## 8. Parse Issues",
        "",
        f"- parse_issue_rows: {len(issues)}",
        f"- issue_type_counts: {json.dumps(dict(issue_counts), ensure_ascii=False, sort_keys=True)}",
        "",
        "## 9. Not Approved / Not Final Statement",
        "",
        "All rows remain `pending`. This stage does not generate bill_code, mappings, approved records, internal_price_library, database rows, or enterprise templates.",
        "",
        "## 10. Human Review Guidance",
        "",
        "- Confirm source_code, raw_name, unit, and pricing fields against the normalized Excel source.",
        "- Treat `quota_name_candidate` as a review candidate only.",
        "- Use `section_inventory_full` to decide which source-code prefixes need official section-name confirmation.",
        "",
    ]
    path.write_text("\n".join(lines), encoding="utf-8")


def artifact_row_count(path: Path) -> str:
    if not path.exists():
        return ""
    if path.suffix.lower() == ".csv":
        return str(csv_row_count(path))
    if path.suffix.lower() == ".xlsx":
        workbook = load_workbook(path, read_only=True, data_only=True)
        return str(sum(max(0, ws.max_row - 1) for ws in workbook.worksheets))
    return ""


def manifest_row(stage: str, artifact: str, path: Path, source_file: str, project_root: Path) -> Dict[str, str]:
    exists = path.exists()
    return {
        "stage_name": stage,
        "artifact_name": artifact,
        "expected_path": rel(path, project_root),
        "exists": str(exists).lower(),
        "file_size_bytes": str(path.stat().st_size) if exists else "",
        "row_count": artifact_row_count(path) if exists else "",
        "sha256": sha256_file(path) if exists else "",
        "created_or_modified_time": datetime.fromtimestamp(path.stat().st_mtime).isoformat(timespec="seconds") if exists else "",
        "source_file": source_file,
        "can_regenerate": "true",
        "backup_required": "true",
        "backup_path": "construction_cost_knowledge_engine/data/private/reference_extraction/backups/runs_backup_after_SOURCE_BASELINE_LOCK_1",
        "status": "generated" if exists else "missing",
        "remark": "source baseline lock artifact; private artifact; pending human review; no approved data",
    }


def update_manifest(project_root: Path, artifacts: Sequence[Path], source_file: str) -> None:
    manifest_path = project_root / DOCS_REF_REL / "reference_artifact_manifest.csv"
    manifest_path.parent.mkdir(parents=True, exist_ok=True)
    existing: List[Dict[str, str]] = read_csv(manifest_path) if manifest_path.exists() else []
    replacement = {
        (STAGE_NAME, path.name): manifest_row(STAGE_NAME, path.name, path, source_file, project_root)
        for path in artifacts
    }
    filtered = [row for row in existing if (row.get("stage_name"), row.get("artifact_name")) not in replacement]
    filtered.extend(replacement.values())
    write_csv(manifest_path, MANIFEST_FIELDS, filtered)
    write_manifest_md(project_root, filtered)


def write_manifest_md(project_root: Path, rows: Sequence[Dict[str, str]]) -> None:
    latest = [row for row in rows if row.get("stage_name") == STAGE_NAME]
    registered = len(rows)
    existing = sum(1 for row in rows if row.get("exists") == "true")
    lines = [
        "# Reference Artifact Manifest",
        "",
        "## Governance",
        "",
        "- `construction_cost_knowledge_engine/data/private/` is intentionally not tracked by Git.",
        "- Every private artifact must be registered with `row_count` and `sha256` after generation.",
        "- Each completed stage must back up its `runs` output directory after validation.",
        "- Mock SQLite and seed CSV artifacts must not be used as source of truth.",
        "- Source baseline lock outputs are pending review artifacts only and do not approve mappings.",
        "",
        "## Current Manifest Summary",
        "",
        f"- registered_artifacts: {registered}",
        f"- existing_artifacts: {existing}",
        f"- missing_artifacts: {registered - existing}",
        "",
        "## Manifest CSV",
        "",
        "`construction_cost_knowledge_engine/docs/reference_extraction/reference_artifact_manifest.csv`",
        "",
        "## Latest Source Baseline Lock Outputs",
        "",
    ]
    for row in latest:
        lines.append(f"- `{row.get('expected_path')}` ({row.get('row_count') or 'n/a'} rows)")
    lines.extend(
        [
            "",
            "## Backup Requirement",
            "",
            "`construction_cost_knowledge_engine/data/private/reference_extraction/backups/runs_backup_after_SOURCE_BASELINE_LOCK_1/`",
            "",
        ]
    )
    (project_root / DOCS_REF_REL / "REFERENCE_ARTIFACT_MANIFEST.md").write_text("\n".join(lines), encoding="utf-8")


def write_summary_report(path: Path, gb_result: Dict[str, Any], gd_result: Dict[str, Any], recommendation: str) -> None:
    gb_candidates = gb_result["candidates"]
    gd_all = gd_result["all_rows"]
    gd_quota = gd_result["quota_rows"]
    lines = [
        "# Source Baseline Lock Summary",
        "",
        "## 1. Task Scope",
        "",
        "Generate two independent source-data baseline review packages before continuing mapping or enterprise-standard work.",
        "",
        "## 2. Why This Stage Is Needed",
        "",
        "Previous stages had already entered mapping. This lock stage gives reviewers two complete, independent source baselines: the national bill standard parsed from DOCX and the Guangdong normalized quota workbook parsed from XLSX.",
        "",
        "## 3. GB50854 Baseline Output",
        "",
        f"- xlsx: `{gb_result['xlsx_path']}`",
        f"- bill_item_rows: {len(gb_candidates)}",
        f"- context_rule_rows: {len(gb_result['rules'])}",
        f"- source_block_rows: {len(gb_result['source_blocks'])}",
        f"- missing_quantity_calculation_rule_rows: {sum(1 for row in gb_candidates if not row.get('quantity_calculation_rule'))}",
        f"- missing_work_content_rows: {sum(1 for row in gb_candidates if not row.get('work_content_raw'))}",
        "",
        "## 4. GD2018 Baseline Output",
        "",
        f"- xlsx: `{gd_result['xlsx_path']}`",
        f"- all_rows_review_rows: {len(gd_all)}",
        f"- quota_item_rows: {len(gd_quota)}",
        f"- pricing_review_rows: {len(gd_result['pricing_rows'])}",
        f"- section_inventory_rows: {len(gd_result['inventory'])}",
        "",
        "## 5. Completeness Checks",
        "",
        f"- GB bill_code_9 invalid count: {sum(1 for row in gb_candidates if not re.fullmatch(r'\\d{9}', row.get('bill_code_9', '') or ''))}",
        f"- GB duplicate bill_code_9 count: {sum(1 for _, count in Counter(row.get('bill_code_9', '') for row in gb_candidates).items() if count > 1)}",
        f"- GB non_pending rows: {sum(1 for row in gb_candidates if row.get('review_status') != REVIEW_STATUS)}",
        f"- GD source_code_normalized blank count: {sum(1 for row in gd_all if not row.get('source_code_normalized'))}",
        f"- GD raw_name blank count: {sum(1 for row in gd_all if not row.get('raw_name'))}",
        f"- GD raw_unit blank count: {sum(1 for row in gd_all if not row.get('raw_unit'))}",
        f"- GD quota non_pending rows: {sum(1 for row in gd_quota if row.get('review_status') != REVIEW_STATUS)}",
        "",
        "## 6. Governance Notes",
        "",
        "- `quota_name_candidate` and `standard_name_candidate` style fields are candidate descriptions only.",
        "- No source baseline row is approved.",
        "- No source baseline row should be treated as final enterprise standard naming.",
        "- Private artifacts must be backed up after manual validation.",
        "",
        "## 7. What Is Not Done",
        "",
        "- No database write.",
        "- No migration or schema change.",
        "- No new quota-to-bill mapping.",
        "- No bill_code write-back.",
        "- No internal_price_library.",
        "- No enterprise template generation.",
        "",
        "## 8. Recommendation",
        "",
        recommendation,
        "",
        f"Generated at: {datetime.now().isoformat(timespec='seconds')}",
        "",
    ]
    path.write_text("\n".join(lines), encoding="utf-8")


def determine_recommendation(gb_result: Dict[str, Any], gd_result: Dict[str, Any]) -> str:
    gb_candidates = gb_result["candidates"]
    gd_all = gd_result["all_rows"]
    gd_quota = gd_result["quota_rows"]
    if not gb_result["xlsx_path"].exists() or not gd_result["xlsx_path"].exists():
        return "blocked_xlsx_generation_failed"
    if not gb_candidates:
        return "gb50854_baseline_failed"
    if not gd_all or not gd_quota:
        return "gd2018_baseline_failed"
    gb_invalid = sum(1 for row in gb_candidates if not re.fullmatch(r"\d{9}", row.get("bill_code_9", "") or ""))
    gb_duplicates = sum(1 for _, count in Counter(row.get("bill_code_9", "") for row in gb_candidates).items() if count > 1)
    gd_non_pending = sum(1 for row in gd_quota if row.get("review_status") != REVIEW_STATUS)
    if gb_invalid or gb_duplicates or gd_non_pending:
        return "source_baseline_partial_manual_intervention_required"
    return "source_baseline_ready_for_human_confirmation"


def main() -> int:
    args = parse_args()
    project_root = args.project_root
    docx_path = project_root / SOURCE_DOCX_REL
    excel_path = project_root / SOURCE_EXCEL_REL
    if not docx_path.exists() or not excel_path.exists():
        missing = [str(path) for path in [docx_path, excel_path] if not path.exists()]
        raise SystemExit("blocked_missing_inputs: " + ";".join(missing))

    gb_result = build_gb_package(project_root)
    gd_result = build_gd_package(project_root)
    recommendation = determine_recommendation(gb_result, gd_result)
    output_root = project_root / OUTPUT_ROOT_REL
    output_root.mkdir(parents=True, exist_ok=True)
    summary_path = output_root / "source_baseline_lock_summary.md"
    write_summary_report(summary_path, gb_result, gd_result, recommendation)

    artifacts = [
        gb_result["xlsx_path"],
        gb_result["output_dir"] / "gb50854_bill_items_full_review.csv",
        gb_result["output_dir"] / "gb50854_context_rules_full_review.csv",
        gb_result["output_dir"] / "gb50854_appendix_table_registry_full.csv",
        gb_result["output_dir"] / "gb50854_source_blocks_full.csv",
        gb_result["output_dir"] / "gb50854_full_parse_issues.csv",
        gb_result["output_dir"] / "gb50854_full_standard_parse_review_report.md",
        gd_result["xlsx_path"],
        gd_result["output_dir"] / "gd2018_normalized_all_rows_review.csv",
        gd_result["output_dir"] / "gd2018_normalized_quota_items_full_review.csv",
        gd_result["output_dir"] / "gd2018_normalized_section_inventory_full.csv",
        gd_result["output_dir"] / "gd2018_normalized_pricing_fields_full_review.csv",
        gd_result["output_dir"] / "gd2018_normalized_full_parse_issues.csv",
        gd_result["output_dir"] / "gd2018_normalized_full_quota_parse_review_report.md",
        summary_path,
    ]
    update_manifest(
        project_root,
        artifacts,
        f"{rel(docx_path, project_root)};{rel(excel_path, project_root)}",
    )

    print(f"recommendation={recommendation}")
    print(f"gb_xlsx_exists={gb_result['xlsx_path'].exists()}")
    print(f"gd_xlsx_exists={gd_result['xlsx_path'].exists()}")
    print(f"gb_bill_item_rows={len(gb_result['candidates'])}")
    print(f"gb_context_rule_rows={len(gb_result['rules'])}")
    print(f"gb_source_block_rows={len(gb_result['source_blocks'])}")
    print(f"gb_missing_quantity_rule={sum(1 for row in gb_result['candidates'] if not row.get('quantity_calculation_rule'))}")
    print(f"gb_missing_work_content={sum(1 for row in gb_result['candidates'] if not row.get('work_content_raw'))}")
    print(f"gd_all_rows={len(gd_result['all_rows'])}")
    print(f"gd_quota_item_rows={len(gd_result['quota_rows'])}")
    print(f"gd_pricing_rows={len(gd_result['pricing_rows'])}")
    print(f"output_root={output_root}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
