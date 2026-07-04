#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Stage B-DOCX-1 extraction for GB/T 50854-2024 Appendix A.

This script is intentionally limited to Appendix A bill item reference
candidates. It does not write any database, does not modify migrations or
pipeline code, and does not generate quota_to_bill_mapping.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import re
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple

from docx import Document

DEFAULT_DOCX_PATH = Path(
    r"E:\workspace\01_Projects\ai-construction-system\construction_cost_knowledge_engine\data\private\reference_extraction\source_standards\房屋建筑与装饰工程工程量计算标准 GB_T 50854-2024.docx"
)

SOURCE_TYPE = "official_bill_standard_docx"
SOURCE_NAME = "房屋建筑与装饰工程工程量计算标准 GB/T 50854-2024"
APPENDIX_CODE = "A"
APPENDIX_NAME = "土石方工程"
REVIEW_STATUS = "pending"
EXTRACTION_METHOD = "docx_table"
BILL_CODE_POLICY = "标准表提供一至九位编码；实际工程清单十二位编码的十至十二位应按具体工程项目名称和项目特征另行编制，本阶段不生成完整十二位编码。"

PROFILE_FIELDS = [
    "profile_id",
    "source_type",
    "source_name",
    "source_file",
    "source_file_hash",
    "file_exists",
    "file_size_bytes",
    "paragraph_count",
    "table_count",
    "can_read_text",
    "can_read_tables",
    "contains_toc",
    "contains_appendix_A",
    "contains_A1",
    "contains_A2",
    "contains_A3",
    "contains_A4",
    "remark",
]

REGISTRY_FIELDS = [
    "registry_id",
    "source_type",
    "source_name",
    "source_file",
    "source_file_hash",
    "source_heading_path",
    "appendix_code",
    "appendix_name",
    "section_code",
    "section_name",
    "table_code",
    "table_name",
    "table_base_code",
    "source_table_index",
    "start_paragraph_index",
    "end_paragraph_index",
    "extraction_method",
    "extraction_confidence",
    "parse_issue",
    "remark",
]

CANDIDATE_FIELDS = [
    "bill_reference_id",
    "source_type",
    "source_name",
    "source_file",
    "source_file_hash",
    "source_heading_path",
    "source_table_index",
    "source_row_index",
    "appendix_code",
    "appendix_name",
    "section_code",
    "section_name",
    "table_code",
    "table_name",
    "table_base_code",
    "bill_code_9",
    "bill_code_full_policy",
    "bill_name",
    "project_feature_raw",
    "project_feature_items_json",
    "unit",
    "quantity_calculation_rule",
    "work_content_raw",
    "work_content_items_json",
    "keywords",
    "extraction_method",
    "extraction_confidence",
    "review_status",
    "reviewer",
    "remark",
]

RULE_FIELDS = [
    "rule_id",
    "source_type",
    "source_name",
    "source_file",
    "source_file_hash",
    "source_heading_path",
    "appendix_code",
    "appendix_name",
    "rule_code",
    "rule_text",
    "related_bill_codes",
    "extraction_method",
    "extraction_confidence",
    "review_status",
    "remark",
]

ISSUE_FIELDS = [
    "issue_id",
    "source_location",
    "appendix_code",
    "section_code",
    "table_code",
    "bill_code_9",
    "issue_type",
    "issue_detail",
    "severity",
    "suggested_action",
]

TABLE_META = {
    "010101": {
        "section_code": "A.1",
        "section_name": "单独土石方",
        "table_code": "表A.1.1",
        "table_name": "单独土石方",
        "heading": "附录A 土石方工程 > A.1 单独土石方 > 表A.1.1 单独土石方（编码：010101）",
    },
    "010102": {
        "section_code": "A.2",
        "section_name": "基础土石方",
        "table_code": "表A.2.1",
        "table_name": "基础土石方",
        "heading": "附录A 土石方工程 > A.2 基础土石方 > 表A.2.1 基础土石方（编码：010102）",
    },
    "010103": {
        "section_code": "A.3",
        "section_name": "平整场地及其他",
        "table_code": "表A.3.1",
        "table_name": "平整场地及其他",
        "heading": "附录A 土石方工程 > A.3 平整场地及其他 > 表A.3.1 平整场地及其他（编码：010103）",
    },
}


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def norm(value: Any) -> str:
    if value is None:
        return ""
    text = str(value).replace("\r", "\n").replace("\u3000", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n+", "\n", text)
    return text.strip()


def compact(value: Any) -> str:
    return re.sub(r"\s+", " ", norm(value)).strip()


def para_texts(doc: Document) -> List[str]:
    return [norm(p.text) for p in doc.paragraphs]


def find_first(texts: Sequence[str], predicate) -> Optional[int]:
    for idx, text in enumerate(texts):
        if predicate(text):
            return idx
    return None


def has_text(texts: Sequence[str], needle: str) -> bool:
    return any(needle in text for text in texts)


def table_rows(table) -> List[List[str]]:
    return [[norm(cell.text) for cell in row.cells] for row in table.rows]


def table_base_code(rows: Sequence[Sequence[str]]) -> str:
    for row in rows[1:]:
        if row and re.fullmatch(r"\d{9}", compact(row[0])):
            return compact(row[0])[:6]
    return ""


def split_numbered_items(raw: str) -> List[str]:
    text = compact(raw)
    if not text:
        return []
    pattern = r"(?:^|\s)(\d+[\.．、]\s*[^0-9]+?)(?=\s+\d+[\.．、]\s*|$)"
    matches = [compact(m.group(1)) for m in re.finditer(pattern, text)]
    if matches:
        return matches
    return [text]


def keywords(*parts: str) -> str:
    text = " ".join(parts)
    tokens = [t for t in re.split(r"[\s、，；;（）()]+", text) if len(t) >= 2]
    seen = set()
    result = []
    for token in tokens:
        if token not in seen:
            seen.add(token)
            result.append(token)
    return ";".join(result[:10])


def validate_bill_code(code: str) -> bool:
    return bool(re.fullmatch(r"\d{9}", code))


def make_issue(
    issues: List[Dict[str, Any]],
    source_location: str,
    section_code: str,
    table_code: str,
    bill_code_9: str,
    issue_type: str,
    detail: str,
    severity: str,
    action: str,
) -> None:
    issues.append(
        {
            "issue_id": f"ISSUE_GB50854_A_{len(issues) + 1:03d}",
            "source_location": source_location,
            "appendix_code": APPENDIX_CODE,
            "section_code": section_code,
            "table_code": table_code,
            "bill_code_9": bill_code_9,
            "issue_type": issue_type,
            "issue_detail": detail,
            "severity": severity,
            "suggested_action": action,
        }
    )


def build_profile(docx_path: Path, doc: Document, source_hash: str, texts: Sequence[str]) -> List[Dict[str, Any]]:
    nonempty_text = [t for t in texts if t]
    return [
        {
            "profile_id": "GB50854_2024_DOCX",
            "source_type": SOURCE_TYPE,
            "source_name": SOURCE_NAME,
            "source_file": docx_path.name,
            "source_file_hash": source_hash,
            "file_exists": str(docx_path.exists()).lower(),
            "file_size_bytes": docx_path.stat().st_size if docx_path.exists() else "",
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "can_read_text": str(bool(nonempty_text)).lower(),
            "can_read_tables": str(len(doc.tables) > 0).lower(),
            "contains_toc": str(has_text(texts, "目  次")).lower(),
            "contains_appendix_A": str(has_text(texts, "附录A 土石方工程")).lower(),
            "contains_A1": str(has_text(texts, "A.1 单独土石方")).lower(),
            "contains_A2": str(has_text(texts, "A.2 基础土石方")).lower(),
            "contains_A3": str(has_text(texts, "A.3 平整场地及其他")).lower(),
            "contains_A4": str(has_text(texts, "A.4 其他规定")).lower(),
            "remark": "DOCX-first extraction; no OCR and no PDF parsing used.",
        }
    ]


def build_registry(docx_path: Path, source_hash: str, texts: Sequence[str], doc: Document) -> List[Dict[str, Any]]:
    idx_app = find_first(texts, lambda t: t == "附录A 土石方工程")
    idx_a1 = find_first(texts, lambda t: t == "A.1 单独土石方")
    idx_t1 = find_first(texts, lambda t: t.startswith("表A.1.1"))
    idx_a2 = find_first(texts, lambda t: t == "A.2 基础土石方")
    idx_t2 = find_first(texts, lambda t: t.startswith("表A.2.1"))
    idx_a3 = find_first(texts, lambda t: t == "A.3 平整场地及其他")
    idx_t3 = find_first(texts, lambda t: t.startswith("表A.3.1"))
    idx_a4 = find_first(texts, lambda t: t == "A.4 其他规定")
    idx_b = find_first(texts, lambda t: t.startswith("附录B "))

    table_index_by_base: Dict[str, int] = {}
    for table_idx, table in enumerate(doc.tables):
        base = table_base_code(table_rows(table))
        if base in TABLE_META and base not in table_index_by_base:
            table_index_by_base[base] = table_idx

    def row(
        rid: str,
        heading: str,
        section_code: str,
        section_name: str,
        table_code: str,
        table_name: str,
        table_base: str,
        table_index: Any,
        start: Any,
        end: Any,
        rtype: str,
        issue: str = "",
        remark: str = "",
    ) -> Dict[str, Any]:
        return {
            "registry_id": rid,
            "source_type": SOURCE_TYPE,
            "source_name": SOURCE_NAME,
            "source_file": docx_path.name,
            "source_file_hash": source_hash,
            "source_heading_path": heading,
            "appendix_code": APPENDIX_CODE,
            "appendix_name": APPENDIX_NAME,
            "section_code": section_code,
            "section_name": section_name,
            "table_code": table_code,
            "table_name": table_name,
            "table_base_code": table_base,
            "source_table_index": table_index,
            "start_paragraph_index": start,
            "end_paragraph_index": end,
            "extraction_method": "docx_paragraph_and_table_index",
            "extraction_confidence": "0.96" if not issue else "0.70",
            "parse_issue": issue,
            "remark": remark or rtype,
        }

    return [
        row("GB50854_A_APPENDIX", "附录A 土石方工程", "A", APPENDIX_NAME, "", "", "", "", idx_app, (idx_b - 1) if idx_b else "", "appendix"),
        row("GB50854_A1_SECTION", "附录A 土石方工程 > A.1 单独土石方", "A.1", "单独土石方", "", "", "", "", idx_a1, (idx_a2 - 1) if idx_a2 else "", "section"),
        row("GB50854_A1_TABLE", TABLE_META["010101"]["heading"], "A.1", "单独土石方", "表A.1.1", "单独土石方", "010101", table_index_by_base.get("010101", ""), idx_t1, idx_t1, "table"),
        row("GB50854_A2_SECTION", "附录A 土石方工程 > A.2 基础土石方", "A.2", "基础土石方", "", "", "", "", idx_a2, (idx_a3 - 1) if idx_a3 else "", "section"),
        row("GB50854_A2_TABLE", TABLE_META["010102"]["heading"], "A.2", "基础土石方", "表A.2.1", "基础土石方", "010102", table_index_by_base.get("010102", ""), idx_t2, idx_t2, "table"),
        row("GB50854_A3_SECTION", "附录A 土石方工程 > A.3 平整场地及其他", "A.3", "平整场地及其他", "", "", "", "", idx_a3, (idx_a4 - 1) if idx_a4 else "", "section"),
        row("GB50854_A3_TABLE", TABLE_META["010103"]["heading"], "A.3", "平整场地及其他", "表A.3.1", "平整场地及其他", "010103", table_index_by_base.get("010103", ""), idx_t3, idx_t3, "table"),
        row("GB50854_A4_RULES", "附录A 土石方工程 > A.4 其他规定", "A.4", "其他规定", "", "", "", "", idx_a4, (idx_b - 1) if idx_b else "", "context_rules"),
    ]


def extract_candidates(docx_path: Path, source_hash: str, doc: Document, issues: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    for table_idx, table in enumerate(doc.tables):
        rows_raw = table_rows(table)
        if not rows_raw:
            continue
        base = table_base_code(rows_raw)
        if base not in TABLE_META:
            continue
        meta = TABLE_META[base]
        for row_idx, cells in enumerate(rows_raw[1:], start=1):
            padded = list(cells) + [""] * (6 - len(cells))
            bill_code, bill_name, feature, unit, quantity_rule, work_content = [norm(c) for c in padded[:6]]
            if not any([bill_code, bill_name, feature, unit, quantity_rule, work_content]):
                continue
            source_location = f"table={table_idx};row={row_idx}"
            if not bill_code:
                make_issue(issues, source_location, meta["section_code"], meta["table_code"], bill_code, "missing_bill_code", "Missing bill code.", "high", "Verify source row before use.")
            elif not validate_bill_code(bill_code):
                make_issue(issues, source_location, meta["section_code"], meta["table_code"], bill_code, "invalid_bill_code", "Bill code is not 9 digits.", "high", "Do not import until corrected.")
            if re.search(r"A1-1-\d+", bill_code):
                make_issue(issues, source_location, meta["section_code"], meta["table_code"], bill_code, "invalid_bill_code", "Unexpected Guangdong quota code pattern found in bill_code.", "critical", "Remove row from bill reference candidates.")
            if not bill_name:
                make_issue(issues, source_location, meta["section_code"], meta["table_code"], bill_code, "missing_bill_name", "Missing bill name.", "high", "Verify source table.")
            if not feature:
                make_issue(issues, source_location, meta["section_code"], meta["table_code"], bill_code, "missing_project_feature", "Missing project feature.", "medium", "Manual QA required.")
            if not unit:
                make_issue(issues, source_location, meta["section_code"], meta["table_code"], bill_code, "missing_unit", "Missing unit.", "high", "Manual QA required.")
            if not quantity_rule:
                make_issue(issues, source_location, meta["section_code"], meta["table_code"], bill_code, "missing_quantity_rule", "Missing quantity calculation rule.", "high", "Manual QA required.")
            if not work_content:
                make_issue(issues, source_location, meta["section_code"], meta["table_code"], bill_code, "missing_work_content", "Missing work content.", "medium", "Manual QA required.")
            feature_items = split_numbered_items(feature)
            work_items = split_numbered_items(work_content)
            if feature and not feature_items:
                make_issue(issues, source_location, meta["section_code"], meta["table_code"], bill_code, "multiline_feature_parse_issue", "Could not split project feature items.", "low", "Keep raw field as source of truth.")
            if work_content and not work_items:
                make_issue(issues, source_location, meta["section_code"], meta["table_code"], bill_code, "multiline_work_content_parse_issue", "Could not split work content items.", "low", "Keep raw field as source of truth.")
            confidence = "0.96" if validate_bill_code(bill_code) and bill_name and unit and quantity_rule and work_content else "0.70"
            rows.append(
                {
                    "bill_reference_id": f"GB50854_2024_A_{bill_code}",
                    "source_type": SOURCE_TYPE,
                    "source_name": SOURCE_NAME,
                    "source_file": docx_path.name,
                    "source_file_hash": source_hash,
                    "source_heading_path": meta["heading"],
                    "source_table_index": table_idx,
                    "source_row_index": row_idx,
                    "appendix_code": APPENDIX_CODE,
                    "appendix_name": APPENDIX_NAME,
                    "section_code": meta["section_code"],
                    "section_name": meta["section_name"],
                    "table_code": meta["table_code"],
                    "table_name": meta["table_name"],
                    "table_base_code": base,
                    "bill_code_9": bill_code,
                    "bill_code_full_policy": BILL_CODE_POLICY,
                    "bill_name": bill_name,
                    "project_feature_raw": feature,
                    "project_feature_items_json": json.dumps(feature_items, ensure_ascii=False),
                    "unit": unit,
                    "quantity_calculation_rule": quantity_rule,
                    "work_content_raw": work_content,
                    "work_content_items_json": json.dumps(work_items, ensure_ascii=False),
                    "keywords": keywords(bill_name, feature),
                    "extraction_method": EXTRACTION_METHOD,
                    "extraction_confidence": confidence,
                    "review_status": REVIEW_STATUS,
                    "reviewer": "",
                    "remark": "DOCX table extraction candidate; not enterprise final standard name.",
                }
            )
    return rows


def extract_rules(docx_path: Path, source_hash: str, texts: Sequence[str]) -> List[Dict[str, Any]]:
    rules: List[Dict[str, Any]] = []
    for idx, text in enumerate(texts):
        match = re.match(r"^(A\.4\.\d+)\s+(.+)", compact(text))
        if not match:
            continue
        rule_code, rule_text = match.groups()
        related_codes = []
        if "单独土石方" in rule_text:
            related_codes.extend(["010101001", "010101002", "010101003"])
        if "基础土石方" in rule_text or "沟槽" in rule_text or "基坑" in rule_text:
            related_codes.extend(["010102001", "010102002", "010102005", "010102006", "010102007"])
        if "平整场地" in rule_text:
            related_codes.append("010103001")
        if "余方弃置" in rule_text:
            related_codes.append("010103002")
        rules.append(
            {
                "rule_id": f"GB50854_A_{rule_code.replace('.', '_')}",
                "source_type": SOURCE_TYPE,
                "source_name": SOURCE_NAME,
                "source_file": docx_path.name,
                "source_file_hash": source_hash,
                "source_heading_path": "附录A 土石方工程 > A.4 其他规定",
                "appendix_code": APPENDIX_CODE,
                "appendix_name": APPENDIX_NAME,
                "rule_code": rule_code,
                "rule_text": rule_text,
                "related_bill_codes": ";".join(sorted(set(related_codes))),
                "extraction_method": "docx_paragraph",
                "extraction_confidence": "0.95",
                "review_status": REVIEW_STATUS,
                "remark": f"source_paragraph_index={idx}; A.4 context rule only; excluded from bill_item_reference candidates.",
            }
        )
    return rules


def write_csv(path: Path, fields: Sequence[str], rows: Sequence[Dict[str, Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8-sig") as fh:
        writer = csv.DictWriter(fh, fieldnames=fields)
        writer.writeheader()
        writer.writerows(rows)


def count_table(counter: Counter) -> str:
    lines = ["| Item | Count |", "|---|---:|"]
    for key in sorted(counter):
        lines.append(f"| {key} | {counter[key]} |")
    return "\n".join(lines)


def write_report(
    path: Path,
    docx_path: Path,
    profile: List[Dict[str, Any]],
    registry: List[Dict[str, Any]],
    candidates: List[Dict[str, Any]],
    rules: List[Dict[str, Any]],
    issues: List[Dict[str, Any]],
) -> None:
    section_counts = Counter(row["section_code"] for row in candidates)
    issue_counts = Counter(row["issue_type"] for row in issues)
    missing_required = []
    required = ["bill_reference_id", "bill_code_9", "bill_name", "unit", "quantity_calculation_rule", "work_content_raw", "review_status"]
    for idx, row in enumerate(candidates, start=1):
        for field in required:
            if not compact(row.get(field, "")):
                missing_required.append(f"row {idx} missing {field}")
    invalid_codes = [row["bill_code_9"] for row in candidates if not validate_bill_code(row["bill_code_9"])]
    quota_codes = [row["bill_code_9"] for row in candidates if re.search(r"A1-1-\d+", row["bill_code_9"])]
    non_pending = [row["bill_code_9"] for row in candidates if row["review_status"] != REVIEW_STATUS]
    go = len(candidates) == 12 and not missing_required and not invalid_codes and not quota_codes and not non_pending

    lines = [
        "# Stage B-DOCX-1 Report - GB/T 50854-2024 Appendix A Bill Item Reference",
        "",
        "## 1. Task Scope",
        "",
        "DOCX-first extraction of Appendix A bill item reference candidates only. No OCR, no PDF parsing, no database writes, no pipeline changes, no quota_to_bill_mapping, and no approved/internal_price_library generation.",
        "",
        "## 2. Input File Profile",
        "",
        f"- source_file: `{docx_path}`",
        f"- source_file_hash: `{profile[0]['source_file_hash']}`",
        f"- file_size_bytes: {profile[0]['file_size_bytes']}",
        f"- paragraph_count: {profile[0]['paragraph_count']}",
        f"- table_count: {profile[0]['table_count']}",
        f"- can_read_text: {profile[0]['can_read_text']}",
        f"- can_read_tables: {profile[0]['can_read_tables']}",
        "",
        "## 3. DOCX Structure Detection",
        "",
        f"- contains_toc: {profile[0]['contains_toc']}",
        f"- contains_appendix_A: {profile[0]['contains_appendix_A']}",
        f"- contains_A1: {profile[0]['contains_A1']}",
        f"- contains_A2: {profile[0]['contains_A2']}",
        f"- contains_A3: {profile[0]['contains_A3']}",
        f"- contains_A4: {profile[0]['contains_A4']}",
        "",
        "## 4. Appendix A Registry",
        "",
        "| Registry | Section | Table | Base Code | Table Index | Paragraph Start |",
        "|---|---|---|---|---:|---:|",
    ]
    for row in registry:
        lines.append(f"| {row['registry_id']} | {row['section_code']} {row['section_name']} | {row['table_code']} {row['table_name']} | {row['table_base_code']} | {row['source_table_index']} | {row['start_paragraph_index']} |")
    lines.extend(
        [
            "",
            "## 5. Extracted Bill Item Summary",
            "",
            f"- candidate_count: {len(candidates)}",
            count_table(section_counts),
            "",
            "## 6. Field Completeness Check",
            "",
            f"- missing_required_fields: {'; '.join(missing_required) if missing_required else 'none'}",
            f"- invalid_bill_code_9: {'; '.join(invalid_codes) if invalid_codes else 'none'}",
            f"- unexpected_A1_1_quota_codes: {'; '.join(quota_codes) if quota_codes else 'none'}",
            f"- non_pending_review_status: {'; '.join(non_pending) if non_pending else 'none'}",
            "",
            "## 7. A.4 Context Rules",
            "",
            f"- context_rule_count: {len(rules)}",
            "- A.4 rules were written to `bill_context_rules_A.csv` and excluded from `bill_item_reference_A_candidate.csv`.",
            "",
            "## 8. Issues and Risks",
            "",
            count_table(issue_counts) if issues else "No extraction issues were generated by structural checks.",
            "- Word extraction results are still reference candidates and require human QA before downstream use.",
            "",
            "## 9. Manual QA Checklist",
            "",
            "- Verify `bill_code_9` truly comes from GB/T 50854-2024 Appendix A.",
            "- Verify `bill_name` is correct.",
            "- Verify `unit` is correct.",
            "- Verify `project_feature_raw` is complete.",
            "- Verify `quantity_calculation_rule` is complete.",
            "- Verify `work_content_raw` is complete.",
            "- Verify no `A1-1-*` Guangdong quota codes are mixed in.",
            "- Verify all `review_status` values are `pending`.",
            "- Verify A.4 rules did not enter the bill item candidate main table.",
            "",
            "## 10. Go / No-Go Recommendation for Stage B-DOCX-2",
            "",
            "Go for Stage B-DOCX-2 after manual QA of the 12 candidates and A.4 context rules." if go else "No-Go until the issues above are resolved.",
            "",
            f"Generated at: {datetime.now().isoformat(timespec='seconds')}",
        ]
    )
    path.write_text("\n".join(lines), encoding="utf-8")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Extract GB/T 50854-2024 Appendix A bill item references from DOCX.")
    parser.add_argument("--docx", type=Path, default=DEFAULT_DOCX_PATH)
    parser.add_argument(
        "--project-root",
        type=Path,
        default=Path(__file__).resolve().parents[2],
        help="construction_cost_knowledge_engine project root",
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    docx_path = args.docx
    project_root = args.project_root
    output_dir = project_root / "data" / "private" / "reference_extraction" / "runs" / "GB50854_2024_stageB_docx_A"
    output_dir.mkdir(parents=True, exist_ok=True)

    if not docx_path.exists():
        issues = []
        make_issue(issues, str(docx_path), "", "", "", "docx_not_found", "DOCX file not found at required path.", "critical", "Place the DOCX at the required project path and rerun.")
        write_csv(output_dir / "bill_extraction_issues_A.csv", ISSUE_FIELDS, issues)
        raise SystemExit(f"DOCX not found: {docx_path}")

    source_hash = sha256_file(docx_path)
    doc = Document(str(docx_path))
    texts = para_texts(doc)
    issues: List[Dict[str, Any]] = []

    if not any(texts):
        make_issue(issues, str(docx_path), "", "", "", "no_docx_text", "DOCX has no readable paragraph text.", "critical", "Verify source document.")
    if not doc.tables:
        make_issue(issues, str(docx_path), "", "", "", "no_docx_tables", "DOCX has no readable Word tables.", "critical", "Verify source document.")
    if not has_text(texts, "附录A 土石方工程"):
        make_issue(issues, "paragraph_scan", APPENDIX_CODE, "", "", "appendix_A_not_found", "Could not locate Appendix A heading.", "critical", "Verify DOCX structure.")

    profile = build_profile(docx_path, doc, source_hash, texts)
    registry = build_registry(docx_path, source_hash, texts, doc)
    for required_registry in ["GB50854_A1_TABLE", "GB50854_A2_TABLE", "GB50854_A3_TABLE"]:
        row = next((r for r in registry if r["registry_id"] == required_registry), None)
        if not row or row["source_table_index"] == "":
            make_issue(issues, required_registry, APPENDIX_CODE, row["section_code"] if row else "", row["table_code"] if row else "", "table_not_found", "Required Appendix A table was not located.", "critical", "Verify DOCX table structure.")

    candidates = extract_candidates(docx_path, source_hash, doc, issues)
    rules = extract_rules(docx_path, source_hash, texts)
    if len(candidates) != 12:
        make_issue(issues, "Appendix A candidate count", APPENDIX_CODE, "", "", "unexpected_candidate_count", f"Expected 12 candidates, extracted {len(candidates)}.", "high", "Review table detection before Stage B-DOCX-2.")

    write_csv(output_dir / "docx_profile.csv", PROFILE_FIELDS, profile)
    write_csv(output_dir / "bill_appendix_registry_A.csv", REGISTRY_FIELDS, registry)
    write_csv(output_dir / "bill_item_reference_A_candidate.csv", CANDIDATE_FIELDS, candidates)
    write_csv(output_dir / "bill_context_rules_A.csv", RULE_FIELDS, rules)
    write_csv(output_dir / "bill_extraction_issues_A.csv", ISSUE_FIELDS, issues)
    write_report(output_dir / "stageB_docx_A_report.md", docx_path, profile, registry, candidates, rules, issues)

    print(f"profile_rows={len(profile)}")
    print(f"registry_rows={len(registry)}")
    print(f"candidate_rows={len(candidates)}")
    print(f"context_rule_rows={len(rules)}")
    print(f"issue_rows={len(issues)}")
    print("candidate_section_counts=" + json.dumps(dict(Counter(r["section_code"] for r in candidates)), ensure_ascii=False, sort_keys=True))
    print("output_dir=" + str(output_dir))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
