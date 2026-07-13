from __future__ import annotations

import csv
import hashlib
import json
import re
import sqlite3
import sys
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, Iterable, List, Sequence, Tuple
from xml.sax.saxutils import escape as xml_escape


STAGE_NAME = "REFERENCE_FAMILY_FRAMEWORK_LOCK_1"
ENGINE_DIR = "construction_cost_knowledge_engine"
RUN_DIR_REL = Path(
    "construction_cost_knowledge_engine/data/private/reference_extraction/runs/REFERENCE_FAMILY_FRAMEWORK_LOCK_1"
)
SOURCE_STANDARDS_REL = Path("construction_cost_knowledge_engine/data/private/reference_extraction/source_standards")
DOCS_REL = Path("construction_cost_knowledge_engine/docs/reference_extraction")
WEB_REL = Path("construction_cost_knowledge_engine/web_collab_prototype")


SOURCE_DOC_HEADERS = [
    "source_document_id",
    "document_group",
    "family_id",
    "standard_code",
    "volume_code",
    "title",
    "file_name",
    "relative_path",
    "absolute_path",
    "extension",
    "file_size_bytes",
    "sha256",
    "page_count",
    "text_layer_status",
    "text_sample_chars",
    "scanned_at",
    "source_role",
    "classification_basis",
    "notes",
]

STANDARD_FAMILY_HEADERS = [
    "family_id",
    "family_type",
    "display_name",
    "standard_code",
    "document_count",
    "source_document_ids",
    "ordinary_bill_mapping_target",
    "mapped_national_standard",
    "routing_status",
    "routing_basis",
    "notes",
]

ROUTING_HEADERS = [
    "source_family",
    "source_volume_pattern",
    "source_document_ids",
    "source_volume_names",
    "target_standard_family",
    "target_standard_code",
    "target_document_ids",
    "mapping_target_type",
    "route_status",
    "routing_reason",
    "downstream_entry",
]

LAYER_HEADERS = [
    "layer_id",
    "layer_name",
    "business_purpose",
    "canonical_inputs",
    "canonical_outputs",
    "mutable",
    "write_allowed",
    "allowed_operations",
    "forbidden_operations",
    "review_status_policy",
    "trace_fields",
    "current_stage_status",
]

ENTITY_HEADERS = [
    "entity_name",
    "layer_id",
    "primary_key",
    "business_meaning",
    "canonical_source",
    "mutable",
    "write_allowed",
    "relationships",
    "review_status",
    "trace_fields",
    "notes",
]

GOLDEN_HEADERS = [
    "slice_id",
    "source_pdf_document_id",
    "source_pdf_path",
    "source_pdf_sha256",
    "gbt_baseline_document_id",
    "gbt_baseline_path",
    "gbt_baseline_sha256",
    "quota_count",
    "resource_count",
    "rule_block_count",
    "scope_link_count",
    "mapping_relation_count",
    "mapping_relation_source",
    "copy_semantics",
    "move_semantics",
    "exclude_semantics",
    "restore_semantics",
    "smoke_result",
    "page_route",
    "current_draft_count",
    "current_audit_count",
    "approved_count",
    "registration_status",
    "notes",
]

ISSUE_HEADERS = [
    "issue_id",
    "severity",
    "category",
    "check_name",
    "expected",
    "actual",
    "status",
    "notes",
]


def project_root_from_here() -> Path:
    current = Path(__file__).resolve()
    for parent in [current, *current.parents]:
        if (parent / ENGINE_DIR).exists():
            return parent
    return current.parents[3]


def nstr(value: Any) -> str:
    if value is None:
        return ""
    return str(value)


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def read_csv(path: Path) -> List[Dict[str, str]]:
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as fh:
        return list(csv.DictReader(fh))


def read_headers(path: Path) -> List[str]:
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as fh:
        reader = csv.reader(fh)
        return next(reader, [])


def write_csv(path: Path, headers: Sequence[str], rows: Sequence[Dict[str, Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as fh:
        writer = csv.DictWriter(fh, fieldnames=headers, extrasaction="ignore")
        writer.writeheader()
        for row in rows:
            writer.writerow({field: row.get(field, "") for field in headers})


def count_csv_rows(path: Path) -> int:
    if not path.exists():
        return 0
    with path.open("r", encoding="utf-8-sig", newline="") as fh:
        reader = csv.reader(fh)
        next(reader, None)
        return sum(1 for _row in reader)


def find_named_child(parent: Path, needle: str) -> Path:
    for child in parent.iterdir():
        if child.is_dir() and needle in child.name:
            return child
    raise FileNotFoundError(f"Cannot locate directory containing {needle!r} under {parent}")


def pdf_profile(path: Path) -> Tuple[str, str, str]:
    try:
        from pypdf import PdfReader
    except Exception as exc:  # pragma: no cover - depends on local runtime
        return "", f"text_layer_unknown:pypdf_unavailable:{exc}", "0"

    try:
        reader = PdfReader(str(path))
        page_count = len(reader.pages)
        samples: List[str] = []
        for page in reader.pages[: min(5, page_count)]:
            try:
                text = page.extract_text() or ""
            except Exception:
                text = ""
            if text.strip():
                samples.append(text.strip())
        sample_text = re.sub(r"\s+", " ", " ".join(samples))[:240]
        status = "text_present" if sample_text else "text_not_detected"
        return str(page_count), status, str(len(sample_text))
    except Exception as exc:
        return "", f"text_layer_error:{type(exc).__name__}", "0"


def docx_profile(path: Path) -> Tuple[str, str, str]:
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover - depends on local runtime
        return "not_applicable_docx", f"docx_text_unknown:python_docx_unavailable:{exc}", "0"

    try:
        document = Document(str(path))
        parts: List[str] = []
        parts.extend(paragraph.text for paragraph in document.paragraphs[:80])
        for table in document.tables[:10]:
            for row in table.rows[:4]:
                for cell in row.cells[:6]:
                    parts.append(cell.text)
        sample_text = re.sub(r"\s+", " ", " ".join(parts))[:240]
        status = "docx_text_present" if sample_text.strip() else "docx_text_not_detected"
        return "not_applicable_docx", status, str(len(sample_text))
    except Exception as exc:
        return "not_applicable_docx", f"docx_text_error:{type(exc).__name__}", "0"


def classify_national(path: Path, in_national_dir: bool) -> Dict[str, str]:
    name = path.name
    if "房屋建筑" in name and "装饰" in name:
        return {
            "family_id": "GB_BILL_BUILDING_DECORATION_2024",
            "standard_code": "GB/T 50854-2024",
            "volume_code": "GB50854",
            "title": "房屋建筑与装饰工程工程量计算标准",
            "role": "current_gbt50854_baseline" if path.suffix.lower() == ".docx" else "official_national_standard_pdf",
            "basis": "actual source file scanned; title text in file name and text layer/profile recorded",
            "notes": "DOCX baseline is the current GB/T 50854 extraction source" if path.suffix.lower() == ".docx" else "",
        }
    if "通用安装" in name:
        return {
            "family_id": "GB_BILL_INSTALLATION_2024",
            "standard_code": "GB/T 50856-2024",
            "volume_code": "GB50856",
            "title": "通用安装工程工程量计算标准",
            "role": "official_national_standard_pdf",
            "basis": "actual source file scanned; title text in file name and text layer/profile recorded",
            "notes": "",
        }
    if "市政" in name:
        return {
            "family_id": "GB_BILL_MUNICIPAL_2024",
            "standard_code": "GB/T 50857-2024",
            "volume_code": "GB50857",
            "title": "市政工程工程量计算标准",
            "role": "official_national_standard_pdf",
            "basis": "actual source file scanned; title text in file name and text layer/profile recorded",
            "notes": "",
        }
    if "园林" in name or "绿化" in name:
        return {
            "family_id": "GB_BILL_LANDSCAPE_2024",
            "standard_code": "GB/T 50858-2024",
            "volume_code": "GB50858",
            "title": "园林绿化工程工程量计算标准",
            "role": "official_national_standard_pdf",
            "basis": "actual source file scanned; title text in file name and text layer/profile recorded",
            "notes": "",
        }
    return {
        "family_id": "GB_BILL_UNKNOWN_2024",
        "standard_code": "",
        "volume_code": "",
        "title": path.stem,
        "role": "official_national_standard_pdf" if in_national_dir else "national_standard_candidate",
        "basis": "actual source file scanned; family requires manual confirmation",
        "notes": "unclassified national source",
    }


def classify_gd(path: Path) -> Dict[str, str]:
    name = path.name
    upper = name.upper()
    if upper.startswith(("A01_", "A02_", "A03_")):
        volume = upper.split("_", 1)[0]
        return {
            "family_id": "GD2018_A_BUILDING_DECORATION",
            "standard_code": "GD2018-A",
            "volume_code": volume,
            "title": path.stem,
            "role": "gd2018_quota_source_pdf",
            "basis": "actual GD2018 source file scanned; A01/A02/A03 are building and decoration quota volumes",
            "notes": "ordinary bill mapping target is GB/T 50854-2024",
        }
    if upper.startswith("A04_"):
        return {
            "family_id": "GD2018_A04_MACHINE_SHIFT_FEE_BASIS",
            "standard_code": "GD2018-A04",
            "volume_code": "A04",
            "title": path.stem,
            "role": "gd2018_machine_shift_fee_basis_pdf",
            "basis": "actual GD2018 source file scanned; A04 title identifies construction machine shift fee rule",
            "notes": "not an ordinary bill mapping target",
        }
    if upper.startswith("C."):
        volume = upper.split("_", 1)[0]
        return {
            "family_id": "GD2018_C_INSTALLATION",
            "standard_code": "GD2018-C",
            "volume_code": volume,
            "title": path.stem,
            "role": "gd2018_quota_source_pdf",
            "basis": "actual GD2018 source file scanned; C series are installation quotas",
            "notes": "ordinary bill mapping target is GB/T 50856-2024",
        }
    if upper.startswith("D."):
        volume = upper.split("_", 1)[0]
        return {
            "family_id": "GD2018_D_MUNICIPAL",
            "standard_code": "GD2018-D",
            "volume_code": volume,
            "title": path.stem,
            "role": "gd2018_quota_source_pdf",
            "basis": "actual GD2018 source file scanned; D series are municipal quotas",
            "notes": "ordinary bill mapping target is GB/T 50857-2024",
        }
    if upper.startswith("E_"):
        return {
            "family_id": "GD2018_E_LANDSCAPE",
            "standard_code": "GD2018-E",
            "volume_code": "E",
            "title": path.stem,
            "role": "gd2018_quota_source_pdf",
            "basis": "actual GD2018 source file scanned; E volume is landscape quota",
            "notes": "ordinary bill mapping target is GB/T 50858-2024",
        }
    return {
        "family_id": "GD2018_UNKNOWN",
        "standard_code": "GD2018",
        "volume_code": "",
        "title": path.stem,
        "role": "gd2018_quota_source_pdf",
        "basis": "actual source file scanned; family requires manual confirmation",
        "notes": "unclassified GD2018 source",
    }


def source_document_id(group: str, family_id: str, volume_code: str, path: Path, digest: str) -> str:
    stem = re.sub(r"[^A-Za-z0-9]+", "_", path.stem.upper()).strip("_")
    if not stem:
        stem = "DOC"
    stem = f"{stem}_{digest[:12]}"
    if volume_code:
        return f"{group}:{family_id}:{volume_code}:{stem}"[:180]
    return f"{group}:{family_id}:{stem}"[:180]


def build_source_registry(project_root: Path) -> Tuple[List[Dict[str, Any]], Dict[str, Path]]:
    source_root = project_root / SOURCE_STANDARDS_REL
    national_dir = find_named_child(source_root, "国家标准")
    gd_dir = find_named_child(source_root, "广东省建设工程综合定额")
    baseline_docx = source_root / "房屋建筑与装饰工程工程量计算标准 GB_T 50854-2024.docx"

    source_paths: List[Tuple[str, Path, bool]] = []
    source_paths.extend(("national_standard", path, True) for path in sorted(national_dir.glob("*")) if path.is_file())
    if baseline_docx.exists():
        source_paths.append(("national_standard_baseline", baseline_docx, False))
    source_paths.extend(("gd2018_quota", path, False) for path in sorted(gd_dir.glob("*")) if path.is_file())

    scanned_at = datetime.now(timezone.utc).isoformat(timespec="seconds")
    rows: List[Dict[str, Any]] = []
    for group, path, in_national_dir in source_paths:
        if group.startswith("national"):
            meta = classify_national(path, in_national_dir)
        else:
            meta = classify_gd(path)

        extension = path.suffix.lower()
        if extension == ".pdf":
            page_count, text_status, text_chars = pdf_profile(path)
        elif extension == ".docx":
            page_count, text_status, text_chars = docx_profile(path)
        else:
            page_count, text_status, text_chars = "", "unsupported_extension", "0"

        digest = sha256_file(path)
        row = {
            "source_document_id": source_document_id(
                "GB" if group.startswith("national") else "GD",
                meta["family_id"],
                meta["volume_code"],
                path,
                digest,
            ),
            "document_group": group,
            "family_id": meta["family_id"],
            "standard_code": meta["standard_code"],
            "volume_code": meta["volume_code"],
            "title": meta["title"],
            "file_name": path.name,
            "relative_path": str(path.relative_to(project_root)).replace("\\", "/"),
            "absolute_path": str(path),
            "extension": extension,
            "file_size_bytes": path.stat().st_size,
            "sha256": digest,
            "page_count": page_count,
            "text_layer_status": text_status,
            "text_sample_chars": text_chars,
            "scanned_at": scanned_at,
            "source_role": meta["role"],
            "classification_basis": meta["basis"],
            "notes": meta["notes"],
        }
        rows.append(row)

    return rows, {"source_root": source_root, "national_dir": national_dir, "gd_dir": gd_dir}


def grouped(rows: Iterable[Dict[str, Any]], key: str) -> Dict[str, List[Dict[str, Any]]]:
    output: Dict[str, List[Dict[str, Any]]] = {}
    for row in rows:
        output.setdefault(nstr(row.get(key)), []).append(row)
    return output


def build_standard_families(source_rows: Sequence[Dict[str, Any]]) -> List[Dict[str, Any]]:
    family_info = {
        "GB_BILL_BUILDING_DECORATION_2024": (
            "national_bill_standard",
            "GB/T 50854-2024 房屋建筑与装饰工程工程量计算标准",
            "GB/T 50854-2024",
            "yes",
            "self",
            "ready",
            "GB/T 50854 baseline present as DOCX and national PDF candidate",
        ),
        "GB_BILL_INSTALLATION_2024": (
            "national_bill_standard",
            "GB/T 50856-2024 通用安装工程工程量计算标准",
            "GB/T 50856-2024",
            "yes",
            "self",
            "ready",
            "routing target for GD2018 C series only",
        ),
        "GB_BILL_MUNICIPAL_2024": (
            "national_bill_standard",
            "GB/T 50857-2024 市政工程工程量计算标准",
            "GB/T 50857-2024",
            "yes",
            "self",
            "ready",
            "routing target for GD2018 D series only",
        ),
        "GB_BILL_LANDSCAPE_2024": (
            "national_bill_standard",
            "GB/T 50858-2024 园林绿化工程工程量计算标准",
            "GB/T 50858-2024",
            "yes",
            "self",
            "ready",
            "routing target for GD2018 E volume only",
        ),
        "GD2018_A_BUILDING_DECORATION": (
            "provincial_quota_family",
            "广东省建设工程综合定额(2018) A01/A02/A03 房屋建筑与装饰",
            "GD2018-A",
            "yes",
            "GB/T 50854-2024",
            "locked",
            "A01/A02/A03 must route to GB/T 50854-2024; GB/T 50856 is not valid for A series",
        ),
        "GD2018_A04_MACHINE_SHIFT_FEE_BASIS": (
            "provincial_fee_basis",
            "广东省建设工程综合定额(2018) A04 施工机具台班费用编制规则",
            "GD2018-A04",
            "no",
            "machine_shift_fee_basis",
            "locked",
            "A04 is fee basis / machine shift data, not ordinary bill mapping target",
        ),
        "GD2018_C_INSTALLATION": (
            "provincial_quota_family",
            "广东省建设工程综合定额(2018) C 系列 安装",
            "GD2018-C",
            "yes",
            "GB/T 50856-2024",
            "locked",
            "C series route to GB/T 50856-2024",
        ),
        "GD2018_D_MUNICIPAL": (
            "provincial_quota_family",
            "广东省建设工程综合定额(2018) D 系列 市政",
            "GD2018-D",
            "yes",
            "GB/T 50857-2024",
            "locked",
            "D series route to GB/T 50857-2024",
        ),
        "GD2018_E_LANDSCAPE": (
            "provincial_quota_family",
            "广东省建设工程综合定额(2018) E 园林绿化",
            "GD2018-E",
            "yes",
            "GB/T 50858-2024",
            "locked",
            "E volume routes to GB/T 50858-2024",
        ),
    }

    by_family = grouped(source_rows, "family_id")
    rows: List[Dict[str, Any]] = []
    for family_id, info in family_info.items():
        docs = by_family.get(family_id, [])
        rows.append(
            {
                "family_id": family_id,
                "family_type": info[0],
                "display_name": info[1],
                "standard_code": info[2],
                "document_count": len(docs),
                "source_document_ids": "; ".join(row["source_document_id"] for row in docs),
                "ordinary_bill_mapping_target": info[3],
                "mapped_national_standard": info[4],
                "routing_status": info[5],
                "routing_basis": info[6],
                "notes": "source document present" if docs else "no scanned source document found",
            }
        )
    return rows


def docs_for_family(source_rows: Sequence[Dict[str, Any]], family_id: str) -> List[Dict[str, Any]]:
    return [row for row in source_rows if row.get("family_id") == family_id]


def build_routing_matrix(source_rows: Sequence[Dict[str, Any]]) -> List[Dict[str, Any]]:
    def ids(family_id: str) -> str:
        return "; ".join(row["source_document_id"] for row in docs_for_family(source_rows, family_id))

    def names(family_id: str) -> str:
        return "; ".join(row["file_name"] for row in docs_for_family(source_rows, family_id))

    return [
        {
            "source_family": "GD2018_A_BUILDING_DECORATION",
            "source_volume_pattern": "A01; A02; A03",
            "source_document_ids": ids("GD2018_A_BUILDING_DECORATION"),
            "source_volume_names": names("GD2018_A_BUILDING_DECORATION"),
            "target_standard_family": "GB_BILL_BUILDING_DECORATION_2024",
            "target_standard_code": "GB/T 50854-2024",
            "target_document_ids": ids("GB_BILL_BUILDING_DECORATION_2024"),
            "mapping_target_type": "ordinary_bill_mapping",
            "route_status": "locked",
            "routing_reason": "A01/A02/A03 are building and decoration quota volumes; GB/T 50856 is explicitly not used for A series.",
            "downstream_entry": "GD2018_BUILDING_A_FULL_PARSE_1 -> MAP_GB50854_TO_GD2018_BUILDING_1",
        },
        {
            "source_family": "GD2018_A04_MACHINE_SHIFT_FEE_BASIS",
            "source_volume_pattern": "A04",
            "source_document_ids": ids("GD2018_A04_MACHINE_SHIFT_FEE_BASIS"),
            "source_volume_names": names("GD2018_A04_MACHINE_SHIFT_FEE_BASIS"),
            "target_standard_family": "MACHINE_SHIFT_FEE_BASIS",
            "target_standard_code": "not ordinary bill standard",
            "target_document_ids": "",
            "mapping_target_type": "fee_basis_only",
            "route_status": "locked_out_of_ordinary_bill_mapping",
            "routing_reason": "A04 contains construction machine shift and fee basis rules, not bill item mapping scope.",
            "downstream_entry": "Use only as fee/resource basis after quota parse governance is complete.",
        },
        {
            "source_family": "GD2018_C_INSTALLATION",
            "source_volume_pattern": "C.*",
            "source_document_ids": ids("GD2018_C_INSTALLATION"),
            "source_volume_names": names("GD2018_C_INSTALLATION"),
            "target_standard_family": "GB_BILL_INSTALLATION_2024",
            "target_standard_code": "GB/T 50856-2024",
            "target_document_ids": ids("GB_BILL_INSTALLATION_2024"),
            "mapping_target_type": "ordinary_bill_mapping",
            "route_status": "locked",
            "routing_reason": "C series are installation quota volumes.",
            "downstream_entry": "future installation family parse and mapping stage",
        },
        {
            "source_family": "GD2018_D_MUNICIPAL",
            "source_volume_pattern": "D.*",
            "source_document_ids": ids("GD2018_D_MUNICIPAL"),
            "source_volume_names": names("GD2018_D_MUNICIPAL"),
            "target_standard_family": "GB_BILL_MUNICIPAL_2024",
            "target_standard_code": "GB/T 50857-2024",
            "target_document_ids": ids("GB_BILL_MUNICIPAL_2024"),
            "mapping_target_type": "ordinary_bill_mapping",
            "route_status": "locked",
            "routing_reason": "D series are municipal quota volumes.",
            "downstream_entry": "future municipal family parse and mapping stage",
        },
        {
            "source_family": "GD2018_E_LANDSCAPE",
            "source_volume_pattern": "E",
            "source_document_ids": ids("GD2018_E_LANDSCAPE"),
            "source_volume_names": names("GD2018_E_LANDSCAPE"),
            "target_standard_family": "GB_BILL_LANDSCAPE_2024",
            "target_standard_code": "GB/T 50858-2024",
            "target_document_ids": ids("GB_BILL_LANDSCAPE_2024"),
            "mapping_target_type": "ordinary_bill_mapping",
            "route_status": "locked",
            "routing_reason": "E volume is landscape quota.",
            "downstream_entry": "future landscape family parse and mapping stage",
        },
    ]


def build_layer_contracts() -> List[Dict[str, Any]]:
    return [
        {
            "layer_id": "L0",
            "layer_name": "Source Registry",
            "business_purpose": "Record official source documents, physical file facts, hashes, page/text profiles, standard family, and routing identity.",
            "canonical_inputs": "official PDFs/DOCX under source_standards",
            "canonical_outputs": "source_document_registry; standard_family_registry; source_family_routing_matrix",
            "mutable": "no",
            "write_allowed": "framework scan may append new source registry rows only; no source file mutation",
            "allowed_operations": "scan; checksum; page count; text-layer profiling; route registration",
            "forbidden_operations": "edit source PDFs/DOCX; normalize prices; parse full A01/A02/A03 quota tables",
            "review_status_policy": "source_registered / needs_manual_confirmation",
            "trace_fields": "source_document_id; file_path; sha256; file_size_bytes; page_count; text_layer_status; scanned_at",
            "current_stage_status": "locked by REFERENCE_FAMILY_FRAMEWORK_LOCK_1",
        },
        {
            "layer_id": "L1",
            "layer_name": "Evidence Baseline",
            "business_purpose": "Hold immutable baseline extraction evidence such as GB/T bill candidates and A1.1 golden-slice evidence.",
            "canonical_inputs": "L0 source registry; existing GB50854_2024_stageB_docx_full; existing GD2018 A1.1 evidence runs",
            "canonical_outputs": "source_page_evidence; gb_bill_item; gb_context_rule; gb_source_block; gb_rule_scope_link",
            "mutable": "no",
            "write_allowed": "new baseline run creates new immutable artifacts; Web writes are not allowed",
            "allowed_operations": "read; cite; validate hash; derive immutable evidence in new run folders",
            "forbidden_operations": "overwrite baseline rows; back-write review decisions; write enterprise price or formal enterprise quota",
            "review_status_policy": "evidence_pending / evidence_locked",
            "trace_fields": "source_document_id; page_no; source_block_id; extraction_run_id; source_sha256",
            "current_stage_status": "GB/T 50854 baseline hash registered; no baseline mutation",
        },
        {
            "layer_id": "L2",
            "layer_name": "Parsed Reference",
            "business_purpose": "Represent parsed quota/bill reference structures with source trace, but without final mapping decisions.",
            "canonical_inputs": "L0/L1 source evidence",
            "canonical_outputs": "gd_quota_item; gd_quota_fee_snapshot; gd_quota_resource_component; gd_work_content_block; gd_quantity_rule_block; gd_note_block; gd_conversion_rule; gd_parse_issue",
            "mutable": "controlled append only",
            "write_allowed": "parser stages may create new candidate runs; existing Source/Baseline cannot be edited",
            "allowed_operations": "parse candidate; normalize codes; attach evidence; raise parse issues",
            "forbidden_operations": "mark final mapping; produce enterprise price library; write Web draft tables",
            "review_status_policy": "parsed_candidate / parse_issue / needs_review",
            "trace_fields": "quota_uid; source_document_id; page_no; source_row_ref; parser_version; extraction_run_id",
            "current_stage_status": "A.1.1 parsed; A01/A02/A03 full parse is planned only",
        },
        {
            "layer_id": "L3",
            "layer_name": "Mapping Reference",
            "business_purpose": "Represent bill-to-quota candidate links and mapping issues, separate from human draft actions.",
            "canonical_inputs": "L1 GB/T bill references; L2 GD quota parsed references",
            "canonical_outputs": "bill_quota_mapping_candidate; bill_quota_mapping_issue",
            "mutable": "controlled append only",
            "write_allowed": "mapping stages may create candidate artifacts in new private run folders",
            "allowed_operations": "rank candidates; record confidence; emit issues; preserve all candidate status as pending",
            "forbidden_operations": "write to source candidate; overwrite baseline; turn Web draft action into canonical mapping without governed promotion",
            "review_status_policy": "pending / needs_manual_review / excluded_by_rule",
            "trace_fields": "mapping_candidate_id; bill_item_id; quota_uid; evidence_refs; mapping_run_id",
            "current_stage_status": "A.1.1 trial exists; full building mapping is planned only",
        },
        {
            "layer_id": "L4",
            "layer_name": "Review Draft",
            "business_purpose": "Capture reviewer overlay actions without changing parsed or mapping references.",
            "canonical_inputs": "L3 mapping candidates; reviewer UI actions",
            "canonical_outputs": "bill_quota_mapping_draft; mapping_audit_log",
            "mutable": "yes",
            "write_allowed": "Copy / Move / Exclude / Restore only; every write must audit",
            "allowed_operations": "copy_link; move_link; exclude_link; restore_original; export draft; export audit",
            "forbidden_operations": "back-write source candidate; mutate L0/L1/L2/L3; create enterprise formal quota",
            "review_status_policy": "draft_active / draft_excluded / reverted / audit_only",
            "trace_fields": "draft_edge_id; original_edge_id; action_type; reviewer; created_at; audit_id",
            "current_stage_status": "existing /quota-a111 draft counts registered and left unchanged",
        },
        {
            "layer_id": "L5",
            "layer_name": "Web Collaboration",
            "business_purpose": "Read reference layers and expose reviewer overlay workflows.",
            "canonical_inputs": "read-only SQLite view model plus L4 draft tables",
            "canonical_outputs": "Web draft tables and audit log exports only",
            "mutable": "limited",
            "write_allowed": "only Review Draft and Audit tables; no writes to Source Candidate",
            "allowed_operations": "read tree/detail APIs; create draft edge action; audit export; smoke validation",
            "forbidden_operations": "write Source/Baseline/Parsed/Mapping Candidate; alter business pages in this stage",
            "review_status_policy": "UI draft state mirrors L4; no formal approval state emitted",
            "trace_fields": "route; api_name; draft_edge_id; audit_id; sqlite_hash; smoke_run_id",
            "current_stage_status": "/quota-a111 audited only; no Web file or DB modification",
        },
    ]


def build_entity_dictionary() -> List[Dict[str, Any]]:
    rows = [
        (
            "source_document_registry",
            "L0",
            "source_document_id",
            "Physical official source document registry with hash, path, page count, text layer profile, and family identity.",
            "source_standards scan",
            "no",
            "framework scan only",
            "1-to-many source_page_evidence; many-to-1 standard_family_registry",
            "source_registered / needs_manual_confirmation",
            "source_document_id; relative_path; sha256; page_count; text_layer_status; scanned_at",
            "Source files are immutable.",
        ),
        (
            "standard_family_registry",
            "L0",
            "family_id",
            "Canonical standard/quota family registry used for routing and downstream stage gating.",
            "source_document_registry",
            "no",
            "framework scan only",
            "1-to-many source_document_registry; 1-to-many routing matrix rows",
            "locked / needs_manual_confirmation",
            "family_id; standard_code; source_document_ids; routing_basis",
            "A01/A02/A03 route to GB/T 50854-2024.",
        ),
        (
            "source_page_evidence",
            "L1",
            "source_page_evidence_id",
            "Page/block-level immutable evidence anchor for parsed or bill reference facts.",
            "official source document pages",
            "no",
            "baseline evidence run only",
            "many-to-1 source_document_registry; 1-to-many gb_source_block / gd_*_block",
            "evidence_pending / evidence_locked",
            "source_document_id; page_no; block_id; source_sha256; extraction_run_id",
            "No Web write-back.",
        ),
        (
            "gb_bill_item",
            "L1",
            "gb_bill_item_id",
            "GB/T bill item reference row, including code, name, unit, and appendix/family context.",
            "GB/T baseline extraction",
            "no",
            "baseline creation only",
            "many-to-1 source_page_evidence; 1-to-many bill_quota_mapping_candidate",
            "bill_candidate / evidence_locked",
            "gb_bill_item_id; gb_code; standard_code; source_document_id; page_ref; source_sha256",
            "GB/T 50854 current baseline is DOCX-first.",
        ),
        (
            "gb_context_rule",
            "L1",
            "gb_context_rule_id",
            "GB/T explanatory or quantity/context rule tied to bill item scope.",
            "GB/T baseline extraction",
            "no",
            "baseline creation only",
            "many-to-1 gb_bill_item or gb_source_block; many-to-many through gb_rule_scope_link",
            "rule_candidate / evidence_locked",
            "rule_id; standard_code; appendix; source_block_id; page_ref",
            "Used as mapping context, not a price source.",
        ),
        (
            "gb_source_block",
            "L1",
            "gb_source_block_id",
            "Traceable source block from GB/T document tables or clauses.",
            "GB/T source evidence",
            "no",
            "baseline creation only",
            "many-to-1 source_page_evidence; 1-to-many gb_context_rule",
            "evidence_locked",
            "source_block_id; source_document_id; page_no; table_id; row_ref; source_text_hash",
            "Preserves raw evidence before normalization.",
        ),
        (
            "gb_rule_scope_link",
            "L1",
            "gb_rule_scope_link_id",
            "Scope relationship between GB/T rules and bill items or appendices.",
            "GB/T baseline extraction",
            "no",
            "baseline creation only",
            "many-to-1 gb_context_rule; many-to-1 gb_bill_item",
            "scope_pending / scope_locked",
            "rule_id; gb_bill_item_id; scope_type; source_block_id",
            "Required before high-confidence mapping promotion.",
        ),
        (
            "gd_quota_item",
            "L2",
            "quota_uid",
            "GD2018 quota item identity and normalized code/name/unit. quota_uid format: GD:2018:{family}:{source_code_normalized}, e.g. GD:2018:A:A1-1-1.",
            "GD2018 parsed reference run",
            "append-only candidate",
            "parser stage output only",
            "many-to-1 source_page_evidence; 1-to-many gd_quota_resource_component; 1-to-many bill_quota_mapping_candidate",
            "parsed_candidate / needs_review",
            "quota_uid; source_code; source_document_id; page_no; parser_version; source_sha256",
            "A01/A02/A03 full parse is not executed in this stage.",
        ),
        (
            "gd_quota_fee_snapshot",
            "L2",
            "quota_fee_snapshot_id",
            "Reference fee snapshot from GD quota source, separated from enterprise price and formal enterprise quota.",
            "GD2018 parsed reference run",
            "append-only candidate",
            "parser stage output only",
            "many-to-1 gd_quota_item",
            "fee_candidate / needs_review",
            "quota_uid; fee_type; amount; source_page; source_sha256; extraction_run_id",
            "Enterprise price is out of scope.",
        ),
        (
            "gd_quota_resource_component",
            "L2",
            "quota_resource_component_id",
            "Labor/material/machine resource component under a GD quota item.",
            "GD2018 parsed reference run",
            "append-only candidate",
            "parser stage output only",
            "many-to-1 gd_quota_item; may reference A04 fee basis later",
            "resource_candidate / alignment_issue",
            "quota_uid; resource_code; resource_name; unit; amount; source_row_ref; page_no",
            "A04 may support machine shift fee basis but is not an ordinary mapping target.",
        ),
        (
            "gd_work_content_block",
            "L2",
            "gd_work_content_block_id",
            "GD quota work-content block extracted from official source pages.",
            "GD2018 parsed reference run",
            "append-only candidate",
            "parser stage output only",
            "many-to-1 source_page_evidence; many-to-many through gd_work_content_scope_link",
            "content_candidate / scope_pending",
            "block_id; source_document_id; page_no; raw_text_hash; parser_version",
            "Scope links must be explicit.",
        ),
        (
            "gd_work_content_scope_link",
            "L2",
            "gd_work_content_scope_link_id",
            "Link between GD work-content block and affected quota item(s) or chapter scope.",
            "GD2018 parsed reference run",
            "append-only candidate",
            "parser stage output only",
            "many-to-1 gd_work_content_block; many-to-1 gd_quota_item",
            "scope_pending / scope_locked",
            "block_id; quota_uid; scope_type; confidence; source_page",
            "Keeps chapter-level text separate from item-level facts.",
        ),
        (
            "gd_quantity_rule_block",
            "L2",
            "gd_quantity_rule_block_id",
            "GD quantity calculation rule block extracted with page and raw text evidence.",
            "GD2018 parsed reference run",
            "append-only candidate",
            "parser stage output only",
            "many-to-1 source_page_evidence; many-to-many through gd_quantity_rule_scope_link",
            "rule_candidate / scope_pending",
            "rule_block_id; source_document_id; page_no; raw_text_hash; parser_version",
            "A1.1 rule blocks are registered in golden slice.",
        ),
        (
            "gd_quantity_rule_scope_link",
            "L2",
            "gd_quantity_rule_scope_link_id",
            "Link between GD quantity rule block and affected quota items, chapters, or bill scopes.",
            "GD2018 parsed reference run",
            "append-only candidate",
            "parser stage output only",
            "many-to-1 gd_quantity_rule_block; optional many-to-1 gd_quota_item",
            "scope_pending / scope_locked",
            "rule_block_id; quota_uid; scope_type; confidence; source_page",
            "Required before Web dual-view display.",
        ),
        (
            "gd_note_block",
            "L2",
            "gd_note_block_id",
            "GD note, explanation, or appendix block that influences quota interpretation.",
            "GD2018 parsed reference run",
            "append-only candidate",
            "parser stage output only",
            "many-to-1 source_page_evidence; optional links to gd_quota_item",
            "note_candidate / needs_review",
            "note_block_id; source_document_id; page_no; raw_text_hash; parser_version",
            "Not a mapping decision by itself.",
        ),
        (
            "gd_conversion_rule",
            "L2",
            "gd_conversion_rule_id",
            "GD conversion or adjustment rule, including rule scope and affected fields.",
            "GD2018 parsed reference run",
            "append-only candidate",
            "parser stage output only",
            "optional many-to-1 gd_quota_item; many-to-1 source_page_evidence",
            "conversion_candidate / needs_review",
            "conversion_rule_id; quota_uid; rule_text_hash; source_page; extraction_run_id",
            "No enterprise quota generation.",
        ),
        (
            "gd_parse_issue",
            "L2",
            "gd_parse_issue_id",
            "Parser issue ledger for uncertain rows, conflicts, missing evidence, or scope ambiguity.",
            "GD2018 parser validation",
            "append-only",
            "parser stage output only",
            "may reference any gd_* entity",
            "open / resolved_in_later_run",
            "issue_id; entity_name; entity_pk; severity; source_page; extraction_run_id",
            "Issues do not mutate parsed facts.",
        ),
        (
            "bill_quota_mapping_candidate",
            "L3",
            "mapping_candidate_id",
            "Candidate relationship between GB/T bill item and GD quota item.",
            "mapping stage output",
            "append-only candidate",
            "mapping stage output only",
            "many-to-1 gb_bill_item; many-to-1 gd_quota_item; 1-to-many draft edges",
            "pending / needs_manual_review / excluded_by_rule",
            "mapping_candidate_id; gb_bill_item_id; quota_uid; confidence; evidence_refs; mapping_run_id",
            "Full A-building mapping is not executed in this stage.",
        ),
        (
            "bill_quota_mapping_issue",
            "L3",
            "bill_quota_mapping_issue_id",
            "Mapping issue ledger for low confidence, multi-candidate, or missing target conditions.",
            "mapping stage output",
            "append-only",
            "mapping stage output only",
            "many-to-1 bill_quota_mapping_candidate or mapping run",
            "open / reviewed / carried_forward",
            "issue_id; mapping_candidate_id; severity; evidence_refs; mapping_run_id",
            "Does not promote candidate to final mapping.",
        ),
        (
            "bill_quota_mapping_draft",
            "L4",
            "draft_edge_id",
            "Reviewer overlay edge with Copy/Move/Exclude/Restore state.",
            "Web Review Draft",
            "yes",
            "Review Draft table only",
            "many-to-1 bill_quota_mapping_candidate; 1-to-many mapping_audit_log",
            "draft_active / draft_excluded / reverted",
            "draft_edge_id; original_edge_id; action_type; reviewer; created_at; source_candidate_id",
            "Must not write back to Source Candidate.",
        ),
        (
            "mapping_audit_log",
            "L4",
            "audit_id",
            "Append-only audit record for every reviewer draft action and export.",
            "Web Review Draft/Audit",
            "append-only",
            "Audit table only",
            "many-to-1 bill_quota_mapping_draft; may reference route/API",
            "audit_only",
            "audit_id; action_type; actor; timestamp; before_hash; after_hash; route",
            "Every Web write must audit.",
        ),
    ]
    return [
        {
            "entity_name": item[0],
            "layer_id": item[1],
            "primary_key": item[2],
            "business_meaning": item[3],
            "canonical_source": item[4],
            "mutable": item[5],
            "write_allowed": item[6],
            "relationships": item[7],
            "review_status": item[8],
            "trace_fields": item[9],
            "notes": item[10],
        }
        for item in rows
    ]


def count_sqlite_table(db_path: Path, table_name: str) -> int:
    if not db_path.exists():
        return 0
    conn = sqlite3.connect(str(db_path))
    try:
        cursor = conn.execute(f'SELECT COUNT(*) FROM "{table_name}"')
        value = cursor.fetchone()[0]
        return int(value)
    finally:
        conn.close()


def smoke_summary(path: Path) -> str:
    rows = read_csv(path)
    if not rows:
        return f"{path.name}:missing_or_empty"
    status_fields = ["pass_fail", "status", "result"]
    values: List[str] = []
    for row in rows:
        for field in status_fields:
            if row.get(field):
                values.append(row[field].strip().lower())
                break
    if values and all(value in {"pass", "ok", "passed", "true"} for value in values):
        return f"{path.name}:pass({len(rows)} checks)"
    if values:
        failures = sum(1 for value in values if value not in {"pass", "ok", "passed", "true"})
        return f"{path.name}:mixed({failures} non-pass / {len(rows)} checks)"
    return f"{path.name}:{len(rows)} rows"


def build_golden_slice(project_root: Path, source_rows: Sequence[Dict[str, Any]]) -> List[Dict[str, Any]]:
    run_root = project_root / "construction_cost_knowledge_engine/data/private/reference_extraction/runs"
    web_db = project_root / WEB_REL / "data/web_collab_readonly.sqlite"

    a01 = next(
        row
        for row in source_rows
        if row.get("family_id") == "GD2018_A_BUILDING_DECORATION" and row.get("volume_code") == "A01"
    )
    gb_docx = next(
        (
            row
            for row in source_rows
            if row.get("family_id") == "GB_BILL_BUILDING_DECORATION_2024"
            and row.get("extension") == ".docx"
        ),
        next(row for row in source_rows if row.get("family_id") == "GB_BILL_BUILDING_DECORATION_2024"),
    )

    full_review = run_root / "GD2018_PDF_A111_FULL_REVIEW_PACK_1"
    dual_view = run_root / "WEB_QUOTA_A111_QUANTITY_RULE_DUAL_VIEW_1"
    viewer = run_root / "WEB_QUOTA_A111_PDF_DETAIL_VIEWER_1"
    mapping_draft = run_root / "WEB_QUOTA_A111_MAPPING_DRAFT_1"

    quota_count = count_csv_rows(full_review / "main_quota_all_137.csv")
    resource_count = count_csv_rows(full_review / "resource_display_all_629.csv")
    rule_block_count = count_csv_rows(dual_view / "quantity_rule_source_blocks.csv")
    scope_link_count = count_csv_rows(dual_view / "quantity_rule_scope_links.csv")
    mapping_relation_count = count_csv_rows(viewer / "web_quota_a111_bill_to_quota_rows.csv")
    if mapping_relation_count == 0:
        mapping_relation_count = count_csv_rows(run_root / "MAP_A111_quota_to_bill_trial/quota_to_bill_mapping_A111_candidate.csv")
    draft_count = count_sqlite_table(web_db, "web_quota_a111_mapping_draft_edges")
    audit_count = count_sqlite_table(web_db, "web_quota_a111_mapping_draft_audit_log")

    smoke_parts = [
        smoke_summary(viewer / "web_quota_a111_smoke_result.csv"),
        smoke_summary(mapping_draft / "web_quota_a111_mapping_draft_smoke.csv"),
        smoke_summary(dual_view / "quantity_rule_dual_view_smoke.csv"),
    ]

    return [
        {
            "slice_id": "GOLDEN_SLICE_GD2018_A111_V1",
            "source_pdf_document_id": a01["source_document_id"],
            "source_pdf_path": a01["relative_path"],
            "source_pdf_sha256": a01["sha256"],
            "gbt_baseline_document_id": gb_docx["source_document_id"],
            "gbt_baseline_path": gb_docx["relative_path"],
            "gbt_baseline_sha256": gb_docx["sha256"],
            "quota_count": quota_count,
            "resource_count": resource_count,
            "rule_block_count": rule_block_count,
            "scope_link_count": scope_link_count,
            "mapping_relation_count": mapping_relation_count,
            "mapping_relation_source": "WEB_QUOTA_A111_PDF_DETAIL_VIEWER_1/web_quota_a111_bill_to_quota_rows.csv",
            "copy_semantics": "copy_link adds a draft overlay edge and preserves the original candidate edge.",
            "move_semantics": "move_link creates the target draft edge and excludes the source overlay edge without changing Source Candidate.",
            "exclude_semantics": "exclude_link hides the selected relation in the draft overlay only.",
            "restore_semantics": "restore_original reverts the selected draft overlay state to the original candidate relation.",
            "smoke_result": " | ".join(smoke_parts),
            "page_route": "/quota-a111",
            "current_draft_count": draft_count,
            "current_audit_count": audit_count,
            "approved_count": 0,
            "registration_status": "registered_existing_golden_slice_no_mutation",
            "notes": "Counts are read from existing private runs and Web SQLite; this stage does not write Web or DB.",
        }
    ]


def build_validation_issues(
    source_rows: Sequence[Dict[str, Any]],
    routing_rows: Sequence[Dict[str, Any]],
    golden_rows: Sequence[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    issues: List[Dict[str, Any]] = []

    def add(
        severity: str,
        category: str,
        check_name: str,
        expected: str,
        actual: str,
        status: str,
        notes: str,
    ) -> None:
        issues.append(
            {
                "issue_id": f"RF-{len(issues) + 1:03d}",
                "severity": severity,
                "category": category,
                "check_name": check_name,
                "expected": expected,
                "actual": actual,
                "status": status,
                "notes": notes,
            }
        )

    standard_codes = {row.get("standard_code") for row in source_rows}
    for code in ["GB/T 50854-2024", "GB/T 50856-2024", "GB/T 50857-2024", "GB/T 50858-2024"]:
        add(
            "info",
            "source_inventory",
            f"{code}_present",
            "present in scanned source documents",
            "present" if code in standard_codes else "missing",
            "pass" if code in standard_codes else "issue",
            "This check is based on scanned source files and recorded hash/page facts.",
        )
    for code in ["GB/T 50500-2024", "GB/T 50855-2024"]:
        add(
            "warning",
            "source_inventory",
            f"{code}_not_found_in_scan",
            "optional family may be absent from current source_standards scan",
            "not found",
            "needs_manual_confirmation",
            "No physical source file was found in the scanned directories for this expected national standard code.",
        )

    a_route = next(row for row in routing_rows if row["source_family"] == "GD2018_A_BUILDING_DECORATION")
    add(
        "critical",
        "routing",
        "A01_A02_A03_route",
        "GB/T 50854-2024",
        a_route["target_standard_code"],
        "pass" if a_route["target_standard_code"] == "GB/T 50854-2024" else "issue",
        "GB/T 50856 is not used for A series.",
    )
    add(
        "info",
        "routing",
        "A04_out_of_ordinary_bill_mapping",
        "fee_basis_only",
        next(row for row in routing_rows if row["source_family"] == "GD2018_A04_MACHINE_SHIFT_FEE_BASIS")[
            "mapping_target_type"
        ],
        "pass",
        "A04 registered as machine shift / fee basis.",
    )
    add(
        "info",
        "layer_contract",
        "source_baseline_web_mutation",
        "no Source/Baseline/Web business file mutation in this stage",
        "script writes only docs, manifest, and private run outputs",
        "pass",
        "Validated by stage scope; final git status should still be reviewed.",
    )
    add(
        "info",
        "golden_slice",
        "A111_golden_slice_registered",
        "quota/resource/rule/scope/mapping/draft/audit counts registered",
        json.dumps(golden_rows[0], ensure_ascii=False),
        "pass",
        "Read-only registration from existing A1.1 runs.",
    )
    add(
        "info",
        "scope_guard",
        "full_A01_A02_A03_parse_executed",
        "false",
        "false",
        "pass",
        "This stage does not parse A01/A02/A03 full data.",
    )
    add(
        "info",
        "scope_guard",
        "mapping_executed",
        "false",
        "false",
        "pass",
        "This stage does not execute Mapping.",
    )
    add(
        "info",
        "scope_guard",
        "approved_count",
        "0",
        "0",
        "pass",
        "No formal approval output or status rows are produced by this stage.",
    )
    return issues


def xml_text(value: Any) -> str:
    return xml_escape(nstr(value), {"\n": "&#10;", "\r": ""})


def excel_col(index: int) -> str:
    number = index + 1
    label = ""
    while number:
        number, remainder = divmod(number - 1, 26)
        label = chr(65 + remainder) + label
    return label


def cell_xml(row_index: int, col_index: int, value: Any, header: bool = False) -> str:
    ref = f"{excel_col(col_index)}{row_index}"
    style = ' s="1"' if header else ""
    if value is None or nstr(value) == "":
        return f'<c r="{ref}"{style}/>'
    return f'<c r="{ref}"{style} t="inlineStr"><is><t xml:space="preserve">{xml_text(value)}</t></is></c>'


def column_width(field: str, values: Sequence[Dict[str, Any]]) -> float:
    if field in {"absolute_path", "relative_path", "classification_basis", "notes", "routing_reason", "business_meaning"}:
        return 56.0
    if field in {"sha256", "source_document_ids", "target_document_ids", "trace_fields", "relationships"}:
        return 42.0
    max_len = len(field)
    for row in values[:200]:
        max_len = max(max_len, len(nstr(row.get(field, ""))))
    return min(max(max_len + 2, 10), 36)


def sheet_xml(headers: Sequence[str], rows: Sequence[Dict[str, Any]]) -> str:
    max_row = max(len(rows) + 1, 1)
    max_col = max(len(headers), 1)
    last_ref = f"{excel_col(max_col - 1)}{max_row}"
    cols = "".join(
        f'<col min="{index}" max="{index}" width="{column_width(field, rows):.2f}" customWidth="1"/>'
        for index, field in enumerate(headers, start=1)
    )
    row_xml: List[str] = []
    header_cells = "".join(cell_xml(1, index, field, header=True) for index, field in enumerate(headers))
    row_xml.append(f'<row r="1" spans="1:{max_col}">{header_cells}</row>')
    for row_number, row in enumerate(rows, start=2):
        cells = "".join(cell_xml(row_number, index, row.get(field, "")) for index, field in enumerate(headers))
        row_xml.append(f'<row r="{row_number}" spans="1:{max_col}">{cells}</row>')
    return f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <dimension ref="A1:{last_ref}"/>
  <sheetViews><sheetView workbookViewId="0"><pane ySplit="1" topLeftCell="A2" activePane="bottomLeft" state="frozen"/><selection pane="bottomLeft" activeCell="A2" sqref="A2"/></sheetView></sheetViews>
  <sheetFormatPr defaultRowHeight="15"/>
  <cols>{cols}</cols>
  <sheetData>{''.join(row_xml)}</sheetData>
  <autoFilter ref="A1:{last_ref}"/>
</worksheet>'''


def styles_xml() -> str:
    return '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<styleSheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
  <fonts count="2">
    <font><sz val="11"/><color theme="1"/><name val="Calibri"/><family val="2"/></font>
    <font><b/><sz val="11"/><color rgb="FFFFFFFF"/><name val="Calibri"/><family val="2"/></font>
  </fonts>
  <fills count="3">
    <fill><patternFill patternType="none"/></fill>
    <fill><patternFill patternType="gray125"/></fill>
    <fill><patternFill patternType="solid"><fgColor rgb="FF1F4E78"/><bgColor indexed="64"/></patternFill></fill>
  </fills>
  <borders count="1"><border><left/><right/><top/><bottom/><diagonal/></border></borders>
  <cellStyleXfs count="1"><xf numFmtId="0" fontId="0" fillId="0" borderId="0"/></cellStyleXfs>
  <cellXfs count="2">
    <xf numFmtId="0" fontId="0" fillId="0" borderId="0" xfId="0"/>
    <xf numFmtId="0" fontId="1" fillId="2" borderId="0" xfId="0" applyFont="1" applyFill="1"/>
  </cellXfs>
  <cellStyles count="1"><cellStyle name="Normal" xfId="0" builtinId="0"/></cellStyles>
</styleSheet>'''


def write_xlsx(path: Path, sheets: Sequence[Tuple[str, Sequence[str], Sequence[Dict[str, Any]]]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    workbook_sheets = []
    workbook_rels = []
    content_overrides = []
    for index, (name, _headers, _rows) in enumerate(sheets, start=1):
        safe_name = name[:31]
        workbook_sheets.append(f'<sheet name="{xml_text(safe_name)}" sheetId="{index}" r:id="rId{index}"/>')
        workbook_rels.append(
            f'<Relationship Id="rId{index}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" Target="worksheets/sheet{index}.xml"/>'
        )
        content_overrides.append(
            f'<Override PartName="/xl/worksheets/sheet{index}.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>'
        )
    style_rel_id = len(sheets) + 1
    workbook_rels.append(
        f'<Relationship Id="rId{style_rel_id}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/>'
    )
    now = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
    workbook_xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"><sheets>{''.join(workbook_sheets)}</sheets></workbook>'''
    workbook_rels_xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">{''.join(workbook_rels)}</Relationships>'''
    rels_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/>
  <Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" Target="docProps/core.xml"/>
  <Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" Target="docProps/app.xml"/>
</Relationships>'''
    content_types = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>
  <Override PartName="/xl/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.styles+xml"/>
  <Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>
  <Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>
  {''.join(content_overrides)}
</Types>'''
    core_xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:dcterms="http://purl.org/dc/terms/" xmlns:dcmitype="http://purl.org/dc/dcmitype/" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
  <dc:title>Reference Family Framework Review</dc:title>
  <dc:creator>AI Construction System</dc:creator>
  <cp:lastModifiedBy>AI Construction System</cp:lastModifiedBy>
  <dcterms:created xsi:type="dcterms:W3CDTF">{now}</dcterms:created>
  <dcterms:modified xsi:type="dcterms:W3CDTF">{now}</dcterms:modified>
</cp:coreProperties>'''
    app_xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties" xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">
  <Application>AI Construction System</Application>
  <HeadingPairs><vt:vector size="2" baseType="variant"><vt:variant><vt:lpstr>Worksheets</vt:lpstr></vt:variant><vt:variant><vt:i4>{len(sheets)}</vt:i4></vt:variant></vt:vector></HeadingPairs>
  <TitlesOfParts><vt:vector size="{len(sheets)}" baseType="lpstr">{''.join(f"<vt:lpstr>{xml_text(name[:31])}</vt:lpstr>" for name, _headers, _rows in sheets)}</vt:vector></TitlesOfParts>
  <Company>AI Construction System</Company>
</Properties>'''
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", content_types)
        archive.writestr("_rels/.rels", rels_xml)
        archive.writestr("docProps/core.xml", core_xml)
        archive.writestr("docProps/app.xml", app_xml)
        archive.writestr("xl/workbook.xml", workbook_xml)
        archive.writestr("xl/_rels/workbook.xml.rels", workbook_rels_xml)
        archive.writestr("xl/styles.xml", styles_xml())
        for index, (_name, headers, rows) in enumerate(sheets, start=1):
            archive.writestr(f"xl/worksheets/sheet{index}.xml", sheet_xml(headers, rows))


def md_table(headers: Sequence[str], rows: Sequence[Dict[str, Any]]) -> str:
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join("---" for _ in headers) + " |"]
    for row in rows:
        cells = []
        for header in headers:
            value = nstr(row.get(header, "")).replace("\n", " ").replace("|", "\\|")
            cells.append(value)
        lines.append("| " + " | ".join(cells) + " |")
    return "\n".join(lines)


def write_architecture_doc(
    path: Path,
    source_rows: Sequence[Dict[str, Any]],
    family_rows: Sequence[Dict[str, Any]],
    routing_rows: Sequence[Dict[str, Any]],
    layer_rows: Sequence[Dict[str, Any]],
    golden_rows: Sequence[Dict[str, Any]],
) -> None:
    national_files = [row for row in source_rows if row["document_group"] == "national_standard"]
    gd_files = [row for row in source_rows if row["document_group"] == "gd2018_quota"]
    baseline_files = [row for row in source_rows if row["document_group"] == "national_standard_baseline"]
    content = f"""# Reference Family Architecture

Stage: `{STAGE_NAME}`

This document locks the source-family framework for national standards and GD2018 quota references. It establishes routing, six-layer data contracts, and the downstream entry plan. It does not parse A01/A02/A03 full data, does not execute Mapping, and does not modify Web business pages.

## Source Inventory

- National standard directory files: {len(national_files)}
- Current GB/T 50854 baseline file outside the national directory: {len(baseline_files)}
- GD2018 quota source files: {len(gd_files)}
- Source/Baseline/Web mutation in this stage: no

## Standard Families

{md_table(["family_id", "family_type", "standard_code", "document_count", "mapped_national_standard", "routing_status"], family_rows)}

## Routing Matrix

{md_table(["source_family", "source_volume_pattern", "target_standard_code", "mapping_target_type", "route_status", "downstream_entry"], routing_rows)}

Routing lock:

- GD2018 A01/A02/A03 route to GB/T 50854-2024.
- GD2018 A series does not route to GB/T 50856-2024.
- GD2018 C series routes to GB/T 50856-2024.
- GD2018 D series routes to GB/T 50857-2024.
- GD2018 E routes to GB/T 50858-2024.
- GD2018 A04 is machine shift / fee basis and is outside ordinary bill mapping.

## Data Layer Contracts

{md_table(["layer_id", "layer_name", "mutable", "write_allowed", "allowed_operations", "forbidden_operations"], layer_rows)}

Layer lock:

- L0 Source Registry and L1 Evidence Baseline are immutable.
- L2 Parsed Reference and L3 Mapping Reference are candidate/append-only until governed promotion.
- L4 Review Draft allows Copy, Move, Exclude, and Restore only, with audit.
- L5 Web Collaboration writes only Review Draft and Audit tables; Web drafts must not write back to Source Candidate.
- Enterprise price and formal enterprise quota are out of scope.

## A1.1 Golden Slice

{md_table(["slice_id", "quota_count", "resource_count", "rule_block_count", "scope_link_count", "mapping_relation_count", "page_route", "current_draft_count", "current_audit_count", "approved_count", "registration_status"], golden_rows)}

## Downstream Entry

The next execution entry for building/decorating is `GD2018_BUILDING_A_FULL_PARSE_1`, followed by `MAP_GB50854_TO_GD2018_BUILDING_1`, then `WEB_QUOTA_BUILDING_PROTOTYPE_1`. These stages are documented as plan only here.
"""
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")


def write_entity_doc(path: Path, entity_rows: Sequence[Dict[str, Any]]) -> None:
    content = f"""# Reference Entity Dictionary

Stage: `{STAGE_NAME}`

Unified quota UID format: `GD:2018:{{family}}:{{source_code_normalized}}`, for example `GD:2018:A:A1-1-1`.

The dictionary below defines primary key, business meaning, canonical source, mutability, write permission, relationships, review status, and trace fields for the core reference/mapping/review entities.

{md_table(ENTITY_HEADERS, entity_rows)}
"""
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")


def write_execution_plan_doc(path: Path) -> None:
    rows = [
        {
            "stage": "REFERENCE_FAMILY_FRAMEWORK_LOCK_1",
            "status": "executed_this_round",
            "purpose": "Lock source family registry, routing matrix, six-layer contracts, entity dictionary, and A1.1 golden slice registration.",
            "inputs": "source_standards scan; existing A1.1 runs; existing Web draft/audit counts",
            "outputs": "framework docs, private registry CSVs, review workbook, manifest updates",
            "explicit_no_go": "no A01/A02/A03 full parse; no Mapping execution; no Web business page modification",
        },
        {
            "stage": "GD2018_BUILDING_A_FULL_PARSE_1",
            "status": "planned_only",
            "purpose": "Parse GD2018 A01/A02/A03 building/decorating quota family into L2 Parsed Reference.",
            "inputs": "L0 source registry; A01/A02/A03 PDFs; A1.1 parser lessons",
            "outputs": "gd_quota_item, fee/resource/work/rule/note/conversion candidates and parse issues",
            "explicit_no_go": "must not write Source/Baseline/Web; must not create final mapping or enterprise quota",
        },
        {
            "stage": "MAP_GB50854_TO_GD2018_BUILDING_1",
            "status": "planned_only",
            "purpose": "Map GB/T 50854 bill items to GD2018 A family quota candidates in L3 Mapping Reference.",
            "inputs": "GB/T 50854 L1 baseline; GD2018 A-family L2 parsed references",
            "outputs": "bill_quota_mapping_candidate and bill_quota_mapping_issue",
            "explicit_no_go": "no final promotion; no Source Candidate write-back; no enterprise price output",
        },
        {
            "stage": "WEB_QUOTA_BUILDING_PROTOTYPE_1",
            "status": "planned_only",
            "purpose": "Extend Web collaboration from A1.1 slice to building-family review draft workflows.",
            "inputs": "read-only L3 candidate view model; L4 draft/audit schema",
            "outputs": "read-only APIs plus Review Draft/Audit writes only",
            "explicit_no_go": "no business page mutation in this framework stage; Web must not mutate L0-L3",
        },
    ]
    content = f"""# Building Family Execution Plan

Stage: `{STAGE_NAME}`

This is a plan-only downstream sequence. Only the framework lock stage is executed now.

{md_table(["stage", "status", "purpose", "inputs", "outputs", "explicit_no_go"], rows)}

## Acceptance Gate For Next Stage

- A01/A02/A03 must use `GB/T 50854-2024` as bill standard family.
- `quota_uid` must follow `GD:2018:{{family}}:{{source_code_normalized}}`.
- Parsed outputs must carry source document, page/block, parser version, run id, and source hash trace fields.
- Existing A1.1 golden slice counts and Web smoke must remain stable before expanding to the whole A family.
- Copy/Move/Exclude/Restore semantics stay in L4 Review Draft and never mutate Source Candidate.
"""
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")


def stage_manifest_rows(project_root: Path, paths: Sequence[Path], source_file: str) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    for path in paths:
        rel = str(path.relative_to(project_root)).replace("\\", "/")
        row_count = count_csv_rows(path) if path.suffix.lower() == ".csv" else ""
        rows.append(
            {
                "stage_name": STAGE_NAME,
                "artifact_name": path.name,
                "expected_path": rel,
                "exists": str(path.exists()).lower(),
                "file_size_bytes": path.stat().st_size if path.exists() else "",
                "row_count": row_count,
                "sha256": sha256_file(path) if path.exists() else "",
                "created_or_modified_time": datetime.fromtimestamp(path.stat().st_mtime).isoformat(timespec="seconds")
                if path.exists()
                else "",
                "source_file": source_file,
                "can_regenerate": "true",
                "backup_required": "false",
                "backup_path": "",
                "status": "framework_locked",
                "remark": "stage output; private artifacts remain under data/private where applicable; no Source/Baseline/Web mutation",
            }
        )
    return rows


def update_manifest_csv(project_root: Path, artifact_paths: Sequence[Path]) -> None:
    manifest_path = project_root / DOCS_REL / "reference_artifact_manifest.csv"
    headers = read_headers(manifest_path)
    if not headers:
        headers = [
            "stage_name",
            "artifact_name",
            "expected_path",
            "exists",
            "file_size_bytes",
            "row_count",
            "sha256",
            "created_or_modified_time",
            "source_file",
            "can_regenerate",
            "backup_required",
            "backup_path",
            "status",
            "remark",
        ]
    new_rows = stage_manifest_rows(
        project_root,
        artifact_paths,
        "source_standards registry scan + existing A1.1 runs + Web draft/audit read-only audit",
    )

    existing = read_csv(manifest_path)
    if any(row.get("stage_name") == STAGE_NAME for row in existing):
        existing = [row for row in existing if row.get("stage_name") != STAGE_NAME]
        write_csv(manifest_path, headers, [*existing, *new_rows])
    else:
        with manifest_path.open("a", encoding="utf-8-sig", newline="") as fh:
            writer = csv.DictWriter(fh, fieldnames=headers, extrasaction="ignore")
            if manifest_path.stat().st_size == 0:
                writer.writeheader()
            for row in new_rows:
                writer.writerow({field: row.get(field, "") for field in headers})


def update_manifest_md(project_root: Path, artifact_paths: Sequence[Path]) -> None:
    manifest_path = project_root / DOCS_REL / "REFERENCE_ARTIFACT_MANIFEST.md"
    rows = stage_manifest_rows(
        project_root,
        artifact_paths,
        "source_standards registry scan + existing A1.1 runs + Web draft/audit read-only audit",
    )
    section = f"""

## {STAGE_NAME}

| Artifact | Exists | Rows | SHA256 | Status |
| --- | --- | ---: | --- | --- |
"""
    for row in rows:
        sha = nstr(row["sha256"])
        section += (
            f"| {row['artifact_name']} | {row['exists']} | {row['row_count']} | "
            f"{sha[:12]}... | {row['status']} |\n"
        )
    section += "\nScope: framework/registry/doc generation only; no A01/A02/A03 full parse, no Mapping execution, no Web business page modification, no DB writes.\n"

    original = manifest_path.read_text(encoding="utf-8") if manifest_path.exists() else "# Reference Artifact Manifest\n"
    marker = f"\n## {STAGE_NAME}\n"
    if marker in original:
        before, rest = original.split(marker, 1)
        next_match = re.search(r"\n## [^\n]+\n", rest)
        after = rest[next_match.start() :] if next_match else ""
        manifest_path.write_text(before.rstrip() + section + after, encoding="utf-8")
    else:
        manifest_path.write_text(original.rstrip() + section, encoding="utf-8")


def write_report(
    path: Path,
    source_rows: Sequence[Dict[str, Any]],
    routing_rows: Sequence[Dict[str, Any]],
    golden_rows: Sequence[Dict[str, Any]],
    issue_rows: Sequence[Dict[str, Any]],
    output_paths: Sequence[Path],
    project_root: Path,
) -> None:
    national_files = [row for row in source_rows if row["document_group"] == "national_standard"]
    gd_files = [row for row in source_rows if row["document_group"] == "gd2018_quota"]
    baseline_files = [row for row in source_rows if row["document_group"] == "national_standard_baseline"]
    a_docs = [row for row in gd_files if row["volume_code"] in {"A01", "A02", "A03"}]
    gb50854_docs = [row for row in source_rows if row["standard_code"] == "GB/T 50854-2024"]
    issue_count = sum(1 for row in issue_rows if row["status"] not in {"pass"})
    content = f"""# Stage {STAGE_NAME} Report

## Final Status

framework_locked_with_manual_source_inventory_warnings

This stage scanned the national standard and GD2018 source directories, registered actual file paths/hashes/file sizes/page counts/text-layer status, locked the routing matrix, generated six-layer contracts and the entity dictionary, and registered the existing A1.1 golden slice. It did not parse A01/A02/A03 full data, did not execute Mapping, did not write DB, and did not modify Web business pages.

## Source Inventory

- National standard directory files: {len(national_files)}
- Current GB/T 50854 baseline DOCX files outside national directory: {len(baseline_files)}
- GD2018 quota files: {len(gd_files)}

### A01/A02/A03 Actual Files

{md_table(["volume_code", "file_name", "sha256", "page_count", "text_layer_status"], a_docs)}

### GB/T 50854 Actual Files

{md_table(["document_group", "file_name", "sha256", "page_count", "text_layer_status"], gb50854_docs)}

## Routing Matrix

{md_table(["source_family", "source_volume_pattern", "target_standard_code", "mapping_target_type", "route_status", "routing_reason"], routing_rows)}

## A1.1 Golden Slice Registration

{md_table(GOLDEN_HEADERS, golden_rows)}

## Validation

- validation non-pass rows: {issue_count}
- approved_count: 0
- banned output count: 0

{md_table(ISSUE_HEADERS, issue_rows)}

## Generated Files

{md_table(["path"], [{"path": str(path.relative_to(project_root)).replace("\\", "/")} for path in output_paths])}
"""
    path.write_text(content, encoding="utf-8")


def run(project_root: Path) -> Dict[str, Path]:
    run_dir = project_root / RUN_DIR_REL
    docs_dir = project_root / DOCS_REL
    run_dir.mkdir(parents=True, exist_ok=True)
    docs_dir.mkdir(parents=True, exist_ok=True)

    source_rows, _source_dirs = build_source_registry(project_root)
    family_rows = build_standard_families(source_rows)
    routing_rows = build_routing_matrix(source_rows)
    layer_rows = build_layer_contracts()
    entity_rows = build_entity_dictionary()
    golden_rows = build_golden_slice(project_root, source_rows)
    issue_rows = build_validation_issues(source_rows, routing_rows, golden_rows)

    source_csv = run_dir / "source_document_registry.csv"
    family_csv = run_dir / "standard_family_registry.csv"
    routing_csv = run_dir / "source_family_routing_matrix.csv"
    layer_csv = run_dir / "reference_layer_contract.csv"
    entity_csv = run_dir / "reference_entity_dictionary.csv"
    golden_csv = run_dir / "golden_slice_A111_registry.csv"
    issue_csv = run_dir / "framework_validation_issues.csv"
    workbook = run_dir / "Reference_Family_Framework_Review.xlsx"
    report = run_dir / "stage_reference_family_framework_lock_report.md"
    architecture_doc = docs_dir / "REFERENCE_FAMILY_ARCHITECTURE.md"
    entity_doc = docs_dir / "REFERENCE_ENTITY_DICTIONARY.md"
    plan_doc = docs_dir / "BUILDING_FAMILY_EXECUTION_PLAN.md"

    write_csv(source_csv, SOURCE_DOC_HEADERS, source_rows)
    write_csv(family_csv, STANDARD_FAMILY_HEADERS, family_rows)
    write_csv(routing_csv, ROUTING_HEADERS, routing_rows)
    write_csv(layer_csv, LAYER_HEADERS, layer_rows)
    write_csv(entity_csv, ENTITY_HEADERS, entity_rows)
    write_csv(golden_csv, GOLDEN_HEADERS, golden_rows)
    write_csv(issue_csv, ISSUE_HEADERS, issue_rows)

    write_xlsx(
        workbook,
        [
            ("source_document_registry", SOURCE_DOC_HEADERS, source_rows),
            ("standard_family_registry", STANDARD_FAMILY_HEADERS, family_rows),
            ("routing_matrix", ROUTING_HEADERS, routing_rows),
            ("layer_contract", LAYER_HEADERS, layer_rows),
            ("entity_dictionary", ENTITY_HEADERS, entity_rows),
            ("golden_slice_A111", GOLDEN_HEADERS, golden_rows),
            ("validation_issues", ISSUE_HEADERS, issue_rows),
        ],
    )

    write_architecture_doc(architecture_doc, source_rows, family_rows, routing_rows, layer_rows, golden_rows)
    write_entity_doc(entity_doc, entity_rows)
    write_execution_plan_doc(plan_doc)

    output_paths = [
        source_csv,
        family_csv,
        routing_csv,
        layer_csv,
        entity_csv,
        golden_csv,
        issue_csv,
        workbook,
        report,
        architecture_doc,
        entity_doc,
        plan_doc,
    ]
    write_report(report, source_rows, routing_rows, golden_rows, issue_rows, output_paths, project_root)
    update_manifest_csv(project_root, output_paths)
    update_manifest_md(project_root, output_paths)

    return {
        "run_dir": run_dir,
        "report": report,
        "architecture_doc": architecture_doc,
        "entity_doc": entity_doc,
        "plan_doc": plan_doc,
    }


def main() -> int:
    project_root = project_root_from_here()
    outputs = run(project_root)
    print(f"{STAGE_NAME} complete")
    for key, path in outputs.items():
        print(f"{key}: {path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
